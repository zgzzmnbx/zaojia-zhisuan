from __future__ import annotations

from datetime import date
from pathlib import Path
from typing import Any

from docx import Document
from docx.shared import Pt
from docx.text.paragraph import Paragraph
from docx.oxml.ns import qn
from docx.oxml import OxmlElement
from openpyxl import load_workbook

from .experience_warning import build_warning_report_lines
from .formula_resolver import WorkbookFormulaResolver
from .paths import DEFAULT_REPORT_TEMPLATE_PATH
from .schemas import FillSummary

REPORT_TEMPLATE_PATH = DEFAULT_REPORT_TEMPLATE_PATH


def write_report(
    output_path: str | Path,
    input_name: str,
    summary: FillSummary,
    output_excel_path: str | Path | None = None,
    input_excel_path: str | Path | None = None,
    report_date: date | None = None,
    report_template_path: str | Path | None = None,
) -> Path:
    report_day = report_date or date.today()
    output_path = _dated_output_path(Path(output_path), report_day)
    output_path.parent.mkdir(parents=True, exist_ok=True)
    markdown_path = output_path.with_suffix(".md")
    markdown_text = build_report_markdown(input_name, summary)
    markdown_path.write_text(markdown_text, encoding="utf-8")

    template_path = _resolve_template_path(report_day, report_template_path)
    document = Document(template_path) if template_path.exists() else Document()
    style = document.styles["Normal"]
    style.font.name = "Microsoft YaHei"
    style.font.size = Pt(10.5)

    fee_summary = _extract_fee_summary(output_excel_path) if output_excel_path else {}
    if input_excel_path and _has_missing_fee_summary(fee_summary):
        fallback_summary = _extract_fee_summary(input_excel_path)
        fee_summary = {
            key: fee_summary.get(key) or fallback_summary.get(key, "")
            for key in {"project_name", "total_with_tax", "total_without_tax", "vat"}
        }
    _replace_doc_placeholders(document, _build_placeholder_values(report_day, fee_summary))

    target = find_paragraph(document, "【匹配报告】")
    if target:
        set_paragraph_text(target, "造价智算匹配报告", kind="subheading")
        anchor = target
        for line in build_price_log_lines(input_name, summary):
            anchor = insert_after(anchor, line, kind="body")
    else:
        target = find_paragraph(document, "五、其他需要注意的事项")
        if target:
            anchor = insert_after(target, "造价智算匹配报告", kind="subheading")
            for line in build_price_log_lines(input_name, summary):
                anchor = insert_after(anchor, line, kind="body")
        else:
            document.add_heading("造价智算处理报告", level=1)
            target = document.add_heading("五、其他需要注意的事项", level=1)
            anchor = insert_after(target, "造价智算匹配报告", kind="subheading")
            for line in build_price_log_lines(input_name, summary):
                paragraph = document.add_paragraph(line)
                apply_official_paragraph_format(paragraph, kind="body")

    document.save(output_path)
    return output_path


def append_risk_report(output_path: str | Path, risk_text: str) -> Path:
    output_path = Path(output_path)
    document = Document(output_path)
    target = find_paragraph(document, "五、其他需要注意的事项")
    if target:
        anchor = insert_after(target, "大模型风险提示", kind="subheading")
        for line in risk_text.splitlines():
            if line.strip():
                anchor = insert_after(anchor, line.strip(), kind="body")
    else:
        document.add_heading("大模型风险提示", level=2)
        for line in risk_text.splitlines():
            if line.strip():
                paragraph = document.add_paragraph(line.strip())
                apply_official_paragraph_format(paragraph, kind="body")
    document.save(output_path)
    return output_path


def build_report_markdown(input_name: str, summary: FillSummary) -> str:
    lines = [
        "# 造价智算处理报告",
        "",
        f"- 输入文件：{input_name}",
        f"- 总数据行数：{summary.total_data_rows}",
        f"- 价格列：{summary.price_column}",
        f"- 成功回填行数：{summary.filled_rows}",
        f"- 精确匹配行数：{summary.matched_rows}",
        f"- 原有价格保留行数：{summary.unchanged_rows}",
        f"- 待复核行数：{summary.review_rows}",
        f"- 冲突行数：{summary.conflict_rows}",
        f"- 实物工作费调整系数第一层命中：{summary.physical_matched_rows}",
        f"- 实物工作费调整系数第二层经验：{summary.physical_experience_rows}",
        f"- 实物工作费调整系数待复核：{summary.physical_review_rows}",
        f"- 技术工作费调整系数第一层命中：{summary.technical_matched_rows}",
        f"- 技术工作费调整系数第二层经验：{summary.technical_experience_rows}",
        f"- 技术工作费调整系数待复核：{summary.technical_review_rows}",
        "",
        "## 经验池预警",
        "",
    ]
    lines.extend(f"- {line}" for line in build_warning_report_lines(summary.warning_summary, summary.warning_details, limit=10))
    lines.extend([
        "",
        "## 价格识别日志",
        "",
    ])
    lines.extend(f"- {line}" for line in summary.price_logs[:200])
    return "\n".join(lines) + "\n"


def build_price_log_lines(input_name: str, summary: FillSummary) -> list[str]:
    lines = [
        f"输入文件：{input_name}",
        (
            f"处理概况：共识别有效输入 {summary.total_data_rows} 行，"
            f"基价/单价匹配 {summary.matched_rows} 行，保留原有价格 {summary.unchanged_rows} 行，"
            f"待复核 {summary.review_rows} 行。"
        ),
        (
            f"实物工作费调整系数：第一层命中 {summary.physical_matched_rows} 行，"
            f"第二层经验 {summary.physical_experience_rows} 行，待复核 {summary.physical_review_rows} 行。"
        ),
        (
            f"技术工作费调整系数：第一层命中 {summary.technical_matched_rows} 行，"
            f"第二层经验 {summary.technical_experience_rows} 行，待复核 {summary.technical_review_rows} 行。"
        ),
        "说明：第一层为标准规则命中；第二层为单价知识库经验提示，适合辅助复核，后续仍建议结合标准依据确认。",
        "经验池预警：",
    ]
    lines.extend(build_warning_report_lines(summary.warning_summary, summary.warning_details, limit=8))
    if summary.review_details:
        lines.append("待复核提示：以下列出部分待复核行供人工抽查。")
        for row in summary.review_details[:10]:
            element = "，".join(str(value) for value in row.values.values() if str(value or "").strip())
            lines.append(f"第 {row.excel_row} 行：{row.message}；要素：{element}")
        if len(summary.review_details) > 10:
            lines.append(f"其余 {len(summary.review_details) - 10} 行待复核明细请查看输出 Excel 的逐行匹配报告列。")
    else:
        lines.append("待复核提示：本次未发现需要人工复核的基价/单价行，仍建议抽查第二层经验提示系数与项目实际条件是否一致。")
    return lines


def find_paragraph(document: Document, text: str) -> Paragraph | None:
    for paragraph in document.paragraphs:
        if text in paragraph.text:
            return paragraph
    return None


def insert_after(paragraph: Paragraph, text: str, kind: str = "body") -> Paragraph:
    new_p = OxmlElement("w:p")
    paragraph._p.addnext(new_p)
    new_paragraph = Paragraph(new_p, paragraph._parent)
    new_paragraph.add_run(text)
    apply_official_paragraph_format(new_paragraph, kind=kind)
    return new_paragraph


def _resolve_template_path(report_day: date, report_template_path: str | Path | None = None) -> Path:
    template_path = Path(report_template_path) if report_template_path else REPORT_TEMPLATE_PATH
    dated = template_path.with_name(
        template_path.name.replace("yyyy-mm-dd", report_day.isoformat())
    )
    return dated if dated.exists() else template_path


def _dated_output_path(output_path: Path, report_day: date) -> Path:
    return output_path.with_name(output_path.name.replace("yyyy-mm-dd", report_day.isoformat()))


def _build_placeholder_values(report_day: date, fee_summary: dict[str, str]) -> dict[str, str]:
    return {
        "yyyy-mm-dd": report_day.isoformat(),
        "【yyyy】": f"{report_day:%Y}",
        "【mm】": f"{report_day:%m}",
        "【dd】": f"{report_day:%d}",
        "【项目名称】": fee_summary.get("project_name", ""),
        "【费用汇总-合计（不含税）】": fee_summary.get("total_without_tax", ""),
        "【费用汇总-增值税】": fee_summary.get("vat", ""),
        "【采购计划金额】": fee_summary.get("total_without_tax", ""),
    }


def _replace_doc_placeholders(document: Document, values: dict[str, str]) -> None:
    for paragraph in document.paragraphs:
        _replace_in_paragraph(paragraph, values)
    for table in document.tables:
        for row in table.rows:
            for cell in row.cells:
                for paragraph in cell.paragraphs:
                    _replace_in_paragraph(paragraph, values)


def _replace_in_paragraph(paragraph: Paragraph, values: dict[str, str]) -> None:
    full_text = paragraph.text
    if not any(key in full_text for key in values):
        return
    for run in paragraph.runs:
        for key, value in values.items():
            if key in run.text:
                run.text = run.text.replace(key, value)
    if not any(key in paragraph.text for key in values):
        return
    replaced = paragraph.text
    for key, value in values.items():
        replaced = replaced.replace(key, value)
    if paragraph.runs:
        paragraph.runs[0].text = replaced
        for run in paragraph.runs[1:]:
            run.text = ""
    else:
        paragraph.add_run(replaced)


def _extract_fee_summary(output_excel_path: str | Path) -> dict[str, str]:
    path = Path(output_excel_path)
    if not path.exists():
        return {}
    try:
        resolver = WorkbookFormulaResolver(path)
    except Exception:
        resolver = None
    workbook = resolver.value_workbook if resolver is not None else load_workbook(path, data_only=True, read_only=True)
    try:
        if "费用汇总" not in workbook.sheetnames:
            return {}
        sheet = workbook["费用汇总"]
        return {
            "project_name": _find_project_name(sheet),
            "total_with_tax": _format_report_value(_find_label_value(sheet, "合计（含税）", resolver)),
            "total_without_tax": _format_report_value(_find_label_value(sheet, "合计（不含税）", resolver)),
            "vat": _format_report_value(_find_label_value(sheet, "增值税", resolver)),
        }
    finally:
        if resolver is not None:
            resolver.close()
        else:
            workbook.close()


def _find_project_name(sheet: Any) -> str:
    for row in sheet.iter_rows(values_only=True):
        values = list(row)
        for index, value in enumerate(values):
            text = str(value or "").strip()
            if "项目名称" not in text:
                continue
            parts = text.replace("：", ":").split(":", 1)
            if len(parts) == 2 and parts[1].strip():
                return parts[1].strip()
            for next_value in values[index + 1 :]:
                if str(next_value or "").strip():
                    return str(next_value).strip()
    return ""


def _find_label_value(sheet: Any, label: str, resolver: WorkbookFormulaResolver | None = None) -> Any:
    compact_label = label.replace(" ", "")
    for row_index, row in enumerate(sheet.iter_rows(values_only=True), start=1):
        values = list(row)
        for index, value in enumerate(values):
            if compact_label not in str(value or "").replace(" ", ""):
                continue
            for offset, next_value in enumerate(values[index + 1 :], start=index + 2):
                if next_value in (None, "") and resolver is not None:
                    next_value = resolver.cell_value(sheet.title, row_index, offset)
                if next_value not in (None, ""):
                    return next_value
    return None


def _format_report_value(value: Any) -> str:
    if value is None:
        return ""
    if isinstance(value, (int, float)) and not isinstance(value, bool):
        return f"{value:.2f}".rstrip("0").rstrip(".")
    return str(value).strip()


def _has_missing_fee_summary(values: dict[str, str]) -> bool:
    return any(not values.get(key) for key in ["project_name", "total_without_tax", "vat"])


def set_paragraph_text(paragraph: Paragraph, text: str, kind: str = "body") -> None:
    for run in paragraph.runs:
        run.text = ""
    paragraph.add_run(text)
    apply_official_paragraph_format(paragraph, kind=kind)


def apply_official_paragraph_format(paragraph: Paragraph, kind: str = "body") -> None:
    paragraph_format = paragraph.paragraph_format
    paragraph_format.line_spacing = Pt(28)
    paragraph_format.space_before = Pt(0)
    paragraph_format.space_after = Pt(0)
    paragraph_format.first_line_indent = Pt(28)
    font_name = "宋体"
    bold = False
    if kind == "subheading":
        font_name = "方正黑体简体"
        bold = True
    for run in paragraph.runs:
        run.font.name = font_name
        run._element.rPr.rFonts.set(qn("w:eastAsia"), font_name)
        run.font.size = Pt(14)
        run.bold = bold
