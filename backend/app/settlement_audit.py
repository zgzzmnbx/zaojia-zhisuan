from __future__ import annotations

import json
import math
import re
from dataclasses import asdict, dataclass
from datetime import datetime
from pathlib import Path
from typing import Any, Iterable, Literal

from docx import Document
from docx.enum.section import WD_ORIENT
from docx.enum.table import WD_CELL_VERTICAL_ALIGNMENT
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml import OxmlElement
from docx.oxml.ns import qn
from docx.shared import Cm, Pt, RGBColor
from openpyxl import load_workbook
from openpyxl.comments import Comment
from openpyxl.styles import Alignment, Border, Font, PatternFill, Side
from openpyxl.utils import get_column_letter
from openpyxl.worksheet.worksheet import Worksheet


RULE_VERSION = "1.0.0"
SUPPORTED_EXTENSION = ".xlsx"
AUDIT_RESULT_SHEET = "审核结果"
AUDIT_COMMENT = "造价智算辅助审核建议，需人工确认。"
AUDIT_SHEET_FONT_NAME = "等线"
AUDIT_SHEET_BODY_FONT_SIZE = 9

BLUE = "2563EB"
NAVY = "163B65"
TEXT = "172033"
MUTED = "64748B"
BORDER = "D8E1EC"
PALE_BLUE = "EAF2FF"
PALE_GREEN = "EAF7EF"
PALE_YELLOW = "FFF4CC"
PALE_RED = "FDECEC"
WHITE = "FFFFFF"


class SettlementAuditError(ValueError):
    """Raised when a workbook cannot be safely audited."""


@dataclass(frozen=True)
class SheetProfile:
    kind: Literal["measure", "survey", "other"]
    descriptor_columns: tuple[int, ...]
    start_row: int
    quantity_column: int | None
    base_column: int | None
    physical_column: int | None
    technical_column: int | None
    total_column: int
    audit_columns: tuple[int, ...]
    reason_column: int


@dataclass(frozen=True)
class AuditRisk:
    id: str
    severity: Literal["high", "medium", "low"]
    category: str
    rule_id: str
    title: str
    sheet: str
    row: int | None
    coordinate: str
    current_value: Any
    suggested_value: Any
    basis: str
    action: str
    auto_adjusted: bool

    def to_dict(self) -> dict[str, Any]:
        return asdict(self)


MEASURE_PROFILE = SheetProfile(
    kind="measure",
    descriptor_columns=(2, 3, 4, 5, 6),
    start_row=5,
    quantity_column=7,
    base_column=8,
    physical_column=9,
    technical_column=11,
    total_column=13,
    audit_columns=(23, 24, 25, 26, 27, 28),
    reason_column=29,
)

SURVEY_PROFILE = SheetProfile(
    kind="survey",
    descriptor_columns=(2, 3, 4, 5, 6, 7),
    start_row=5,
    quantity_column=8,
    base_column=9,
    physical_column=10,
    technical_column=13,
    total_column=15,
    audit_columns=(25, 26, 27, 28, 29, 30),
    reason_column=31,
)

OTHER_PROFILE = SheetProfile(
    kind="other",
    descriptor_columns=(2,),
    start_row=4,
    quantity_column=None,
    base_column=None,
    physical_column=None,
    technical_column=None,
    total_column=12,
    audit_columns=tuple(range(15, 24)),
    reason_column=25,
)


MANUAL_CHECKLIST = (
    {
        "id": "MANUAL-001",
        "title": "合同与下浮口径",
        "detail": "核验合同、补充协议、框架协议及适用的下浮比例，确认结算阶段和合同边界。",
    },
    {
        "id": "MANUAL-002",
        "title": "工程量与成果一致性",
        "detail": "核验经确认的工作量、统计表、勘察测量成果和现场过程资料是否相互一致。",
    },
    {
        "id": "MANUAL-003",
        "title": "签章、日期与数据真实性",
        "detail": "核验签字、盖章、日期、数据来源及异常数据说明；系统不自动判断资料真实性。",
    },
    {
        "id": "MANUAL-004",
        "title": "其他费用支付证据",
        "detail": "逐项核验协议、发票、收据、转账或付款凭证，以及费用是否真实发生、是否重复计取。",
    },
)


def _clean_text(value: Any) -> str:
    if value is None:
        return ""
    return re.sub(r"\s+", "", str(value)).replace("（", "(").replace("）", ")").lower()


def _display_value(value: Any) -> str:
    if value is None or value == "":
        return "空"
    if isinstance(value, float):
        if math.isfinite(value):
            return f"{value:,.4f}".rstrip("0").rstrip(".")
        return str(value)
    return str(value)


def _numeric(value: Any) -> float | None:
    if isinstance(value, bool):
        return None
    if isinstance(value, (int, float)) and math.isfinite(float(value)):
        return float(value)
    if isinstance(value, str):
        cleaned = value.strip().replace(",", "")
        if not cleaned or cleaned.startswith("="):
            return None
        if cleaned.endswith("%"):
            try:
                return float(cleaned[:-1]) / 100
            except ValueError:
                return None
        try:
            number = float(cleaned)
        except ValueError:
            return None
        return number if math.isfinite(number) else None
    return None


def _different(left: float | None, right: float | None, tolerance: float = 0.01) -> bool:
    if left is None or right is None:
        return left != right
    return abs(left - right) > tolerance


def _is_formula(value: Any) -> bool:
    return isinstance(value, str) and value.startswith("=")


def _effective_value(ws: Worksheet, row: int, column: int) -> Any:
    value = ws.cell(row=row, column=column).value
    if value not in (None, ""):
        return value
    for merged_range in ws.merged_cells.ranges:
        if not (
            merged_range.min_row <= row <= merged_range.max_row
            and merged_range.min_col <= column <= merged_range.max_col
        ):
            continue
        is_vertical = merged_range.min_col == merged_range.max_col
        is_top_left = row == merged_range.min_row and column == merged_range.min_col
        if is_vertical or is_top_left:
            return ws.cell(merged_range.min_row, merged_range.min_col).value
        return None
    return None


def _cell_number(formula_ws: Worksheet, values_ws: Worksheet, row: int, column: int) -> float | None:
    cached = _numeric(_effective_value(values_ws, row, column))
    if cached is not None:
        return cached
    return _numeric(_effective_value(formula_ws, row, column))


def _descriptor(ws: Worksheet, profile: SheetProfile, row: int) -> str:
    return " / ".join(
        part
        for part in (_display_value(_effective_value(ws, row, column)) for column in profile.descriptor_columns)
        if part != "空"
    )


def _descriptor_key(ws: Worksheet, profile: SheetProfile, row: int) -> str:
    return "|".join(_clean_text(_effective_value(ws, row, column)) for column in profile.descriptor_columns)


def _sheet_profile(ws: Worksheet) -> SheetProfile | None:
    marker = _clean_text(ws["A1"].value)
    if "附表1测量费用测算表" in marker:
        return MEASURE_PROFILE
    if "附表2勘察费用测算表" in marker:
        return SURVEY_PROFILE
    if "附表3工程勘察其他费用统计" in marker:
        return OTHER_PROFILE
    return None


def _safe_filename_stem(name: str) -> str:
    stem = Path(name).stem.strip()
    stem = re.sub(r'[<>:"/\\|?*\x00-\x1f]', "-", stem)
    stem = re.sub(r"\s+", " ", stem).strip(" .")
    return stem[:100] or "结算审核"


def _short_sheet_label(name: str) -> str:
    cleaned = name.strip()
    if cleaned.startswith("表1"):
        return "表1 测量费"
    if cleaned.startswith("表2"):
        return "表2-1 勘察费"
    if cleaned.startswith("表3"):
        return "表3 其他费用"
    return cleaned[:20]


def _risk_fill(severity: str, auto_adjusted: bool) -> PatternFill:
    if auto_adjusted:
        return PatternFill("solid", fgColor=PALE_YELLOW)
    if severity == "high":
        return PatternFill("solid", fgColor=PALE_RED)
    return PatternFill("solid", fgColor=PALE_BLUE)


def _set_cell_comment(cell, text: str) -> None:
    cell.comment = Comment(f"{AUDIT_COMMENT}\n{text}", "造价智算")


def _set_docx_cell_shading(cell, fill: str) -> None:
    tc_pr = cell._tc.get_or_add_tcPr()
    shd = tc_pr.find(qn("w:shd"))
    if shd is None:
        shd = OxmlElement("w:shd")
        tc_pr.append(shd)
    shd.set(qn("w:fill"), fill)


def _set_docx_cell_text(cell, text: Any, *, bold: bool = False, color: str = TEXT) -> None:
    cell.text = ""
    paragraph = cell.paragraphs[0]
    paragraph.paragraph_format.space_after = Pt(0)
    run = paragraph.add_run(_display_value(text))
    run.bold = bold
    run.font.name = "等线"
    run.font.size = Pt(9)
    run.font.color.rgb = RGBColor.from_string(color)
    run._element.get_or_add_rPr().rFonts.set(qn("w:eastAsia"), "等线")
    cell.vertical_alignment = WD_CELL_VERTICAL_ALIGNMENT.CENTER


def _prevent_docx_row_split(row) -> None:
    tr_pr = row._tr.get_or_add_trPr()
    if tr_pr.find(qn("w:cantSplit")) is None:
        tr_pr.append(OxmlElement("w:cantSplit"))


def _repeat_docx_table_header(row) -> None:
    tr_pr = row._tr.get_or_add_trPr()
    if tr_pr.find(qn("w:tblHeader")) is None:
        header = OxmlElement("w:tblHeader")
        header.set(qn("w:val"), "true")
        tr_pr.append(header)


class SettlementAuditEngine:
    """Deterministic audit engine for the senior-provided settlement template."""

    def __init__(self, reference_template_path: str | Path):
        self.reference_template_path = Path(reference_template_path)
        if not self.reference_template_path.is_file():
            raise SettlementAuditError("结算审核参考模板不存在，已停止审核以避免使用不完整规则。")
        if self.reference_template_path.suffix.lower() != SUPPORTED_EXTENSION:
            raise SettlementAuditError("结算审核参考模板必须为 .xlsx 文件。")

    def review(
        self,
        input_path: str | Path,
        output_dir: str | Path,
        *,
        source_name: str | None = None,
        project_name: str | None = None,
    ) -> dict[str, Any]:
        input_path = Path(input_path)
        output_dir = Path(output_dir)
        if input_path.suffix.lower() != SUPPORTED_EXTENSION:
            raise SettlementAuditError("当前仅支持前辈统一模板及同结构的 .xlsx 文件，不支持旧版 .xls。")
        if not input_path.is_file():
            raise SettlementAuditError("待审核结算文件不存在。")

        try:
            workbook = load_workbook(input_path, data_only=False, keep_links=True)
            values_workbook = load_workbook(input_path, data_only=True, keep_links=True)
            reference_workbook = load_workbook(self.reference_template_path, data_only=False, keep_links=True)
            reference_values = load_workbook(self.reference_template_path, data_only=True, keep_links=True)
        except Exception as exc:
            raise SettlementAuditError(f"无法读取结算工作簿：{exc}") from exc

        profiles = self._validate_and_collect_profiles(workbook)
        reference_profiles = self._validate_and_collect_profiles(reference_workbook)
        reference_by_kind = {
            profile.kind: (reference_workbook[sheet_name], reference_values[sheet_name], profile)
            for sheet_name, profile in reference_profiles
        }

        risks: list[AuditRisk] = []
        audited_rows = 0
        passed_rows = 0
        reported_total = 0.0
        reviewed_total = 0.0
        sheet_counts = {"measure": 0, "survey": 0, "other": 0}
        sheet_stats: dict[str, dict[str, Any]] = {}

        for sheet_name, profile in profiles:
            sheet_counts[profile.kind] += 1
            ws = workbook[sheet_name]
            values_ws = values_workbook[sheet_name]
            reference_ws, reference_values_ws, reference_profile = reference_by_kind[profile.kind]
            if profile.kind == "other":
                stats = self._audit_other_sheet(ws, values_ws, profile, risks)
            else:
                stats = self._audit_detail_sheet(
                    ws,
                    values_ws,
                    profile,
                    reference_ws,
                    reference_values_ws,
                    reference_profile,
                    risks,
                )
            audited_rows += stats["audited_rows"]
            passed_rows += stats["passed_rows"]
            reported_total += stats["reported_total"]
            reviewed_total += stats["reviewed_total"]
            sheet_stats[sheet_name] = {
                "sheet": sheet_name,
                "audited_rows": stats["audited_rows"],
                "reported_total": round(stats["reported_total"], 2),
                "reviewed_total": round(stats["reviewed_total"], 2),
                "high": 0,
                "medium": 0,
                "low": 0,
                "total": 0,
            }

        self._audit_summary_parameters(workbook, values_workbook, reference_workbook, reference_values, risks)
        for risk in risks:
            row = sheet_stats.setdefault(
                risk.sheet,
                {"sheet": risk.sheet, "audited_rows": 0, "reported_total": 0.0, "reviewed_total": 0.0, "high": 0, "medium": 0, "low": 0, "total": 0},
            )
            row[risk.severity] += 1
            row["total"] += 1

        source_file_name = source_name or input_path.name
        resolved_project_name = project_name or self._project_name(workbook, source_file_name)
        structured_difference = reported_total - reviewed_total
        high_risk_count = sum(risk.severity == "high" for risk in risks)
        manual_review_count = len(MANUAL_CHECKLIST) + sum(not risk.auto_adjusted for risk in risks)

        summary = {
            "sheet_count": len(profiles),
            "sheet_counts": sheet_counts,
            "audited_rows": audited_rows,
            "passed_rows": passed_rows,
            "risk_count": len(risks),
            "high_risk_count": high_risk_count,
            "manual_review_count": manual_review_count,
            "reported_detail_total": round(reported_total, 2),
            "reviewed_detail_total": round(reviewed_total, 2),
            "suggested_difference": round(structured_difference, 2),
            "sheet_summaries": list(sheet_stats.values()),
        }
        created_at = datetime.now().astimezone().isoformat(timespec="seconds")
        result = {
            "module": "settlement-audit",
            "rule_version": RULE_VERSION,
            "created_at": created_at,
            "project_name": resolved_project_name,
            "source_file": source_file_name,
            "scope_note": "金额为结构化明细试算口径，未替代合同下浮、税费、资料真实性及最终结算审定。",
            "summary": summary,
            "risks": [risk.to_dict() for risk in risks],
            "manual_checklist": list(MANUAL_CHECKLIST),
            "sources": [
                "前辈指导：勘察测量最高投标限价及结算编审（2026年7月28日）",
                "勘察测量结算统一报价模板 v1.0",
            ],
        }

        output_dir.mkdir(parents=True, exist_ok=True)
        stem = _safe_filename_stem(source_file_name)
        reviewed_workbook_path = output_dir / f"【审核后】{stem}.xlsx"
        report_path = output_dir / f"【审核报告】{stem}.docx"
        result_path = output_dir / "审核结果.json"

        self._write_result_sheet(workbook, result)
        workbook.calculation.fullCalcOnLoad = True
        workbook.calculation.forceFullCalc = True
        workbook.calculation.calcMode = "auto"
        workbook.save(reviewed_workbook_path)
        self._write_report(report_path, result)

        result["artifacts"] = {
            "excel": reviewed_workbook_path.name,
            "report": report_path.name,
            "result": result_path.name,
        }
        result_path.write_text(json.dumps(result, ensure_ascii=False, indent=2), encoding="utf-8")

        for item in (workbook, values_workbook, reference_workbook, reference_values):
            item.close()
        return result

    def _validate_and_collect_profiles(self, workbook) -> list[tuple[str, SheetProfile]]:
        profiles: list[tuple[str, SheetProfile]] = []
        for ws in workbook.worksheets:
            if ws.title == AUDIT_RESULT_SHEET:
                continue
            profile = _sheet_profile(ws)
            if profile is not None:
                profiles.append((ws.title, profile))
        kinds = {profile.kind for _, profile in profiles}
        missing = {"measure", "survey", "other"} - kinds
        if missing:
            labels = {"measure": "测量费表", "survey": "勘察费表", "other": "其他费用表"}
            missing_text = "、".join(labels[item] for item in sorted(missing))
            raise SettlementAuditError(f"模板结构不完整，未识别到：{missing_text}。请使用统一结算报价模板。")
        return profiles

    def _reference_row(
        self,
        ws: Worksheet,
        profile: SheetProfile,
        row: int,
        reference_ws: Worksheet,
        reference_profile: SheetProfile,
    ) -> int | None:
        key = _descriptor_key(ws, profile, row)
        if not key.strip("|"):
            return None
        if row <= reference_ws.max_row and key == _descriptor_key(reference_ws, reference_profile, row):
            return row
        candidates = [
            candidate
            for candidate in range(reference_profile.start_row, reference_ws.max_row + 1)
            if key == _descriptor_key(reference_ws, reference_profile, candidate)
        ]
        return candidates[0] if len(candidates) == 1 else None

    def _add_risk(
        self,
        risks: list[AuditRisk],
        *,
        severity: Literal["high", "medium", "low"],
        category: str,
        rule_id: str,
        title: str,
        sheet: str,
        row: int | None,
        coordinate: str,
        current_value: Any,
        suggested_value: Any,
        basis: str,
        action: str,
        auto_adjusted: bool,
    ) -> AuditRisk:
        risk = AuditRisk(
            id=f"JS-{len(risks) + 1:04d}",
            severity=severity,
            category=category,
            rule_id=rule_id,
            title=title,
            sheet=sheet,
            row=row,
            coordinate=coordinate,
            current_value=current_value,
            suggested_value=suggested_value,
            basis=basis,
            action=action,
            auto_adjusted=auto_adjusted,
        )
        risks.append(risk)
        return risk

    def _audit_detail_sheet(
        self,
        ws: Worksheet,
        values_ws: Worksheet,
        profile: SheetProfile,
        reference_ws: Worksheet,
        reference_values_ws: Worksheet,
        reference_profile: SheetProfile,
        risks: list[AuditRisk],
    ) -> dict[str, float | int]:
        audited_rows = 0
        passed_rows = 0
        reported_total = 0.0
        reviewed_total = 0.0

        for row in range(profile.start_row, ws.max_row + 1):
            quantity = _cell_number(ws, values_ws, row, profile.quantity_column or 0)
            if quantity is None or quantity <= 0:
                continue

            audited_rows += 1
            row_risks_before = len(risks)
            descriptor = _descriptor(ws, profile, row)
            descriptor_clean = _clean_text(descriptor)
            reference_row = self._reference_row(ws, profile, row, reference_ws, reference_profile)

            base = _cell_number(ws, values_ws, row, profile.base_column or 0)
            physical = _cell_number(ws, values_ws, row, profile.physical_column or 0)
            technical = _cell_number(ws, values_ws, row, profile.technical_column or 0)
            reference_base = (
                _cell_number(reference_ws, reference_values_ws, reference_row, reference_profile.base_column or 0)
                if reference_row
                else None
            )
            reference_physical = (
                _cell_number(reference_ws, reference_values_ws, reference_row, reference_profile.physical_column or 0)
                if reference_row
                else None
            )
            reference_technical = (
                _cell_number(reference_ws, reference_values_ws, reference_row, reference_profile.technical_column or 0)
                if reference_row
                else None
            )
            if reference_row:
                parameter_pairs = (
                    ("base", profile.base_column, reference_profile.base_column, base, reference_base),
                    (
                        "physical",
                        profile.physical_column,
                        reference_profile.physical_column,
                        physical,
                        reference_physical,
                    ),
                    (
                        "technical",
                        profile.technical_column,
                        reference_profile.technical_column,
                        technical,
                        reference_technical,
                    ),
                )
                resolved_parameters: dict[str, float | None] = {}
                for name, column, reference_column, current, reference_value in parameter_pairs:
                    raw_formula = ws.cell(row, column or 1).value
                    reference_formula = reference_ws.cell(reference_row, reference_column or 1).value
                    if (
                        current is None
                        and reference_value is not None
                        and _is_formula(raw_formula)
                        and _clean_text(raw_formula) == _clean_text(reference_formula)
                    ):
                        resolved_parameters[name] = reference_value
                    else:
                        resolved_parameters[name] = current
                base = resolved_parameters["base"]
                physical = resolved_parameters["physical"]
                technical = resolved_parameters["technical"]

            suggested_base = base if base is not None else reference_base
            suggested_physical = physical if physical is not None else reference_physical
            suggested_technical = technical if technical is not None else reference_technical
            reasons: list[str] = []

            is_indoor_test = "室内试验" in descriptor_clean or "室内实验" in descriptor_clean
            is_aerial_or_route_map = any(
                keyword in descriptor_clean for keyword in ("航测", "航空摄影", "航空测量", "走向图", "线路走向")
            )
            is_deep_over_300 = bool(
                re.search(r"300[＜<]", descriptor_clean)
                or re.search(r"[＞>]300", descriptor_clean)
                or "大于300" in descriptor_clean
            )

            if reference_base is not None and _different(base, reference_base):
                self._add_risk(
                    risks,
                    severity="high",
                    category="参数",
                    rule_id="JS-001",
                    title="基价偏离统一模板",
                    sheet=ws.title,
                    row=row,
                    coordinate=f"{get_column_letter(profile.base_column or 1)}{row}",
                    current_value=base,
                    suggested_value=reference_base,
                    basis=f"统一报价模板同项参数：{descriptor}",
                    action="核对收费依据与模板项目；本次已在造价审核栏写入模板建议值。",
                    auto_adjusted=True,
                )
                suggested_base = reference_base
                reasons.append("基价按统一模板建议值调整")

            if reference_physical is not None and _different(physical, reference_physical):
                self._add_risk(
                    risks,
                    severity="high",
                    category="参数",
                    rule_id="JS-001",
                    title="实物工作费系数偏离统一模板",
                    sheet=ws.title,
                    row=row,
                    coordinate=f"{get_column_letter(profile.physical_column or 1)}{row}",
                    current_value=physical,
                    suggested_value=reference_physical,
                    basis=f"统一报价模板同项参数：{descriptor}",
                    action="核对复杂程度与系数来源；本次已在造价审核栏写入模板建议值。",
                    auto_adjusted=True,
                )
                suggested_physical = reference_physical
                reasons.append("实物工作费系数按统一模板建议值调整")

            if (
                not is_indoor_test
                and not is_aerial_or_route_map
                and reference_technical is not None
                and _different(technical, reference_technical)
            ):
                self._add_risk(
                    risks,
                    severity="high",
                    category="参数",
                    rule_id="JS-001",
                    title="技术工作费系数偏离统一模板",
                    sheet=ws.title,
                    row=row,
                    coordinate=f"{get_column_letter(profile.technical_column or 1)}{row}",
                    current_value=technical,
                    suggested_value=reference_technical,
                    basis=f"统一报价模板同项参数：{descriptor}",
                    action="核对技术工作等级与系数来源；本次已在造价审核栏写入模板建议值。",
                    auto_adjusted=True,
                )
                suggested_technical = reference_technical
                reasons.append("技术工作费系数按统一模板建议值调整")

            if is_deep_over_300 and reference_row and profile.base_column:
                reference_formula = reference_ws.cell(reference_row, profile.base_column).value
                if _is_formula(reference_formula) and re.search(r"\*\s*1[.]2", reference_formula, re.IGNORECASE):
                    previous_base = _cell_number(
                        reference_ws,
                        reference_values_ws,
                        max(reference_row - 1, reference_profile.start_row),
                        reference_profile.base_column or 0,
                    )
                    if previous_base is not None:
                        self._add_risk(
                            risks,
                            severity="high",
                            category="专项规则",
                            rule_id="JS-003",
                            title="深度大于 300m 部分存在额外 1.2 基价调整",
                            sheet=ws.title,
                            row=row,
                            coordinate=f"{get_column_letter(profile.base_column)}{row}",
                            current_value=base,
                            suggested_value=previous_base,
                            basis="前辈结算审核经验：钻孔深度大于 300m 部分不再以实物工作费基价乘 1.2。",
                            action="核对深度分段和原始钻探记录；本次按前一深度段基价形成审核建议。",
                            auto_adjusted=True,
                        )
                        suggested_base = previous_base
                        reasons.append("深度大于300m部分取消额外1.2基价调整")

            if is_indoor_test and _different(technical, 1.2, tolerance=0.0001):
                self._add_risk(
                    risks,
                    severity="high",
                    category="专项规则",
                    rule_id="JS-004",
                    title="室内试验技术工作费系数口径异常",
                    sheet=ws.title,
                    row=row,
                    coordinate=f"{get_column_letter(profile.technical_column or 1)}{row}",
                    current_value=technical,
                    suggested_value=1.2,
                    basis="前辈结算审核经验：室内试验技术工作费按 1.2 计取，不重复另收报告编制费 0.1。",
                    action="核对是否存在重复报告编制费；本次按 1.2 形成审核建议。",
                    auto_adjusted=True,
                )
                suggested_technical = 1.2
                reasons.append("室内试验技术工作费按1.2建议")

            if is_aerial_or_route_map and technical is not None and abs(technical) > 0.0001:
                self._add_risk(
                    risks,
                    severity="high",
                    category="专项规则",
                    rule_id="JS-005",
                    title="航测或线路走向图重复计取技术工作费",
                    sheet=ws.title,
                    row=row,
                    coordinate=f"{get_column_letter(profile.technical_column or 1)}{row}",
                    current_value=technical,
                    suggested_value=0,
                    basis="前辈结算审核经验：航测综合单价及线路走向图不再另计技术工作费。",
                    action="核对工作内容是否属于综合单价范围；本次技术工作费系数建议为 0。",
                    auto_adjusted=True,
                )
                suggested_technical = 0.0
                reasons.append("综合单价项目不再另计技术工作费")

            missing_fields = [
                label
                for label, value in (
                    ("基价", suggested_base),
                    ("实物工作费系数", suggested_physical),
                    ("技术工作费系数", suggested_technical),
                )
                if value is None
            ]
            if missing_fields:
                self._add_risk(
                    risks,
                    severity="high",
                    category="结构",
                    rule_id="JS-001",
                    title="计算参数缺失",
                    sheet=ws.title,
                    row=row,
                    coordinate=f"A{row}",
                    current_value="、".join(missing_fields),
                    suggested_value="待补充",
                    basis="统一报价模板要求活动明细具备工程量、基价和两个系数。",
                    action="补充参数及依据后重新审核；当前行未自动计算审核总价。",
                    auto_adjusted=False,
                )
                reasons.append("计算参数缺失，待人工补充")
                suggested_total = None
            else:
                suggested_total = (
                    quantity
                    * float(suggested_physical)
                    * float(suggested_base)
                    * (1 + float(suggested_technical))
                )

            reported_value = _cell_number(ws, values_ws, row, profile.total_column)
            current_computed = (
                quantity * physical * base * (1 + technical)
                if base is not None and physical is not None and technical is not None
                else None
            )
            if reported_value is None:
                reported_value = current_computed
            if reported_value is not None:
                reported_total += reported_value
            if suggested_total is not None:
                reviewed_total += suggested_total

            if (
                reported_value is not None
                and current_computed is not None
                and _different(reported_value, current_computed)
            ):
                self._add_risk(
                    risks,
                    severity="high",
                    category="算术",
                    rule_id="JS-002",
                    title="承包商明细金额与计算参数不一致",
                    sheet=ws.title,
                    row=row,
                    coordinate=f"{get_column_letter(profile.total_column)}{row}",
                    current_value=reported_value,
                    suggested_value=current_computed,
                    basis="统一模板计算结构：工程量 × 实物系数 × 基价 × (1 + 技术系数)。",
                    action="核对是否存在手工覆盖或公式错误；造价审核总价按审核参数重新计算。",
                    auto_adjusted=True,
                )
                reasons.append("承包商金额算术不一致")

            row_risks = risks[row_risks_before:]
            if not row_risks:
                passed_rows += 1
                reasons.append("结构化参数与算术校核通过，工程量及依据仍需人工确认")
            self._write_detail_audit(
                ws,
                profile,
                row,
                quantity,
                suggested_base,
                suggested_physical,
                suggested_technical,
                suggested_total,
                reasons,
                row_risks,
            )

        return {
            "audited_rows": audited_rows,
            "passed_rows": passed_rows,
            "reported_total": reported_total,
            "reviewed_total": reviewed_total,
        }

    def _write_detail_audit(
        self,
        ws: Worksheet,
        profile: SheetProfile,
        row: int,
        quantity: float,
        base: float | None,
        physical: float | None,
        technical: float | None,
        total: float | None,
        reasons: Iterable[str],
        row_risks: list[AuditRisk],
    ) -> None:
        if profile.kind == "measure":
            values = (0.0, quantity, physical, base, technical, total)
        else:
            values = (0.0, quantity, base, physical, technical, total)
        has_adjustment = any(risk.auto_adjusted for risk in row_risks)
        has_manual = any(not risk.auto_adjusted for risk in row_risks)
        fill = (
            PatternFill("solid", fgColor=PALE_YELLOW)
            if has_adjustment
            else PatternFill("solid", fgColor=PALE_RED if has_manual else PALE_GREEN)
        )
        reason_text = "；".join(dict.fromkeys(reasons))
        for column, value in zip(profile.audit_columns, values, strict=True):
            cell = ws.cell(row, column)
            cell.value = value
            cell.fill = fill
            cell.number_format = "#,##0.00"
            if has_adjustment or has_manual:
                _set_cell_comment(cell, reason_text)
        reason_cell = ws.cell(row, profile.reason_column)
        reason_cell.value = reason_text
        reason_cell.fill = fill
        reason_cell.alignment = Alignment(wrap_text=True, vertical="center")
        _set_cell_comment(reason_cell, reason_text)

    def _audit_other_sheet(
        self,
        ws: Worksheet,
        values_ws: Worksheet,
        profile: SheetProfile,
        risks: list[AuditRisk],
    ) -> dict[str, float | int]:
        audited_rows = 0
        passed_rows = 0
        reported_total = 0.0
        reviewed_total = 0.0

        for row in range(profile.start_row, ws.max_row + 1):
            amounts = [_cell_number(ws, values_ws, row, column) or 0.0 for column in range(3, 12)]
            calculated_total = sum(amounts)
            if abs(calculated_total) <= 0.0001:
                continue
            audited_rows += 1
            row_risks_before = len(risks)
            reported_value = _cell_number(ws, values_ws, row, profile.total_column)
            if reported_value is None:
                reported_value = calculated_total
            evidence = _effective_value(ws, row, 13)
            reasons: list[str] = []

            if _different(reported_value, calculated_total):
                self._add_risk(
                    risks,
                    severity="high",
                    category="算术",
                    rule_id="JS-006",
                    title="其他费用合计未覆盖全部费用列",
                    sheet=ws.title,
                    row=row,
                    coordinate=f"L{row}",
                    current_value=reported_value,
                    suggested_value=calculated_total,
                    basis="统一报价模板其他费用 C:K 全列合计；“其他”列也必须纳入。",
                    action="造价审核栏已按全部费用列重算，仍需逐项核验费用真实性。",
                    auto_adjusted=True,
                )
                reasons.append("其他费用按全部费用列重新合计")

            if not _clean_text(evidence):
                self._add_risk(
                    risks,
                    severity="high",
                    category="证据",
                    rule_id="JS-007",
                    title="其他费用缺少费用文件编号",
                    sheet=ws.title,
                    row=row,
                    coordinate=f"M{row}",
                    current_value="空",
                    suggested_value="补充协议、发票、收据或付款证据索引",
                    basis="前辈结算审核经验：其他费用必须结合合同、协议及实际付款资料核验。",
                    action="当前保留上报金额，不自动审减为 0；资料补齐并人工确认后再定案。",
                    auto_adjusted=False,
                )
                reasons.append("费用文件编号缺失，金额保留待人工核验")

            row_risks = risks[row_risks_before:]
            if not row_risks:
                passed_rows += 1
                reasons.append("合计与证据索引结构校核通过，费用真实性仍需人工确认")
            fill = (
                PatternFill("solid", fgColor=PALE_YELLOW)
                if any(risk.auto_adjusted for risk in row_risks)
                else PatternFill("solid", fgColor=PALE_RED if row_risks else PALE_GREEN)
            )
            for source_column, audit_column in zip(range(3, 12), profile.audit_columns, strict=True):
                cell = ws.cell(row, audit_column)
                cell.value = _cell_number(ws, values_ws, row, source_column) or 0.0
                cell.number_format = "#,##0.00"
                cell.fill = fill
            total_cell = ws.cell(row, 24)
            total_cell.value = calculated_total
            total_cell.number_format = "#,##0.00"
            total_cell.fill = fill
            reason_text = "；".join(dict.fromkeys(reasons))
            reason_cell = ws.cell(row, profile.reason_column)
            reason_cell.value = reason_text
            reason_cell.fill = fill
            reason_cell.alignment = Alignment(wrap_text=True, vertical="center")
            _set_cell_comment(reason_cell, reason_text)

            reported_total += reported_value
            reviewed_total += calculated_total

        return {
            "audited_rows": audited_rows,
            "passed_rows": passed_rows,
            "reported_total": reported_total,
            "reviewed_total": reviewed_total,
        }

    def _audit_summary_parameters(
        self,
        workbook,
        values_workbook,
        reference_workbook,
        reference_values,
        risks: list[AuditRisk],
    ) -> None:
        if "汇总表" not in workbook.sheetnames or "汇总表" not in reference_workbook.sheetnames:
            return
        ws = workbook["汇总表"]
        values_ws = values_workbook["汇总表"]
        reference_ws = reference_workbook["汇总表"]
        reference_values_ws = reference_values["汇总表"]
        for coordinate, label in (("M4", "2024 版其他费用系数"), ("N4", "勘察测量框架下浮"), ("O4", "航测框架下浮")):
            current = _cell_number(ws, values_ws, ws[coordinate].row, ws[coordinate].column)
            expected = _cell_number(
                reference_ws,
                reference_values_ws,
                reference_ws[coordinate].row,
                reference_ws[coordinate].column,
            )
            if expected is None or not _different(current, expected, tolerance=0.0001):
                continue
            self._add_risk(
                risks,
                severity="medium",
                category="合同参数",
                rule_id="JS-008",
                title=f"{label}偏离统一模板参考值",
                sheet=ws.title,
                row=ws[coordinate].row,
                coordinate=coordinate,
                current_value=current,
                suggested_value=expected,
                basis="统一报价模板参考参数及前辈结算审核经验。",
                action="核对本项目合同和框架协议后人工确认；系统不自动改写合同参数。",
                auto_adjusted=False,
            )
            ws[coordinate].fill = PatternFill("solid", fgColor=PALE_RED)
            _set_cell_comment(ws[coordinate], "参数偏离参考模板，必须结合本项目合同人工确认。")

    def _project_name(self, workbook, source_name: str) -> str:
        if "汇总表" in workbook.sheetnames:
            title = workbook["汇总表"]["A1"].value
            if title:
                first_line = str(title).splitlines()[0].strip()
                if first_line:
                    return first_line[:120]
        return _safe_filename_stem(source_name)

    def _write_result_sheet(self, workbook, result: dict[str, Any]) -> None:
        if AUDIT_RESULT_SHEET in workbook.sheetnames:
            workbook.remove(workbook[AUDIT_RESULT_SHEET])
        ws = workbook.create_sheet(AUDIT_RESULT_SHEET, 0)
        ws.sheet_view.showGridLines = False
        ws.freeze_panes = "A9"
        ws.page_setup.orientation = "landscape"
        ws.page_setup.fitToWidth = 1
        ws.sheet_properties.pageSetUpPr.fitToPage = True
        ws.print_options.horizontalCentered = True
        ws.sheet_properties.outlinePr.summaryBelow = True

        widths = {
            "A": 12,
            "B": 17,
            "C": 24,
            "D": 18,
            "E": 9,
            "F": 16,
            "G": 16,
            "H": 42,
            "I": 42,
        }
        for column, width in widths.items():
            ws.column_dimensions[column].width = width

        ws.merge_cells("A1:I1")
        title_cell = ws["A1"]
        title_cell.value = "勘察测量结算辅助审核结果"
        title_cell.fill = PatternFill("solid", fgColor=NAVY)
        title_cell.font = Font(name=AUDIT_SHEET_FONT_NAME, size=18, bold=True, color=WHITE)
        title_cell.alignment = Alignment(horizontal="left", vertical="center")
        ws.row_dimensions[1].height = 34

        ws.merge_cells("A2:I2")
        ws["A2"] = "规则辅助审核 · 人工最终审定｜本结果不替代合同判断、资料真实性核验及正式签章"
        ws["A2"].fill = PatternFill("solid", fgColor=PALE_BLUE)
        ws["A2"].font = Font(name=AUDIT_SHEET_FONT_NAME, size=10, color=NAVY, bold=True)
        ws["A2"].alignment = Alignment(vertical="center")
        ws.row_dimensions[2].height = 24

        summary = result["summary"]
        ws.merge_cells("A3:D3")
        ws.merge_cells("E3:I3")
        ws["A3"] = f"项目：{result['project_name']}"
        ws["E3"] = f"源文件：{result['source_file']}"
        for coordinate in ("A3", "E3"):
            cell = ws[coordinate]
            cell.font = Font(name=AUDIT_SHEET_FONT_NAME, size=AUDIT_SHEET_BODY_FONT_SIZE, color=TEXT)
            cell.alignment = Alignment(vertical="center", wrap_text=False)
            cell.border = Border(bottom=Side(style="thin", color=BORDER))
        ws.row_dimensions[3].height = 24

        metrics = (
            ("审核 Sheet", summary["sheet_count"]),
            ("审核明细行", summary["audited_rows"]),
            ("通过行", summary["passed_rows"]),
            ("结构化风险", summary["risk_count"]),
            ("高风险", summary["high_risk_count"]),
            ("待人工核验", summary["manual_review_count"]),
            ("上报明细试算", summary["reported_detail_total"]),
            ("审核建议试算", summary["reviewed_detail_total"]),
            ("建议差额", summary["suggested_difference"]),
        )
        for index, (label, value) in enumerate(metrics, start=1):
            label_cell = ws.cell(4, index)
            value_cell = ws.cell(5, index)
            label_cell.value = label
            label_cell.font = Font(
                name=AUDIT_SHEET_FONT_NAME,
                size=AUDIT_SHEET_BODY_FONT_SIZE,
                color=MUTED,
            )
            label_cell.alignment = Alignment(horizontal="center", vertical="center")
            value_cell.value = value
            value_cell.font = Font(name=AUDIT_SHEET_FONT_NAME, size=11, bold=True, color=TEXT)
            value_cell.alignment = Alignment(horizontal="center", vertical="center", wrap_text=True)
            value_cell.fill = PatternFill("solid", fgColor="F8FAFC")
            value_cell.border = Border(bottom=Side(style="thin", color=BORDER))
            if index >= 7:
                value_cell.number_format = '#,##0.00" 元"'
        ws.row_dimensions[5].height = 35

        ws.merge_cells("A7:I7")
        ws["A7"] = "一、发现问题"
        ws["A7"].font = Font(name=AUDIT_SHEET_FONT_NAME, size=12, bold=True, color=NAVY)
        ws["A7"].fill = PatternFill("solid", fgColor="F1F5F9")
        ws["A7"].alignment = Alignment(vertical="center")

        headers = ("风险等级", "规则", "审核事项", "Sheet", "行号", "当前值", "建议值", "依据", "处置建议")
        for column, header in enumerate(headers, start=1):
            cell = ws.cell(8, column)
            cell.value = header
            cell.fill = PatternFill("solid", fgColor=BLUE)
            cell.font = Font(
                name=AUDIT_SHEET_FONT_NAME,
                size=AUDIT_SHEET_BODY_FONT_SIZE,
                bold=True,
                color=WHITE,
            )
            cell.alignment = Alignment(horizontal="center", vertical="center", wrap_text=True)
        ws.row_dimensions[8].height = 28

        current_row = 9
        for risk in result["risks"]:
            severity_label = {"high": "高", "medium": "中", "low": "低"}[risk["severity"]]
            values = (
                severity_label,
                risk["rule_id"],
                risk["title"],
                _short_sheet_label(risk["sheet"]),
                risk["row"] or "",
                _display_value(risk["current_value"]),
                _display_value(risk["suggested_value"]),
                risk["basis"],
                risk["action"],
            )
            fill = _risk_fill(risk["severity"], risk["auto_adjusted"])
            for column, value in enumerate(values, start=1):
                cell = ws.cell(current_row, column)
                cell.value = value
                cell.fill = fill if column == 1 else PatternFill("solid", fgColor=WHITE)
                cell.font = Font(
                    name=AUDIT_SHEET_FONT_NAME,
                    size=AUDIT_SHEET_BODY_FONT_SIZE,
                    color=TEXT,
                    bold=column == 1,
                )
                cell.alignment = Alignment(vertical="top", wrap_text=True)
                cell.border = Border(bottom=Side(style="thin", color=BORDER))
            if risk["row"]:
                sheet_cell = ws.cell(current_row, 4)
                sheet_cell.hyperlink = f"#'{risk['sheet'].replace(chr(39), chr(39) * 2)}'!A{risk['row']}"
                sheet_cell.font = Font(
                    name=AUDIT_SHEET_FONT_NAME,
                    size=AUDIT_SHEET_BODY_FONT_SIZE,
                    color=BLUE,
                    underline="single",
                )
            ws.row_dimensions[current_row].height = 44
            current_row += 1

        if not result["risks"]:
            ws.merge_cells(start_row=current_row, start_column=1, end_row=current_row, end_column=9)
            ws.cell(current_row, 1).value = "未发现可由系统确定的结构化风险；仍须完成下方人工核验。"
            ws.cell(current_row, 1).fill = PatternFill("solid", fgColor=PALE_GREEN)
            ws.cell(current_row, 1).font = Font(
                name=AUDIT_SHEET_FONT_NAME,
                size=AUDIT_SHEET_BODY_FONT_SIZE,
                color=TEXT,
            )
            ws.cell(current_row, 1).alignment = Alignment(vertical="center")
            current_row += 1

        current_row += 1
        ws.merge_cells(start_row=current_row, start_column=1, end_row=current_row, end_column=9)
        ws.cell(current_row, 1).value = "二、待人工核验"
        ws.cell(current_row, 1).font = Font(
            name=AUDIT_SHEET_FONT_NAME,
            size=12,
            bold=True,
            color=NAVY,
        )
        ws.cell(current_row, 1).fill = PatternFill("solid", fgColor="F1F5F9")
        current_row += 1
        for item in result["manual_checklist"]:
            ws.cell(current_row, 1).value = item["id"]
            ws.cell(current_row, 2).value = item["title"]
            ws.merge_cells(start_row=current_row, start_column=3, end_row=current_row, end_column=9)
            ws.cell(current_row, 3).value = item["detail"]
            for column in range(1, 10):
                cell = ws.cell(current_row, column)
                cell.fill = PatternFill("solid", fgColor=PALE_RED if column == 1 else WHITE)
                cell.font = Font(
                    name=AUDIT_SHEET_FONT_NAME,
                    size=AUDIT_SHEET_BODY_FONT_SIZE,
                    color=TEXT,
                    bold=column in (1, 2),
                )
                cell.alignment = Alignment(vertical="top", wrap_text=True)
                cell.border = Border(bottom=Side(style="thin", color=BORDER))
            ws.row_dimensions[current_row].height = 36
            current_row += 1

        current_row += 1
        ws.merge_cells(start_row=current_row, start_column=1, end_row=current_row, end_column=9)
        ws.cell(current_row, 1).value = result["scope_note"]
        ws.cell(current_row, 1).font = Font(
            name=AUDIT_SHEET_FONT_NAME,
            size=AUDIT_SHEET_BODY_FONT_SIZE,
            italic=True,
            color=MUTED,
        )
        ws.cell(current_row, 1).alignment = Alignment(wrap_text=True, vertical="center")
        ws.cell(current_row, 1).fill = PatternFill("solid", fgColor="F8FAFC")
        ws.row_dimensions[current_row].height = 30
        ws.auto_filter.ref = f"A8:I{max(8, 8 + len(result['risks']))}"
        ws.print_area = f"A1:I{current_row}"

    def _write_report(self, path: Path, result: dict[str, Any]) -> None:
        document = Document()
        section = document.sections[0]
        section.top_margin = Cm(2.2)
        section.bottom_margin = Cm(2.0)
        section.left_margin = Cm(2.1)
        section.right_margin = Cm(2.1)

        normal = document.styles["Normal"]
        normal.font.name = "等线"
        normal.font.size = Pt(10.5)
        normal.font.color.rgb = RGBColor.from_string(TEXT)
        normal._element.get_or_add_rPr().rFonts.set(qn("w:eastAsia"), "等线")
        normal.paragraph_format.space_after = Pt(6)
        normal.paragraph_format.line_spacing = 1.25

        for style_name, size, color in (
            ("Title", 22, NAVY),
            ("Heading 1", 15, NAVY),
            ("Heading 2", 12, BLUE),
        ):
            style = document.styles[style_name]
            style.font.name = "等线"
            style.font.size = Pt(size)
            style.font.color.rgb = RGBColor.from_string(color)
            style._element.get_or_add_rPr().rFonts.set(qn("w:eastAsia"), "等线")

        title = document.add_heading("勘察测量结算辅助审核报告", level=0)
        title.alignment = WD_ALIGN_PARAGRAPH.CENTER
        subtitle = document.add_paragraph()
        subtitle.alignment = WD_ALIGN_PARAGRAPH.CENTER
        run = subtitle.add_run("规则辅助审核 · 待人工审定")
        run.bold = True
        run.font.name = "等线"
        run.font.size = Pt(13)
        run.font.color.rgb = RGBColor.from_string(BLUE)
        run._element.get_or_add_rPr().rFonts.set(qn("w:eastAsia"), "等线")

        notice = document.add_paragraph()
        notice.alignment = WD_ALIGN_PARAGRAPH.CENTER
        notice_run = notice.add_run("本报告不替代合同判断、资料真实性核验、专业复核、内审及正式签章。")
        notice_run.bold = True
        notice_run.font.color.rgb = RGBColor.from_string("B42318")

        info = document.add_table(rows=4, cols=2)
        info.style = "Table Grid"
        info_data = (
            ("项目", result["project_name"]),
            ("源文件", result["source_file"]),
            ("审核时间", result["created_at"]),
            ("规则版本", result["rule_version"]),
        )
        for row, (label, value) in zip(info.rows, info_data, strict=True):
            _prevent_docx_row_split(row)
            _set_docx_cell_text(row.cells[0], label, bold=True, color=NAVY)
            _set_docx_cell_shading(row.cells[0], PALE_BLUE)
            _set_docx_cell_text(row.cells[1], value)

        document.add_heading("一、审核范围与依据", level=1)
        document.add_paragraph(
            "本次审核面向前辈提供的统一勘察测量结算报价模板，使用确定性规则校核承包商上报明细、模板参数、"
            "金额算术及已明确的专项结算经验。系统没有读取合同原件、现场成果和付款原件，因此涉及真实性、"
            "适用合同及签章的事项均保留为待人工核验。"
        )
        for source in result["sources"]:
            document.add_paragraph(source, style="List Bullet")

        document.add_heading("二、审核汇总", level=1)
        summary = result["summary"]
        summary_table = document.add_table(rows=2, cols=7)
        summary_table.style = "Table Grid"
        summary_items = (
            ("审核明细行", summary["audited_rows"]),
            ("通过行", summary["passed_rows"]),
            ("结构化风险", summary["risk_count"]),
            ("高风险", summary["high_risk_count"]),
            ("待人工核验", summary["manual_review_count"]),
            ("审核建议试算", f"{summary['reviewed_detail_total']:,.2f} 元"),
            ("建议差额", f"{summary['suggested_difference']:,.2f} 元"),
        )
        for column, (label, value) in enumerate(summary_items):
            _set_docx_cell_text(summary_table.cell(0, column), label, bold=True, color=WHITE)
            _set_docx_cell_shading(summary_table.cell(0, column), BLUE)
            _set_docx_cell_text(summary_table.cell(1, column), value, bold=True)
        for row in summary_table.rows:
            _prevent_docx_row_split(row)
        scope = document.add_paragraph(result["scope_note"])
        scope.runs[0].italic = True
        scope.runs[0].font.color.rgb = RGBColor.from_string(MUTED)

        if result["risks"]:
            document.add_page_break()
        document.add_heading("三、发现问题", level=1)
        if result["risks"]:
            risk_table = document.add_table(rows=1, cols=7)
            risk_table.style = "Table Grid"
            headings = ("等级", "位置", "审核事项", "当前值", "建议值", "依据", "处置建议")
            for column, heading in enumerate(headings):
                _set_docx_cell_text(risk_table.cell(0, column), heading, bold=True, color=WHITE)
                _set_docx_cell_shading(risk_table.cell(0, column), NAVY)
            _repeat_docx_table_header(risk_table.rows[0])
            _prevent_docx_row_split(risk_table.rows[0])
            for risk in result["risks"]:
                row = risk_table.add_row()
                values = (
                    {"high": "高", "medium": "中", "low": "低"}[risk["severity"]],
                    f"{risk['coordinate']}\n{_short_sheet_label(risk['sheet'])}",
                    f"{risk['rule_id']} {risk['title']}",
                    risk["current_value"],
                    risk["suggested_value"],
                    risk["basis"],
                    risk["action"],
                )
                for column, value in enumerate(values):
                    _set_docx_cell_text(row.cells[column], value, bold=column == 0)
                _set_docx_cell_shading(
                    row.cells[0],
                    PALE_RED if risk["severity"] == "high" else PALE_YELLOW,
                )
                _prevent_docx_row_split(row)
        else:
            document.add_paragraph("未发现可由系统确定的结构化风险。")

        document.add_heading("四、待人工核验资料", level=1)
        for item in result["manual_checklist"]:
            paragraph = document.add_paragraph(style="List Number")
            lead = paragraph.add_run(f"{item['title']}：")
            lead.bold = True
            paragraph.add_run(item["detail"])

        document.add_heading("五、辅助审核结论", level=1)
        document.add_paragraph(
            f"本次共审核 {summary['audited_rows']} 条有值明细，发现 {summary['risk_count']} 项结构化风险，"
            f"其中高风险 {summary['high_risk_count']} 项。审核建议已写入 Excel 副本的“造价审核”栏，"
            "黄色单元格表示系统已形成可追溯建议，红色提示表示仍需资料或合同判断。请由专业人员逐项复核，"
            "完成承包商确认、内部审查和正式签章后再形成最终结算金额。"
        )

        document.add_heading("六、人工复核签认", level=1)
        document.add_paragraph("以下栏位留供专业人员完成资料核验后填写；空白不代表默认同意系统建议。")
        signoff_table = document.add_table(rows=4, cols=2)
        signoff_table.style = "Table Grid"
        signoff_items = (
            ("专业复核意见", "（请人工填写）"),
            ("造价审核意见", "（请人工填写）"),
            ("承包商确认情况", "（请人工填写）"),
            ("最终审定状态", "□ 待复核    □ 已确认    □ 退回补充资料"),
        )
        for row, (label, placeholder) in zip(signoff_table.rows, signoff_items, strict=True):
            _prevent_docx_row_split(row)
            _set_docx_cell_text(row.cells[0], label, bold=True, color=NAVY)
            _set_docx_cell_shading(row.cells[0], PALE_BLUE)
            _set_docx_cell_text(row.cells[1], placeholder, color=MUTED)

        footer = section.footer.paragraphs[0]
        footer.alignment = WD_ALIGN_PARAGRAPH.CENTER
        footer_run = footer.add_run("造价智算 · 结算审核助手 · 辅助审核，人工定案")
        footer_run.font.name = "等线"
        footer_run.font.size = Pt(8)
        footer_run.font.color.rgb = RGBColor.from_string(MUTED)
        footer_run._element.get_or_add_rPr().rFonts.set(qn("w:eastAsia"), "等线")

        if len(result["risks"]) > 7:
            section.orientation = WD_ORIENT.LANDSCAPE
            section.page_width, section.page_height = section.page_height, section.page_width
        document.save(path)
