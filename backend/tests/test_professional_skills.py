from __future__ import annotations

import json
import os
import shutil
from copy import deepcopy
from pathlib import Path

import pytest
from fastapi.testclient import TestClient
from docx import Document
from openpyxl import Workbook, load_workbook

import app.main as main_module
from app.main import app
from app.experience_warning import EXPERIENCE_POOL_HEADERS
from app.paths import BUSINESS_SKILLS_DIR, PROJECT_DEFAULT_SETTINGS_PATH, PROJECT_ROOT
from app.professional_skills import ProfessionalSkillError, ProfessionalSkillRegistry


ACTIVE_ID = "survey-measurement-limit-price"


def _base_manifest(**overrides):
    manifest = {
        "id": ACTIVE_ID,
        "displayName": "勘察测量最高投标限价编制",
        "version": "1.0.0",
        "status": "active",
        "domain": "工程造价/勘察测量",
        "description": "测试专业能力",
        "inputProfile": {"extensions": [".xlsx"], "templateHints": []},
        "capabilities": {
            "pricing": True,
            "workloadCapture": True,
            "experienceWarning": True,
            "knowledgeQa": True,
            "wordReport": True,
        },
        "subSkills": [
            {
                "name": "造价规则匹配 Skill",
                "description": "执行确定性规则匹配。",
                "type": "shared",
                "status": "available",
            }
        ],
        "assets": {
            "knowledgeBase": "kb.xlsx",
            "reportTemplate": "template.docx",
            "technicalRules": "technical.xlsx",
            "physicalRules": "physical.xlsx",
            "physicalOverrides": "overrides.xlsx",
            "experiencePool": "pool.xlsx",
            "experienceWarningSettings": "settings.json",
            "knowledgeSources": "knowledge.md",
            "validationSample": "sample.xlsx",
        },
        "runtime": {
            "processorId": "survey-measurement-v1",
            "knowledgeBaseAsset": "knowledgeBase",
            "ruleAssets": {
                "technicalRules": "technicalRules",
                "physicalRules": "physicalRules",
                "physicalOverrides": "physicalOverrides",
            },
            "riskProfile": {
                "experiencePool": "experiencePool",
                "warningSettings": "experienceWarningSettings",
            },
            "knowledgeSourceAssets": ["knowledgeSources"],
            "reportTemplateAsset": "reportTemplate",
            "validationAsset": "validationSample",
        },
        "validation": {"status": "verified", "sample": "sample", "updatedAt": "2026-07-20"},
    }
    manifest.update(overrides)
    return manifest


def _registry(tmp_path: Path, manifest: dict | None = None) -> tuple[ProfessionalSkillRegistry, Path, Path]:
    project_root = tmp_path / "project"
    skills_root = project_root / "business-skills"
    manifest_path = skills_root / ACTIVE_ID / "manifest.json"
    manifest_path.parent.mkdir(parents=True)
    project_root.mkdir(parents=True, exist_ok=True)
    workbook = Workbook()
    workbook.active.append(["要素1", "要素2", "要素3", "要素4", "要素5", "单位", "基价"])
    workbook.save(project_root / "kb.xlsx")
    workbook.save(project_root / "sample.xlsx")
    workbook.close()
    shutil.copy2(PROJECT_ROOT / "backend/app/rules/technical_fee_rules.xlsx", project_root / "technical.xlsx")
    shutil.copy2(PROJECT_ROOT / "backend/app/rules/physical_factor_rules.xlsx", project_root / "physical.xlsx")
    shutil.copy2(PROJECT_ROOT / "backend/app/rules/physical_factor_overrides.xlsx", project_root / "overrides.xlsx")
    shutil.copy2(PROJECT_ROOT / "05-经验池-预警数据/【经验池】-管勘智算-【codex】.xlsx", project_root / "pool.xlsx")
    shutil.copy2(
        PROJECT_ROOT / "03-知识库-二维数据库制作/01-报告模板-招标控制价报告模板/【模板勿动】控制价报告模板-yyyy-mm-dd.docx",
        project_root / "template.docx",
    )
    (project_root / "settings.json").write_text("{}", encoding="utf-8")
    (project_root / "knowledge.md").write_text("# 测试知识\n\n专业能力测试依据。", encoding="utf-8")
    manifest_path.write_text(json.dumps(manifest or _base_manifest(), ensure_ascii=False), encoding="utf-8")
    settings_path = project_root / "config" / "project-default-settings.json"
    settings_path.parent.mkdir(parents=True)
    settings_path.write_text(
        json.dumps({"professionalSkills": {"defaultSkillId": ACTIVE_ID}}, ensure_ascii=False),
        encoding="utf-8",
    )
    return ProfessionalSkillRegistry(project_root, skills_root, settings_path), project_root, manifest_path


def _write_workbooks(tmp_path: Path) -> tuple[Path, Path]:
    kb_path = tmp_path / "kb.xlsx"
    workbook = Workbook()
    sheet = workbook.active
    sheet.append(["要素1", "要素2", "要素3", "要素4", "要素5", "单位", "基价"])
    sheet.append(["控制测量", "GPS测量E级", "", "中等", "", "点", 3203])
    workbook.save(kb_path)
    workbook.close()

    input_path = tmp_path / "input.xlsx"
    workbook = Workbook()
    sheet = workbook.active
    sheet.title = "表2"
    sheet.append(["要素1", "要素2", "要素3", "要素4", "要素5", "单位", "基价"])
    sheet.append(["控制测量", "GPS测量E级", "", "中等", "", "点", ""])
    workbook.save(input_path)
    workbook.close()
    return kb_path, input_path


def _install_runtime_skill(
    project_root: Path,
    skill_id: str,
    *,
    price: int,
    coefficient: float,
    knowledge_token: str,
    report_marker: str,
    pool_price: int,
) -> None:
    asset_dir = project_root / "fixtures" / skill_id
    asset_dir.mkdir(parents=True, exist_ok=True)

    workbook = Workbook()
    sheet = workbook.active
    sheet.append(["要素1", "要素2", "要素3", "要素4", "要素5", "单位", "基价"])
    sheet.append(["控制测量", "GPS测量E级", "", "中等", "", "点", price])
    workbook.save(asset_dir / "kb.xlsx")
    workbook.save(asset_dir / "sample.xlsx")
    workbook.close()

    workbook = Workbook()
    sheet = workbook.active
    sheet.append([
        "rule_id", "sheet_tokens", "business_keywords", "category_keywords", "coefficient",
        "source_type", "confidence", "basis", "formula_effective", "review_required", "note", "priority",
    ])
    sheet.append([
        f"{skill_id}-technical", "表2", "GPS测量E级", "", str(coefficient),
        "测试规则", "高", f"{skill_id} 独立规则", "true", "false", "隔离验证", "高",
    ])
    workbook.save(asset_dir / "technical.xlsx")
    workbook.close()
    shutil.copy2(PROJECT_ROOT / "backend/app/rules/physical_factor_rules.xlsx", asset_dir / "physical.xlsx")
    shutil.copy2(PROJECT_ROOT / "backend/app/rules/physical_factor_overrides.xlsx", asset_dir / "overrides.xlsx")

    document = Document()
    document.add_paragraph(report_marker)
    document.save(asset_dir / "template.docx")
    (asset_dir / "knowledge.md").write_text(
        f"# {knowledge_token}依据\n\n{knowledge_token} {knowledge_token} {knowledge_token} 专业规则。\n",
        encoding="utf-8",
    )
    (asset_dir / "settings.json").write_text(
        json.dumps({"settings": {"only_check_rows_with_value": True, "value_filter_field": "数量"}}),
        encoding="utf-8",
    )
    workbook = Workbook()
    sheet = workbook.active
    sheet.append(EXPERIENCE_POOL_HEADERS)
    record = {
        "来源文件": f"{skill_id}.xlsx",
        "来源sheet": "表2",
        "来源行": 2,
        "要素1": "控制测量",
        "要素2": "GPS测量E级",
        "要素3": "",
        "要素4": "中等",
        "要素5": "",
        "单位": "点",
        "基价": pool_price,
    }
    sheet.append([record.get(header, "") for header in EXPERIENCE_POOL_HEADERS])
    workbook.save(asset_dir / "pool.xlsx")
    workbook.close()

    relative = lambda name: (asset_dir / name).relative_to(project_root).as_posix()
    manifest = _base_manifest(
        id=skill_id,
        displayName=f"测试能力 {skill_id}",
        assets={
            "knowledgeBase": relative("kb.xlsx"),
            "reportTemplate": relative("template.docx"),
            "technicalRules": relative("technical.xlsx"),
            "physicalRules": relative("physical.xlsx"),
            "physicalOverrides": relative("overrides.xlsx"),
            "experiencePool": relative("pool.xlsx"),
            "experienceWarningSettings": relative("settings.json"),
            "knowledgeSources": relative("knowledge.md"),
            "validationSample": relative("sample.xlsx"),
        },
    )
    manifest_path = project_root / "business-skills" / skill_id / "manifest.json"
    manifest_path.parent.mkdir(parents=True, exist_ok=True)
    manifest_path.write_text(json.dumps(manifest, ensure_ascii=False), encoding="utf-8")


def _write_runtime_input(path: Path) -> Path:
    workbook = Workbook()
    sheet = workbook.active
    sheet.title = "表2"
    sheet.append([
        "要素1", "要素2", "要素3", "要素4", "要素5", "单位", "基价",
        "技术工作费调整系数", "数量",
    ])
    sheet.append(["控制测量", "GPS测量E级", "", "中等", "", "点", "", "", 1])
    workbook.save(path)
    workbook.close()
    return path


def _process_data(**overrides):
    data = {
        "column_mapping": json.dumps(
            {
                "要素1": "A",
                "要素2": "B",
                "要素3": "C",
                "要素4": "D",
                "要素5": "E",
                "单位": "F",
                "输出-价格列": "G",
            },
            ensure_ascii=False,
        ),
        "only_match_rows_with_value": "false",
        "defer_matching": "true",
    }
    data.update(overrides)
    return data


def _post_process(client: TestClient, input_path: Path, **data_overrides):
    with input_path.open("rb") as handle:
        return client.post(
            "/api/process",
            files={"file": ("input.xlsx", handle, "application/vnd.openxmlformats-officedocument.spreadsheetml.sheet")},
            data=_process_data(**data_overrides),
        )


def test_production_registry_loads_active_default_and_planned_card_without_paths():
    registry = ProfessionalSkillRegistry(PROJECT_ROOT, BUSINESS_SKILLS_DIR, PROJECT_DEFAULT_SETTINGS_PATH)

    payload = registry.list_public()
    active = next(item for item in payload["items"] if item["id"] == ACTIVE_ID)
    planned = next(item for item in payload["items"] if item["id"] == "general-service-cost-estimation")
    upcoming = next(item for item in payload["items"] if item["id"] == "more-cost-professions")
    detail = registry.get_public(ACTIVE_ID)

    assert payload["default_skill_id"] == ACTIVE_ID
    assert active["status"] == "active" and active["can_create_task"] is True
    assert planned["status"] == "planned" and planned["can_create_task"] is False
    assert upcoming["status_label"] == "即将开放" and upcoming["can_create_task"] is False
    assert [item["id"] for item in payload["items"]] == [
        ACTIVE_ID,
        "general-service-cost-estimation",
        "more-cost-professions",
    ]
    assert detail["validation"]["status"] == "verified"
    assert len(detail["sub_skills"]) == 10
    assert detail["sub_skills"][0]["name"] == "控制价 Excel 智能识别 Skill"
    assert any(item["type"] == "professional" for item in detail["sub_skills"])
    assert detail["asset_summary"] and "03-知识库" not in json.dumps(detail, ensure_ascii=False)
    assert str(PROJECT_ROOT) not in json.dumps(payload, ensure_ascii=False)


def test_registry_creates_immutable_safe_snapshot_and_rejects_planned_unknown_and_version_mismatch(tmp_path):
    registry, _, manifest_path = _registry(tmp_path)
    manifest = registry.load(ACTIVE_ID)
    snapshot = registry.create_snapshot(manifest)
    manifest["version"] = "9.9.9"
    manifest_path.write_text(json.dumps({**_base_manifest(), "version": "1.1.0"}), encoding="utf-8")

    assert snapshot["version"] == "1.0.0"
    assert snapshot["manifest_hash"]
    assert "runtime_context" in snapshot
    assert "path" not in json.dumps(ProfessionalSkillRegistry.public_snapshot(snapshot)).lower()
    with pytest.raises(ProfessionalSkillError) as version_error:
        registry.resolve_for_task(ACTIVE_ID, "0.9.0")
    assert version_error.value.code == "skill_version_mismatch"
    with pytest.raises(ProfessionalSkillError) as unknown_error:
        registry.resolve_for_task("unknown-skill", None)
    assert unknown_error.value.code == "skill_not_found"

    planned = _base_manifest(status="planned", assets={})
    manifest_path.write_text(json.dumps(planned), encoding="utf-8")
    with pytest.raises(ProfessionalSkillError) as planned_error:
        registry.resolve_for_task(ACTIVE_ID, "1.0.0")
    assert planned_error.value.code == "skill_not_available"


@pytest.mark.parametrize(
    ("mutator", "expected_code"),
    [
        (lambda payload: payload.pop("validation"), "skill_manifest_invalid"),
        (lambda payload: payload.update(status="mystery"), "skill_manifest_invalid"),
        (lambda payload: payload.update(command="run"), "skill_manifest_unsafe"),
        (lambda payload: payload.update(api_key="secret"), "skill_manifest_unsafe"),
        (lambda payload: payload.update(subSkills={}), "skill_manifest_invalid"),
        (lambda payload: payload.update(subSkills=[{"name": "缺少说明", "type": "shared", "status": "available"}]), "skill_manifest_invalid"),
        (lambda payload: payload["runtime"].update(processorId="untrusted-processor"), "skill_processor_not_allowed"),
        (lambda payload: payload.update(assets={"knowledgeBase": "../outside.txt"}), "skill_asset_unsafe"),
        (lambda payload: payload.update(assets={"knowledgeBase": "missing.txt"}), "skill_asset_unavailable"),
        (lambda payload: payload.update(assets={"knowledgeBase": "worker.py"}), "skill_asset_unsafe"),
    ],
)
def test_registry_rejects_incomplete_unsafe_and_missing_manifests(tmp_path, mutator, expected_code):
    payload = deepcopy(_base_manifest())
    mutator(payload)
    registry, _, _ = _registry(tmp_path, payload)

    with pytest.raises(ProfessionalSkillError) as error:
        registry.load(ACTIVE_ID)

    assert error.value.code == expected_code
    assert str(tmp_path) not in error.value.message


def test_registry_rejects_absolute_path_and_symlink_escape(tmp_path):
    outside = tmp_path / "outside.txt"
    outside.write_text("outside", encoding="utf-8")
    absolute_manifest = _base_manifest(assets={"knowledgeBase": str(outside.resolve())})
    registry, project_root, manifest_path = _registry(tmp_path, absolute_manifest)
    with pytest.raises(ProfessionalSkillError) as absolute_error:
        registry.load(ACTIVE_ID)
    assert absolute_error.value.code == "skill_asset_unsafe"

    link_path = project_root / "linked.txt"
    try:
        os.symlink(outside, link_path)
    except OSError:
        pytest.skip("当前 Windows 环境不允许创建测试符号链接")
    manifest_path.write_text(json.dumps(_base_manifest(assets={"knowledgeBase": "linked.txt"})), encoding="utf-8")
    with pytest.raises(ProfessionalSkillError) as symlink_error:
        registry.load(ACTIVE_ID)
    assert symlink_error.value.code == "skill_asset_unavailable"


def test_default_request_can_use_safe_compatibility_snapshot_when_manifest_is_missing(tmp_path):
    registry, _, manifest_path = _registry(tmp_path)
    manifest_path.unlink()

    snapshot = registry.resolve_for_task(None, None)

    assert snapshot["id"] == ACTIVE_ID
    assert snapshot["compatibility_fallback"] is True
    with pytest.raises(ProfessionalSkillError):
        registry.resolve_for_task(ACTIVE_ID, "1.0.0")


def test_professional_skill_api_is_safe_and_process_preserves_snapshot_across_batch_and_download(tmp_path, monkeypatch):
    kb_path, input_path = _write_workbooks(tmp_path)
    registry, project_root, _ = _registry(tmp_path)
    shutil.copy2(kb_path, project_root / "kb.xlsx")
    runtime_dir = tmp_path / "runtime"
    monkeypatch.setattr(main_module, "RUNTIME_DIR", runtime_dir)
    monkeypatch.setattr(main_module, "PROFESSIONAL_SKILL_REGISTRY", registry)
    client = TestClient(app)

    catalog_response = client.get("/api/professional-skills")
    assert catalog_response.status_code == 200
    assert str(PROJECT_ROOT) not in catalog_response.text

    legacy_response = _post_process(client, input_path)
    assert legacy_response.status_code == 200
    legacy_payload = legacy_response.json()
    snapshot = legacy_payload["professional_skill"]
    assert snapshot["id"] == ACTIVE_ID
    assert snapshot["version"] == "1.0.0"

    explicit_response = _post_process(client, input_path, skill_id=ACTIVE_ID, skill_version="1.0.0")
    assert explicit_response.status_code == 200
    assert explicit_response.json()["summary"] == legacy_payload["summary"]

    batch_response = client.post("/api/process/batch-match", json={"job_id": legacy_payload["job_id"]})
    assert batch_response.status_code == 200
    assert batch_response.json()["professional_skill"] == snapshot

    risk_response = client.get(f"/api/risk/summary?job_id={legacy_payload['job_id']}")
    assert risk_response.status_code == 200
    assert risk_response.json()["professional_skill"] == snapshot

    download_response = client.get(f"/api/download/{legacy_payload['job_id']}/excel")
    assert download_response.status_code == 200
    assert download_response.headers["x-professional-skill-id"] == ACTIVE_ID
    assert download_response.headers["x-professional-skill-version"] == "1.0.0"

    output_path = next((runtime_dir / legacy_payload["job_id"]).glob("【输出】-控制价计算表-*.xlsx"))
    workbook = load_workbook(output_path, data_only=True)
    try:
        assert workbook["表2"]["G2"].value == 3203
    finally:
        workbook.close()


def test_runtime_context_freezes_and_isolates_two_test_skills_across_assets(tmp_path, monkeypatch):
    project_root = tmp_path / "project"
    skills_root = project_root / "business-skills"
    settings_path = project_root / "config" / "project-default-settings.json"
    settings_path.parent.mkdir(parents=True)
    first_id = "test-survey-a"
    second_id = "test-survey-b"
    _install_runtime_skill(
        project_root,
        first_id,
        price=100,
        coefficient=0.11,
        knowledge_token="甲专属知识",
        report_marker="甲专属报告模板",
        pool_price=100,
    )
    _install_runtime_skill(
        project_root,
        second_id,
        price=200,
        coefficient=0.22,
        knowledge_token="乙专属知识",
        report_marker="乙专属报告模板",
        pool_price=50,
    )
    settings_path.write_text(
        json.dumps({"professionalSkills": {"defaultSkillId": first_id}}, ensure_ascii=False),
        encoding="utf-8",
    )
    registry = ProfessionalSkillRegistry(project_root, skills_root, settings_path, fallback_default_skill_id=first_id)
    runtime_dir = tmp_path / "runtime"
    monkeypatch.setattr(main_module, "RUNTIME_DIR", runtime_dir)
    monkeypatch.setattr(main_module, "PROFESSIONAL_SKILL_REGISTRY", registry)
    input_path = _write_runtime_input(tmp_path / "runtime-input.xlsx")
    mapping = json.dumps(
        {
            "要素1": "A", "要素2": "B", "要素3": "C", "要素4": "D", "要素5": "E",
            "单位": "F", "输出-价格列": "G", "输出-技术工作费调整系数": "H",
        },
        ensure_ascii=False,
    )
    client = TestClient(app)

    first_response = _post_process(
        client,
        input_path,
        column_mapping=mapping,
        defer_matching="true",
        skill_id=first_id,
        skill_version="1.0.0",
    )
    assert first_response.status_code == 200, first_response.text
    first_job = first_response.json()["job_id"]
    settings_path.write_text(
        json.dumps({"professionalSkills": {"defaultSkillId": second_id}}, ensure_ascii=False),
        encoding="utf-8",
    )
    batch_response = client.post("/api/process/batch-match", json={"job_id": first_job})
    assert batch_response.status_code == 200, batch_response.text

    second_response = _post_process(
        client,
        input_path,
        column_mapping=mapping,
        defer_matching="false",
        skill_id=second_id,
        skill_version="1.0.0",
    )
    assert second_response.status_code == 200, second_response.text
    second_job = second_response.json()["job_id"]

    for job_id, expected_price, expected_coefficient, marker in (
        (first_job, 100, 0.11, "甲专属报告模板"),
        (second_job, 200, 0.22, "乙专属报告模板"),
    ):
        excel_path = next((runtime_dir / job_id).glob("【输出】-控制价计算表-*.xlsx"))
        workbook = load_workbook(excel_path, data_only=True)
        try:
            assert workbook["表2"]["G2"].value == expected_price
            assert workbook["表2"]["H2"].value == expected_coefficient
        finally:
            workbook.close()
        report_path = next((runtime_dir / job_id).glob("【输出】-控制价报告-*.docx"))
        assert marker in "\n".join(paragraph.text for paragraph in Document(report_path).paragraphs)

    first_knowledge = client.post("/api/knowledge/search", json={"job_id": first_job, "question": "甲专属知识依据"})
    second_knowledge = client.post("/api/knowledge/search", json={"job_id": second_job, "question": "乙专属知识依据"})
    crossed_knowledge = client.post("/api/knowledge/search", json={"job_id": first_job, "question": "乙专属知识依据"})
    assert first_knowledge.status_code == second_knowledge.status_code == crossed_knowledge.status_code == 200
    assert first_knowledge.json()["results"]
    assert second_knowledge.json()["results"]
    assert crossed_knowledge.json()["results"] == []

    first_warning = client.post("/api/experience-warnings/run", data={"job_id": first_job})
    second_warning = client.post("/api/experience-warnings/run", data={"job_id": second_job})
    assert first_warning.status_code == second_warning.status_code == 200
    assert first_warning.json()["summary"]["warning_summary"]["warning_rows"] == 0
    assert second_warning.json()["summary"]["warning_summary"]["warning_rows"] == 1

    first_state = json.loads((runtime_dir / first_job / "process-state.json").read_text(encoding="utf-8"))
    second_state = json.loads((runtime_dir / second_job / "process-state.json").read_text(encoding="utf-8"))
    first_runtime = first_state["skill_snapshot"]["runtime_context"]
    second_runtime = second_state["skill_snapshot"]["runtime_context"]
    assert first_runtime["knowledge_base"] != second_runtime["knowledge_base"]
    assert first_runtime["risk_profile"] != second_runtime["risk_profile"]
    assert str(project_root) not in json.dumps(first_runtime, ensure_ascii=False)

    first_download = client.get(f"/api/download/{first_job}/excel")
    assert first_download.status_code == 200
    assert first_download.headers["x-professional-skill-id"] == first_id


def test_legacy_and_explicit_skill_100_row_outputs_are_identical(tmp_path, monkeypatch):
    sample_path = main_module._find_demo_sample_path()
    if not sample_path:
        pytest.skip("当前工作区缺少 100 行标准输入样例")
    runtime_dir = tmp_path / "runtime"
    monkeypatch.setattr(main_module, "RUNTIME_DIR", runtime_dir)
    sheet_configs = json.dumps(main_module._demo_sample_sheet_configs(sample_path), ensure_ascii=False)
    client = TestClient(app)

    legacy_response = _post_process(
        client,
        sample_path,
        sheet_configs=sheet_configs,
        defer_matching="false",
    )
    explicit_response = _post_process(
        client,
        sample_path,
        sheet_configs=sheet_configs,
        defer_matching="false",
        skill_id=ACTIVE_ID,
        skill_version="1.0.0",
    )

    assert legacy_response.status_code == 200, legacy_response.text
    assert explicit_response.status_code == 200, explicit_response.text
    legacy_payload = legacy_response.json()
    explicit_payload = explicit_response.json()
    for field in (
        "total_data_rows",
        "filled_rows",
        "matched_rows",
        "unchanged_rows",
        "review_rows",
        "conflict_rows",
        "physical_matched_rows",
        "physical_experience_rows",
        "physical_review_rows",
        "technical_matched_rows",
        "technical_experience_rows",
        "technical_review_rows",
    ):
        assert explicit_payload["summary"][field] == legacy_payload["summary"][field]
    assert explicit_payload["professional_skill"]["manifest_hash"] == legacy_payload["professional_skill"]["manifest_hash"]

    legacy_excel = next((runtime_dir / legacy_payload["job_id"]).glob("【输出】-控制价计算表-*.xlsx"))
    explicit_excel = next((runtime_dir / explicit_payload["job_id"]).glob("【输出】-控制价计算表-*.xlsx"))
    legacy_book = load_workbook(legacy_excel, data_only=False)
    explicit_book = load_workbook(explicit_excel, data_only=False)
    try:
        assert explicit_book.sheetnames == legacy_book.sheetnames
        for sheet_name in legacy_book.sheetnames:
            legacy_sheet = legacy_book[sheet_name]
            explicit_sheet = explicit_book[sheet_name]
            assert (explicit_sheet.max_row, explicit_sheet.max_column) == (legacy_sheet.max_row, legacy_sheet.max_column)
            for row in legacy_sheet.iter_rows():
                for legacy_cell in row:
                    explicit_cell = explicit_sheet.cell(legacy_cell.row, legacy_cell.column)
                    assert explicit_cell.value == legacy_cell.value, f"{sheet_name}!{legacy_cell.coordinate}"
    finally:
        legacy_book.close()
        explicit_book.close()


@pytest.mark.parametrize(
    ("skill_id", "skill_version", "expected_code"),
    [
        ("unknown-skill", "1.0.0", "skill_not_found"),
        ("general-service-cost-estimation", "0.1.0", "skill_not_available"),
        (ACTIVE_ID, "9.9.9", "skill_version_mismatch"),
    ],
)
def test_process_explicit_invalid_skill_never_silently_falls_back(tmp_path, skill_id, skill_version, expected_code):
    _, input_path = _write_workbooks(tmp_path)
    response = _post_process(TestClient(app), input_path, skill_id=skill_id, skill_version=skill_version)

    assert response.status_code in {404, 409}
    assert response.json()["detail"]["code"] == expected_code
    assert str(PROJECT_ROOT) not in response.text
