from __future__ import annotations

import hashlib
import json
from pathlib import Path

import pytest
from docx import Document
from fastapi.testclient import TestClient
from openpyxl import load_workbook

from app import settlement_audit_api
from app.main import app
from app.settlement_audit import SettlementAuditEngine, SettlementAuditError
from tools.build_settlement_audit_sample import TEMPLATE_PATH, build_sample


PROJECT_ROOT = Path(__file__).resolve().parents[2]
SAMPLE_PATH = (
    PROJECT_ROOT
    / "00-PRD"
    / "01-模块PRD"
    / "10-结算审核助手模块"
    / "evals"
    / "结算审核演示样例.xlsx"
)
SETTLEMENT_UI_PATH = (
    PROJECT_ROOT
    / "frontend"
    / "src"
    / "components"
    / "settlement-audit"
    / "SettlementAuditWorkbench.tsx"
)
SETTLEMENT_UI_CSS_PATH = SETTLEMENT_UI_PATH.with_name("settlement-audit.css")


def _sha256(path: Path) -> str:
    digest = hashlib.sha256()
    with path.open("rb") as stream:
        for chunk in iter(lambda: stream.read(1024 * 1024), b""):
            digest.update(chunk)
    return digest.hexdigest()


def _all_docx_text(path: Path) -> str:
    document = Document(path)
    text = [paragraph.text for paragraph in document.paragraphs]
    for table in document.tables:
        for row in table.rows:
            text.extend(cell.text for cell in row.cells)
    return "\n".join(text)


def test_sample_builder_and_engine_hit_expected_rules_without_touching_template(tmp_path):
    template_hash_before = _sha256(TEMPLATE_PATH)
    sample_path = build_sample(tmp_path / "结算审核演示样例.xlsx")
    sample_hash_before = _sha256(sample_path)

    result = SettlementAuditEngine(TEMPLATE_PATH).review(
        sample_path,
        tmp_path / "result",
        project_name="专项测试项目",
    )

    assert _sha256(TEMPLATE_PATH) == template_hash_before
    assert _sha256(sample_path) == sample_hash_before
    assert result["project_name"] == "专项测试项目"
    assert result["summary"] == {
        "sheet_count": 3,
        "sheet_counts": {"measure": 1, "survey": 1, "other": 1},
        "audited_rows": 8,
        "passed_rows": 1,
        "risk_count": 8,
        "high_risk_count": 7,
        "manual_review_count": 6,
        "reported_detail_total": 461877.02,
        "reviewed_detail_total": 429505.02,
        "suggested_difference": 32372.0,
    }
    observed = {(risk["rule_id"], risk["coordinate"]) for risk in result["risks"]}
    assert observed == {
        ("JS-001", "H11"),
        ("JS-001", "I24"),
        ("JS-002", "M37"),
        ("JS-003", "I40"),
        ("JS-004", "M200"),
        ("JS-006", "L4"),
        ("JS-007", "M5"),
        ("JS-008", "M4"),
    }


def test_settlement_ui_uses_flat_daweiba_shadcn_tokens():
    css = SETTLEMENT_UI_CSS_PATH.read_text(encoding="utf-8")
    component = SETTLEMENT_UI_PATH.read_text(encoding="utf-8")

    assert "gradient" not in css.lower()
    assert "--sa-blue: var(--dws-color-primary)" in css
    assert "background: var(--dws-color-background)" in css
    assert "box-shadow: inset" not in css
    assert 'className="settlement-audit dabawei-shadcn-ui"' in component


def test_aerial_composite_price_does_not_charge_technical_fee(tmp_path):
    input_path = tmp_path / "航测专项规则样例.xlsx"
    workbook = load_workbook(SAMPLE_PATH, data_only=False)
    measure = workbook["表1 测量费，注所有单体都放一张表中"]
    measure["C11"] = "航测综合单价"
    workbook.save(input_path)
    workbook.close()

    output_dir = tmp_path / "result"
    result = SettlementAuditEngine(TEMPLATE_PATH).review(input_path, output_dir)
    risk = next(item for item in result["risks"] if item["rule_id"] == "JS-005")

    assert risk["coordinate"] == "K11"
    assert risk["current_value"] == pytest.approx(0.22)
    assert risk["suggested_value"] == 0
    reviewed = load_workbook(output_dir / result["artifacts"]["excel"], data_only=False)
    assert reviewed["表1 测量费，注所有单体都放一张表中"]["AA11"].value == 0
    reviewed.close()


def test_reviewed_workbook_preserves_inputs_and_writes_audit_columns(tmp_path):
    output_dir = tmp_path / "result"
    result = SettlementAuditEngine(TEMPLATE_PATH).review(SAMPLE_PATH, output_dir)
    reviewed_path = output_dir / result["artifacts"]["excel"]
    workbook = load_workbook(reviewed_path, data_only=False)

    assert workbook.sheetnames[0] == "审核结果"
    assert "演示样例说明" in workbook.sheetnames
    measure = workbook["表1 测量费，注所有单体都放一张表中"]
    survey = workbook["表2-1 XX勘察费，注每个单体一张表"]
    other = workbook["表3 其他费用(一次报送)"]

    assert measure["H11"].value == 3500
    assert measure["I24"].value == 2
    assert measure["M37"].value == 5000
    assert measure["X11"].value == 6
    assert measure["Y11"].value == pytest.approx(0.6)
    assert measure["Z11"].value == pytest.approx(3203)
    assert measure["AA11"].value == pytest.approx(0.22)
    assert measure["AB11"].value == pytest.approx(14067.576)
    assert "统一模板建议值调整" in measure["AC11"].value
    assert measure["Z11"].comment is not None

    assert survey["Z40"].value == 20
    assert survey["AA40"].value == pytest.approx(6761.3761732608)
    assert survey["AC40"].value == pytest.approx(1.2)
    assert "取消额外1.2" in survey["AE40"].value
    assert survey["AC200"].value == pytest.approx(1.2)
    assert other["X4"].value == pytest.approx(1300)
    assert "文件编号缺失" in other["Y5"].value
    assert workbook["审核结果"]["A1"].value == "勘察测量结算辅助审核结果"
    workbook.close()


def test_word_report_and_json_are_valid_and_state_human_boundary(tmp_path):
    output_dir = tmp_path / "result"
    result = SettlementAuditEngine(TEMPLATE_PATH).review(SAMPLE_PATH, output_dir)
    report_path = output_dir / result["artifacts"]["report"]
    result_path = output_dir / result["artifacts"]["result"]

    text = _all_docx_text(report_path)
    assert "勘察测量结算辅助审核报告" in text
    assert "待人工审定" in text
    assert "深度大于 300m" in text
    assert "室内试验技术工作费系数口径异常" in text
    assert "合同与下浮口径" in text
    assert "不替代合同判断" in text
    assert "人工复核签认" in text
    assert "空白不代表默认同意系统建议" in text

    stored = json.loads(result_path.read_text(encoding="utf-8"))
    assert stored["summary"]["risk_count"] == 8
    assert stored["artifacts"]["excel"].startswith("【审核后】")
    assert stored["artifacts"]["report"].startswith("【审核报告】")
    assert not any("Codex" in name or "codex" in name for name in stored["artifacts"].values())


def test_engine_rejects_workbook_missing_required_template_sheet(tmp_path):
    broken_path = tmp_path / "broken.xlsx"
    workbook = load_workbook(SAMPLE_PATH)
    workbook.remove(workbook["表3 其他费用(一次报送)"])
    workbook.save(broken_path)
    workbook.close()

    with pytest.raises(SettlementAuditError, match="其他费用表"):
        SettlementAuditEngine(TEMPLATE_PATH).review(broken_path, tmp_path / "result")


def test_settlement_audit_api_review_and_downloads_are_isolated(tmp_path, monkeypatch):
    runtime_dir = tmp_path / "settlement-audit"
    monkeypatch.setattr(settlement_audit_api, "SETTLEMENT_AUDIT_RUNTIME_DIR", runtime_dir)
    client = TestClient(app)

    profile = client.get("/api/settlement-audit/profile")
    assert profile.status_code == 200
    assert profile.json()["rule_version"] == "1.0.0"
    assert profile.json()["boundary"].startswith("规则辅助审核")

    with SAMPLE_PATH.open("rb") as stream:
        response = client.post(
            "/api/settlement-audit/review",
            files={
                "file": (
                    SAMPLE_PATH.name,
                    stream,
                    "application/vnd.openxmlformats-officedocument.spreadsheetml.sheet",
                )
            },
            data={"project_name": "API 演示项目"},
        )
    assert response.status_code == 200, response.text
    payload = response.json()
    assert payload["project_name"] == "API 演示项目"
    assert payload["summary"]["risk_count"] == 8
    assert len(payload["job_id"]) == 32

    for kind, expected_type in (
        ("excel", "application/vnd.openxmlformats-officedocument.spreadsheetml.sheet"),
        ("report", "application/vnd.openxmlformats-officedocument.wordprocessingml.document"),
        ("result", "application/json"),
    ):
        download = client.get(payload["downloads"][kind])
        assert download.status_code == 200
        assert download.headers["content-type"].startswith(expected_type)
        assert download.content

    assert len(list(runtime_dir.iterdir())) == 1
    assert not (PROJECT_ROOT / "原始上传.xlsx").exists()


def test_settlement_audit_api_rejects_unsupported_or_corrupt_files(tmp_path, monkeypatch):
    monkeypatch.setattr(
        settlement_audit_api,
        "SETTLEMENT_AUDIT_RUNTIME_DIR",
        tmp_path / "settlement-audit",
    )
    client = TestClient(app)

    unsupported = client.post(
        "/api/settlement-audit/review",
        files={"file": ("旧版结算.xls", b"not-an-xls", "application/vnd.ms-excel")},
    )
    assert unsupported.status_code == 400
    assert ".xlsx" in unsupported.json()["detail"]

    corrupt = client.post(
        "/api/settlement-audit/review",
        files={"file": ("损坏结算.xlsx", b"not-an-xlsx", "application/octet-stream")},
    )
    assert corrupt.status_code == 400
    assert "无法读取结算工作簿" in corrupt.json()["detail"]

    monkeypatch.setattr(settlement_audit_api, "MAX_SETTLEMENT_FILE_BYTES", 8)
    oversized = client.post(
        "/api/settlement-audit/review",
        files={"file": ("超限结算.xlsx", b"123456789", "application/octet-stream")},
    )
    assert oversized.status_code == 413
    assert "文件超过" in oversized.json()["detail"]

    missing = client.get("/api/settlement-audit/download/not-a-job/excel")
    assert missing.status_code == 404

    job_id = "a" * 32
    job_dir = settlement_audit_api.SETTLEMENT_AUDIT_RUNTIME_DIR / job_id
    job_dir.mkdir(parents=True)
    (tmp_path / "outside.xlsx").write_bytes(b"outside")
    (job_dir / "审核结果.json").write_text(
        json.dumps({"artifacts": {"excel": "../../outside.xlsx"}}),
        encoding="utf-8",
    )
    escaped = client.get(f"/api/settlement-audit/download/{job_id}/excel")
    assert escaped.status_code == 404
    assert "成果不存在" in escaped.json()["detail"]
