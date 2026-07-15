from pathlib import Path
from datetime import date
import re
from io import BytesIO

from docx import Document
from fastapi.testclient import TestClient
from openpyxl import Workbook, load_workbook
import json
import pytest

from app.main import app, _refresh_preview_from_sheet
import app.main as main_module
import app.llm as llm_module
from app.llm import LlmConfig
import app.report as report_module
from app.report import build_report_markdown, write_report
from app.schemas import FillSummary, ReviewRow
from app.experience_warning import EXPERIENCE_POOL_HEADERS, WARNING_DETAIL_FIELD, WARNING_PARAMETER_FIELD

ROOT = Path(__file__).resolve().parents[2]
DATA_DIR = ROOT / "03-知识库-二维数据库制作"


def find_data_file(*tokens: str, exclude: tuple[str, ...] = (), required: bool = True) -> Path | None:
    for path in DATA_DIR.glob("*.xlsx"):
        if path.name.startswith("~$"):
            continue
        if all(token in path.name for token in tokens) and not any(token in path.name for token in exclude):
            return path
    if not required:
        return None
    raise FileNotFoundError(f"未找到测试数据文件：{tokens}")


INPUT_PATH = find_data_file("输入100", "空单价100", exclude=("答案", " 和 "), required=False)


def test_health_endpoint():
    client = TestClient(app)
    response = client.get("/api/health")

    assert response.status_code == 200
    assert response.json()["status"] == "ok"
    assert response.json()["version"] == "v5.8.18"


def test_project_default_settings_include_zhisuan_window():
    client = TestClient(app)
    response = client.get("/api/project-default-settings")

    assert response.status_code == 200
    payload = response.json()
    assert payload["file_path"].replace("\\", "/").endswith("config/project-default-settings.json")
    assert payload["zhisuanWindow"]["dockWidth"] == 400
    assert payload["zhisuanWindow"]["welcomeMessage"]
    assert payload["zhisuanWindow"]["quickSettings"]["autoHide"] is True
    assert payload["zhisuanWindow"]["quickSettings"]["customPrompts"] == ["@知识库："]


def test_ui_preferences_are_saved_and_loaded(tmp_path, monkeypatch):
    preferences_path = tmp_path / "ui-preferences.json"
    monkeypatch.setattr(main_module, "DEFAULT_UI_PREFERENCES_PATH", preferences_path)
    client = TestClient(app)

    response = client.post(
        "/api/ui-preferences",
        json={
            "preferences": {
                "enabled": True,
                "styles": {
                    "hero": {"fontSize": 44, "paddingX": 28, "opacity": 88},
                    "ignored": {"unknown": 999},
                },
                "text": {"hero.title": "管勘智算测试"},
            }
        },
    )

    assert response.status_code == 200
    payload = response.json()
    assert payload["preferences"]["enabled"] is True
    assert payload["preferences"]["styles"]["hero"] == {"fontSize": 44.0, "paddingX": 28.0, "opacity": 88.0}
    assert "ignored" not in payload["preferences"]["styles"]
    assert payload["preferences"]["text"]["hero.title"] == "管勘智算测试"
    assert preferences_path.exists()

    get_response = client.get("/api/ui-preferences")
    assert get_response.status_code == 200
    assert get_response.json()["preferences"]["text"]["hero.title"] == "管勘智算测试"


def test_preview_column_preferences_post_returns_temporary_preferences_without_saving(tmp_path, monkeypatch):
    preferences_path = tmp_path / "preview-column-preferences.json"
    monkeypatch.setattr(main_module, "DEFAULT_PREVIEW_COLUMN_PREFERENCES_PATH", preferences_path)
    client = TestClient(app)

    response = client.post(
        "/api/preview-column-preferences",
        json={
            "preferences": {
                "defaultLabels": ["要素1", "单价", "", "单价"],
                "sheetOverrides": {"表2": ["要素1", "预警参数"], "空表": []},
                "headerRows": {"表2": 4, "错误": 0, "过大": 1200},
                "maxDisplayChars": 99,
                "columnWidths": {"表2": {"#1": 90, "单价": 9999, "": 120}},
            }
        },
    )

    assert response.status_code == 200
    payload = response.json()
    assert payload["preferences"]["defaultLabels"] == ["要素1", "单价"]
    assert payload["preferences"]["sheetOverrides"] == {"表2": ["要素1", "预警参数"]}
    assert payload["preferences"]["headerRows"] == {"表2": 4, "过大": 999}
    assert payload["preferences"]["maxDisplayChars"] == 40
    assert payload["preferences"]["columnWidths"] == {"表2": {"#1": 90, "单价": 420}}
    assert not preferences_path.exists()

    get_response = client.get("/api/preview-column-preferences")
    assert get_response.status_code == 200
    assert get_response.json()["file_path"].replace("\\", "/").endswith("config/project-default-settings.json")


def test_download_excel_can_hide_empty_core_sheet_rows_by_warning_filter_field(tmp_path, monkeypatch):
    job_dir = tmp_path / "job-download"
    job_dir.mkdir()
    workbook_path = job_dir / "【输出】-控制价计算表-test.xlsx"
    workbook = Workbook()
    sheet = workbook.active
    sheet.title = "表2-通用工程测量费用"
    sheet.append(["标题"])
    sheet.append([""])
    sheet.append([""])
    sheet.append(["要素1", "数量", "基价"])
    sheet.append(["空数量", "", 100])
    sheet.append(["零数量", 0, 200])
    sheet.append(["有数量", 3, 300])
    sheet.append(["合计", "", 600])
    sheet.append(["合计（含税）", 0, 636])
    sheet3 = workbook.create_sheet("表3-地质测绘")
    sheet3.append(["标题"])
    sheet3.append([""])
    sheet3.append([""])
    sheet3.append(["要素1", "工程量", "基价"])
    sheet3.append(["空工程量", "", 100])
    sheet3.append(["有工程量", 2, 200])
    summary = workbook.create_sheet("费用汇总")
    summary.append(["标题"])
    summary.append([""])
    summary.append([""])
    summary.append(["要素1", "数量", "费用"])
    summary.append(["非核心表不隐藏", "", 1])
    workbook.save(workbook_path)
    workbook.close()

    monkeypatch.setattr(main_module, "RUNTIME_DIR", tmp_path)
    client = TestClient(app)
    response = client.get("/api/download/job-download/excel?hide_empty_rows=true&value_filter_field=数量")

    assert response.status_code == 200
    assert "codex" not in response.headers["content-disposition"].lower()
    downloaded = load_workbook(BytesIO(response.content))
    try:
        assert downloaded["表2-通用工程测量费用"].row_dimensions[5].hidden is True
        assert downloaded["表2-通用工程测量费用"].row_dimensions[6].hidden is True
        assert downloaded["表2-通用工程测量费用"].row_dimensions[7].hidden is False
        assert downloaded["表2-通用工程测量费用"].row_dimensions[8].hidden is False
        assert downloaded["表2-通用工程测量费用"].row_dimensions[9].hidden is False
        assert downloaded["表3-地质测绘"].row_dimensions[5].hidden is True
        assert downloaded["表3-地质测绘"].row_dimensions[6].hidden is False
        assert downloaded["费用汇总"].row_dimensions[5].hidden is False
    finally:
        downloaded.close()

    original = load_workbook(workbook_path)
    try:
        assert original["表2-通用工程测量费用"].row_dimensions[5].hidden is False
    finally:
        original.close()


def test_input_field_preferences_affect_inspect_suggestions(tmp_path, monkeypatch):
    preference_path = tmp_path / "input-field-preferences.json"
    monkeypatch.setattr(main_module, "DEFAULT_INPUT_FIELD_PREFERENCES_PATH", preference_path)
    client = TestClient(app)

    save_response = client.post(
        "/api/input/field-preferences",
        json={"preferences": {"输出-价格列": ["金额列"]}},
    )

    assert save_response.status_code == 200
    assert save_response.json()["preferences"]["输出-价格列"] == ["金额列"]
    assert not preference_path.exists()

    workbook_path = tmp_path / "input.xlsx"
    workbook = Workbook()
    sheet = workbook.active
    sheet.title = "表2-测试"
    sheet.append(["项目名称", "工作内容", "计量单位", "金额列"])
    sheet.append(["岩土工程勘察", "地质测绘", "km2", ""])
    workbook.save(workbook_path)

    with workbook_path.open("rb") as handle:
        inspect_response = client.post(
            "/api/inspect",
            data={"field_preferences": json.dumps({"输出-价格列": ["金额列"]}, ensure_ascii=False)},
            files={"file": (workbook_path.name, handle, "application/vnd.openxmlformats-officedocument.spreadsheetml.sheet")},
        )

    assert inspect_response.status_code == 200
    payload = inspect_response.json()
    assert payload["suggested_mapping"]["要素1"] == "A"
    assert payload["suggested_mapping"]["单位"] == "C"
    assert payload["suggested_mapping"]["输出-价格列"] == "D"


def test_report_template_path_uses_current_locked_template():
    assert report_module.REPORT_TEMPLATE_PATH == (
        DATA_DIR
        / "01-报告模板-招标控制价报告模板"
        / "【模板勿动】控制价报告模板-yyyy-mm-dd.docx"
    )
    assert report_module.REPORT_TEMPLATE_PATH.exists()


def test_refresh_preview_uses_sequential_rows_instead_of_random_cell_access():
    class FakeSheet:
        title = "表2-通用工程测量费用"
        max_row = 4
        max_column = 2000

        def iter_rows(self, min_row, max_row, max_col=None, values_only=False):
            rows = {
                1: ["要素1", "要素2", "基价"],
                2: ["地下管线测量", "地下管线测量", 1700],
                3: ["地质测绘", "1:500", 655],
                4: ["工程勘察", "钻孔", 711],
            }
            for row_index in range(min_row, max_row + 1):
                row = rows[row_index]
                if max_col is not None:
                    row = row[:max_col]
                yield row

        def cell(self, row, column):
            raise AssertionError("预览刷新不应在 read_only 工作表上逐格随机读取")

    preview = {
        "sheet_name": "表2-通用工程测量费用",
        "headers": ["要素1", "要素2", "基价"],
        "rows": [["旧", "旧", "旧"]],
    }

    refreshed = _refresh_preview_from_sheet(FakeSheet(), preview)

    assert refreshed["headers"] == ["要素1", "要素2", "基价"]
    assert refreshed["rows"] == [["地下管线测量", "地下管线测量", 1700]]


def test_refresh_table_preview_resolves_missing_formula_cache(tmp_path):
    workbook_path = tmp_path / "formula-preview.xlsx"
    workbook = Workbook()
    sheet = workbook.active
    sheet.title = "费用汇总"
    sheet.append(["项目", "费用（万元）"])
    sheet.append(["通用工程测量费用", 880.4632])
    sheet.append(["通用工程勘察费用", 1890.2881])
    sheet.append(["小计", "=SUM(B2:B3)"])
    sheet.append(["浮动后勘察费", "=ROUND(B4*0.9,4)"])
    workbook.save(workbook_path)
    workbook.close()

    preview = {
        "sheet_name": "费用汇总",
        "headers": ["项目", "费用（万元）"],
        "rows": [["旧", "旧"], ["旧", "旧"], ["旧", "旧"], ["旧", "旧"]],
    }

    refreshed = main_module._refresh_table_preview_from_output(preview, workbook_path)

    assert refreshed["rows"][2][1] == pytest.approx(2770.7513)
    assert refreshed["rows"][3][1] == pytest.approx(2493.6762)


def test_refresh_table_preview_preserves_selected_header_row(tmp_path):
    workbook_path = tmp_path / "header-row-preview.xlsx"
    workbook = Workbook()
    sheet = workbook.active
    sheet.title = "表2-通用工程测量费用"
    sheet.append(["标题", "", "", ""])
    sheet.append(["要素1", "要素2", "基价", "草稿表头"])
    sheet.append(["要素1", "要素2", "基价", "正式表头"])
    sheet.append(["地下管线测量", "管线探测", 1700, "正式数据"])
    workbook.save(workbook_path)
    workbook.close()

    preview = {
        "sheet_name": "表2-通用工程测量费用",
        "header_row": 3,
        "headers": ["要素1", "要素2", "基价", "正式表头"],
        "rows": [["旧", "旧", "旧", "旧"]],
    }

    refreshed = main_module._refresh_table_preview_from_output(preview, workbook_path)

    assert refreshed["header_row"] == 3
    assert refreshed["headers"] == ["要素1", "要素2", "基价", "正式表头"]
    assert refreshed["rows"][0] == ["地下管线测量", "管线探测", 1700, "正式数据"]


def test_preview_refresh_endpoint_uses_saved_sheet_header_rows(tmp_path, monkeypatch):
    job_id = "job-preview-header-row"
    job_dir = tmp_path / job_id
    job_dir.mkdir()
    workbook_path = job_dir / "output.xlsx"
    workbook = Workbook()
    sheet = workbook.active
    sheet.title = "表2-通用工程测量费用"
    sheet.append(["标题", "", ""])
    sheet.append(["要素1", "基价", "草稿列"])
    sheet.append(["要素1", "基价", "正式列"])
    sheet.append(["地下管线测量", 1700, "正式数据"])
    workbook.save(workbook_path)
    workbook.close()

    summary = FillSummary(
        total_data_rows=1,
        price_column="基价",
        filled_rows=1,
        matched_rows=1,
        unchanged_rows=0,
        review_rows=0,
        conflict_rows=0,
        output_excel=workbook_path.name,
        output_report="report.docx",
        report_text="",
        table_preview={
            "sheet_name": "表2-通用工程测量费用",
            "header_row": 2,
            "headers": ["要素1", "基价", "草稿列"],
            "rows": [["旧", "旧", "旧"]],
        },
        review_details=[],
        price_logs=[],
    )
    main_module._save_process_state(job_dir, "input.xlsx", None, workbook_path, job_dir / "report.docx", summary)
    monkeypatch.setattr(main_module, "RUNTIME_DIR", tmp_path)

    client = TestClient(app)
    response = client.post(
        "/api/preview/refresh",
        json={"job_id": job_id, "header_rows": {"表2-通用工程测量费用": 3}},
    )

    assert response.status_code == 200
    preview = response.json()["summary"]["table_preview"]
    assert preview["header_row"] == 3
    assert preview["headers"] == ["要素1", "基价", "正式列"]
    assert preview["rows"][0] == ["地下管线测量", 1700, "正式数据"]
    assert preview["row_numbers"] == [4]


def test_preview_cell_edit_updates_output_workbook_preview_and_log(tmp_path, monkeypatch):
    job_id = "job-preview-edit"
    job_dir = tmp_path / job_id
    job_dir.mkdir()
    workbook_path = job_dir / "output.xlsx"
    workbook = Workbook()
    sheet = workbook.active
    sheet.title = "费用汇总"
    sheet.append(["项目", "费用（万元）"])
    sheet.append(["通用工程测量费用", 100])
    sheet.append(["通用工程勘察费用", 200])
    sheet.append(["合计（不含税）", "=SUM(B2:B3)"])
    workbook.save(workbook_path)
    workbook.close()
    report_path = job_dir / "【输出】-控制价报告-test.docx"

    summary = FillSummary(
        total_data_rows=2,
        price_column="费用（万元）",
        filled_rows=2,
        matched_rows=2,
        unchanged_rows=0,
        review_rows=0,
        conflict_rows=0,
        output_excel=workbook_path.name,
        output_report=report_path.name,
        report_text="",
        table_preview={
            "sheet_name": "费用汇总",
            "header_row": 1,
            "headers": ["项目", "费用（万元）"],
            "rows": [["通用工程测量费用", 100], ["通用工程勘察费用", 200], ["合计（不含税）", ""]],
        },
        review_details=[],
        price_logs=[],
    )
    main_module._save_process_state(job_dir, "input.xlsx", None, workbook_path, report_path, summary)
    monkeypatch.setattr(main_module, "RUNTIME_DIR", tmp_path)

    client = TestClient(app)
    response = client.post(
        "/api/preview/cell",
        json={
            "job_id": job_id,
            "sheet_name": "费用汇总",
            "row_number": 2,
            "column_number": 2,
            "value": "150.5",
            "header_rows": {"费用汇总": 1},
        },
    )

    assert response.status_code == 200
    payload = response.json()
    assert payload["manual_edit"]["original_value"] == 100
    assert payload["manual_edit"]["new_value"] == 150.5
    assert payload["needs_recalculate"] is True
    assert payload["summary"]["table_preview"]["rows"][0][1] == 150.5
    assert payload["summary"]["table_preview"]["rows"][2][1] == ""

    edited = load_workbook(workbook_path)
    try:
        cell = edited["费用汇总"]["B2"]
        assert cell.value == 150.5
        assert str(cell.fill.fgColor.rgb).endswith("DDEBFF")
        assert cell.comment is not None
        assert "人工修改" in cell.comment.text
    finally:
        edited.close()

    log_path = job_dir / main_module.MANUAL_EDIT_LOG_FILENAME
    records = json.loads(log_path.read_text(encoding="utf-8"))
    assert records[0]["sheet"] == "费用汇总"
    assert records[0]["column_letter"] == "B"

    recalculate_response = client.post(
        "/api/preview/recalculate",
        json={"job_id": job_id, "header_rows": {"费用汇总": 1}},
    )
    assert recalculate_response.status_code == 200
    recalculated_payload = recalculate_response.json()
    assert recalculated_payload["needs_recalculate"] is False
    assert recalculated_payload["summary"]["table_preview"]["rows"][2][1] == pytest.approx(350.5)

    report_response = client.get(recalculated_payload["downloads"]["report"])
    assert report_response.status_code == 200


def test_preview_cell_edit_rejects_formula_and_merged_non_start_cells(tmp_path, monkeypatch):
    job_id = "job-preview-edit-readonly"
    job_dir = tmp_path / job_id
    job_dir.mkdir()
    workbook_path = job_dir / "output.xlsx"
    workbook = Workbook()
    sheet = workbook.active
    sheet.title = "表2-通用工程测量费用"
    sheet.append(["要素1", "基价", "匹配说明"])
    sheet.append(["地下管线测量", "=SUM(1,2)", "系统说明"])
    sheet.append(["合并首格", None, ""])
    sheet.merge_cells("A3:B3")
    workbook.save(workbook_path)
    workbook.close()

    summary = FillSummary(
        total_data_rows=2,
        price_column="基价",
        filled_rows=1,
        matched_rows=1,
        unchanged_rows=0,
        review_rows=1,
        conflict_rows=0,
        output_excel=workbook_path.name,
        output_report="report.docx",
        report_text="",
        table_preview={
            "sheet_name": "表2-通用工程测量费用",
            "header_row": 1,
            "headers": ["要素1", "基价", "匹配说明"],
            "rows": [["地下管线测量", "", "系统说明"], ["合并首格", "", ""]],
        },
        review_details=[],
        price_logs=[],
    )
    main_module._save_process_state(job_dir, "input.xlsx", None, workbook_path, job_dir / "report.docx", summary)
    monkeypatch.setattr(main_module, "RUNTIME_DIR", tmp_path)

    client = TestClient(app)
    formula_response = client.post(
        "/api/preview/cell",
        json={
            "job_id": job_id,
            "sheet_name": "表2-通用工程测量费用",
            "row_number": 2,
            "column_number": 2,
            "value": "999",
            "header_rows": {"表2-通用工程测量费用": 1},
        },
    )
    merged_response = client.post(
        "/api/preview/cell",
        json={
            "job_id": job_id,
            "sheet_name": "表2-通用工程测量费用",
            "row_number": 3,
            "column_number": 2,
            "value": "999",
            "header_rows": {"表2-通用工程测量费用": 1},
        },
    )
    system_column_response = client.post(
        "/api/preview/cell",
        json={
            "job_id": job_id,
            "sheet_name": "表2-通用工程测量费用",
            "row_number": 2,
            "column_number": 3,
            "value": "人工说明",
            "header_rows": {"表2-通用工程测量费用": 1},
        },
    )

    assert formula_response.status_code == 400
    assert "公式单元格" in formula_response.json()["detail"]
    assert merged_response.status_code == 400
    assert "合并单元格非左上角" in merged_response.json()["detail"]
    assert system_column_response.status_code == 400
    assert "系统生成列" in system_column_response.json()["detail"]


def test_refresh_table_preview_inherits_vertical_merged_cells_only(tmp_path):
    workbook_path = tmp_path / "merged-preview.xlsx"
    workbook = Workbook()
    sheet = workbook.active
    sheet.title = "表2-通用工程测量费用"
    sheet.append(["要素1", "要素2", "数量"])
    sheet.append(["地下管线测量", "纵向首行", 1])
    sheet.append([None, "纵向第二行", 2])
    sheet.append(["合并标题", None, 3])
    sheet.merge_cells("A2:A3")
    sheet.merge_cells("A4:B4")
    workbook.save(workbook_path)
    workbook.close()

    preview = {
        "sheet_name": "表2-通用工程测量费用",
        "headers": ["要素1", "要素2", "数量"],
        "rows": [["旧", "旧", "旧"], ["旧", "旧", "旧"], ["旧", "旧", "旧"]],
    }

    refreshed = main_module._refresh_table_preview_from_output(preview, workbook_path)

    assert refreshed["rows"][0] == ["地下管线测量", "纵向首行", 1]
    assert refreshed["rows"][1] == ["地下管线测量", "纵向第二行", 2]
    assert refreshed["rows"][2] == ["合并标题", "", 3]


def test_cors_allows_current_frontend_port():
    client = TestClient(app)
    response = client.options(
        "/api/process",
        headers={
            "Origin": "http://127.0.0.1:5174",
            "Access-Control-Request-Method": "POST",
        },
    )

    assert response.status_code == 200
    assert response.headers["access-control-allow-origin"] == "http://127.0.0.1:5174"


def test_import_experience_pool_endpoint_writes_temp_pool(tmp_path, monkeypatch):
    source_path = find_data_file("铜梁江津", "v2.3", required=False)
    if source_path is None:
        pytest.skip("缺少项目例子测试文件")
    pool_path = tmp_path / "经验池.xlsx"
    monkeypatch.setattr(main_module, "DEFAULT_EXPERIENCE_POOL_PATH", pool_path)
    client = TestClient(app)

    with source_path.open("rb") as handle:
        response = client.post(
            "/api/experience-pool/import",
            files={"file": (source_path.name, handle, "application/vnd.openxmlformats-officedocument.spreadsheetml.sheet")},
            data={"selected_fields": json.dumps(["实物工作费调整系数", "技术工作费调整系数"])},
        )

    assert response.status_code == 200
    payload = response.json()
    assert pool_path.exists()
    assert payload["summary"]["imported_rows"] > 0
    assert payload["pool_file"] == str(pool_path)


def test_experience_pool_inspect_endpoint_returns_all_sheets_and_mapping(tmp_path):
    input_path = tmp_path / "experience-inspect.xlsx"
    workbook = Workbook()
    sheet = workbook.active
    sheet.title = "乱版经验表"
    sheet.append(["标题"])
    sheet.append(["序号", "项目名称", "计量单位", "综合单价", "工程数量", "备注A"])
    sheet.append(["", "要素1", "单位", "基价", "工程量", "原表备注1"])
    sheet.append([1, "控制测量", "个", 4274, 3, "批注"])
    other = workbook.create_sheet("其他sheet")
    other.append(["要素1", "单位", "基价"])
    workbook.save(input_path)
    workbook.close()

    client = TestClient(app)
    with input_path.open("rb") as handle:
        response = client.post(
            "/api/experience-pool/inspect",
            files={
                "file": (
                    input_path.name,
                    handle,
                    "application/vnd.openxmlformats-officedocument.spreadsheetml.sheet",
                )
            },
        )

    assert response.status_code == 200
    payload = response.json()
    assert [sheet["sheet_name"] for sheet in payload["sheets"]] == ["乱版经验表", "其他sheet"]
    first = payload["sheets"][0]
    assert first["enabled"] is True
    assert first["header_row"] == 3
    assert first["suggested_mapping"]["要素1"] == "B"
    assert first["suggested_mapping"]["单位"] == "C"
    assert first["suggested_mapping"]["基价"] == "D"
    assert first["suggested_mapping"]["工程量"] == "E"
    assert first["suggested_mapping"]["原表备注1"] == "F"


def test_experience_pool_field_preferences_are_saved_and_used_for_inspect(tmp_path, monkeypatch):
    preferences_path = tmp_path / "experience-field-preferences.json"
    monkeypatch.setattr(main_module, "DEFAULT_EXPERIENCE_FIELD_PREFERENCES_PATH", preferences_path)

    input_path = tmp_path / "custom-experience-inspect.xlsx"
    workbook = Workbook()
    sheet = workbook.active
    sheet.title = "自定义字段表"
    sheet.append(["项目编码", "造价项目名称", "计价单位", "我的价格经验", "自定义实物系数", "自定义技术系数"])
    sheet.append(["001", "控制测量", "点", 1200, 1.1, 0.22])
    workbook.save(input_path)
    workbook.close()

    client = TestClient(app)
    save_response = client.post(
        "/api/experience-pool/field-preferences",
        json={
            "preferences": {
                "要素1": ["造价项目名称"],
                "单位": ["计价单位"],
                "基价": ["我的价格经验"],
                "实物工作费调整系数": ["自定义实物系数"],
                "技术工作费调整系数": ["自定义技术系数"],
            }
        },
    )

    assert save_response.status_code == 200
    assert preferences_path.exists()
    assert save_response.json()["preferences"]["要素1"] == ["造价项目名称"]

    get_response = client.get("/api/experience-pool/field-preferences")
    assert get_response.status_code == 200
    assert get_response.json()["preferences"]["单位"] == ["计价单位"]

    with input_path.open("rb") as handle:
        inspect_response = client.post(
            "/api/experience-pool/inspect",
            files={
                "file": (
                    input_path.name,
                    handle,
                    "application/vnd.openxmlformats-officedocument.spreadsheetml.sheet",
                )
            },
        )

    assert inspect_response.status_code == 200
    mapping = inspect_response.json()["sheets"][0]["suggested_mapping"]
    assert mapping["要素1"] == "B"
    assert mapping["单位"] == "C"
    assert mapping["基价"] == "D"
    assert mapping["实物工作费调整系数"] == "E"
    assert mapping["技术工作费调整系数"] == "F"


def test_import_experience_pool_endpoint_filters_rows_by_mapped_value_field(tmp_path, monkeypatch):
    source_path = tmp_path / "experience-filter.xlsx"
    workbook = Workbook()
    sheet = workbook.active
    sheet.title = "控制价"
    sheet.append(["项目", "单位", "基价", "工程量"])
    sheet.append(["控制测量", "点", 1200, 5])
    sheet.append(["空模板行", "点", 1300, None])
    sheet.append(["零工程量但有值", "点", 1400, 0])
    workbook.save(source_path)
    workbook.close()

    pool_path = tmp_path / "经验池.xlsx"
    monkeypatch.setattr(main_module, "DEFAULT_EXPERIENCE_POOL_PATH", pool_path)
    client = TestClient(app)
    sheet_configs = [
        {
            "sheet_name": "控制价",
            "enabled": True,
            "header_row": 1,
            "column_mapping": {
                "要素1": "A",
                "单位": "B",
                "基价": "C",
                "工程量": "D",
            },
        }
    ]

    with source_path.open("rb") as handle:
        response = client.post(
            "/api/experience-pool/import",
            data={
                "selected_fields": json.dumps(["基价"]),
                "sheet_configs": json.dumps(sheet_configs),
            },
            files={
                "file": (
                    source_path.name,
                    handle,
                    "application/vnd.openxmlformats-officedocument.spreadsheetml.sheet",
                )
            },
        )

    assert response.status_code == 200
    payload = response.json()
    assert payload["summary"]["imported_rows"] == 2
    assert payload["summary"]["skipped_rows"] == 1

    pool_workbook = load_workbook(pool_path, data_only=True)
    try:
        pool_sheet = pool_workbook.active
        rows = list(pool_sheet.iter_rows(values_only=True))
        headers = list(rows[0])
        element_index = headers.index("要素1")
        assert [row[element_index] for row in rows[1:]] == ["控制测量", "零工程量但有值"]
    finally:
        pool_workbook.close()


def test_import_experience_pool_endpoint_stops_after_many_blank_filter_rows(tmp_path, monkeypatch):
    source_path = tmp_path / "experience-long-template.xlsx"
    workbook = Workbook()
    sheet = workbook.active
    sheet.title = "控制价"
    sheet.append(["项目", "单位", "基价", "工程量"])
    sheet.append(["控制测量", "点", 1200, 5])
    for _ in range(210):
        sheet.append(["模板空行", "点", 1300, None])
    sheet.append(["很后面的数据", "点", 1400, 8])
    workbook.save(source_path)
    workbook.close()

    pool_path = tmp_path / "经验池.xlsx"
    monkeypatch.setattr(main_module, "DEFAULT_EXPERIENCE_POOL_PATH", pool_path)
    client = TestClient(app)
    sheet_configs = [
        {
            "sheet_name": "控制价",
            "enabled": True,
            "header_row": 1,
            "column_mapping": {
                "要素1": "A",
                "单位": "B",
                "基价": "C",
                "工程量": "D",
            },
        }
    ]

    with source_path.open("rb") as handle:
        response = client.post(
            "/api/experience-pool/import",
            data={
                "selected_fields": json.dumps(["基价"]),
                "sheet_configs": json.dumps(sheet_configs),
            },
            files={
                "file": (
                    source_path.name,
                    handle,
                    "application/vnd.openxmlformats-officedocument.spreadsheetml.sheet",
                )
            },
        )

    assert response.status_code == 200
    payload = response.json()
    assert payload["summary"]["imported_rows"] == 1
    assert payload["summary"]["skipped_rows"] == 200


def test_workload_capture_inspect_endpoint_detects_source_sheets(tmp_path):
    input_path = tmp_path / "workload-source.xlsx"
    workbook = Workbook()
    sheet = workbook.active
    sheet.title = "标一测量工作量"
    sheet.append(["项目说明"])
    sheet.append(["不是表头"])
    sheet.append(["项目", "内容", "类别", "比例尺", "单位", "数量", "调整系数", "备注"])
    sheet.append(["控制测量", "平面控制", "一等", "1:500", "km", 12, 1.1, "委托方给定"])
    skipped = workbook.create_sheet("勘察测量工作范围")
    skipped.append(["项目", "单位", "数量"])
    workbook.save(input_path)
    workbook.close()

    client = TestClient(app)
    with input_path.open("rb") as handle:
        response = client.post(
            "/api/workload-capture/inspect",
            data={"role": "source"},
            files={
                "file": (
                    input_path.name,
                    handle,
                    "application/vnd.openxmlformats-officedocument.spreadsheetml.sheet",
                )
            },
        )

    assert response.status_code == 200
    payload = response.json()
    first = payload["sheets"][0]
    assert first["sheet_name"] == "标一测量工作量"
    assert first["enabled"] is True
    assert first["header_row"] == 3
    assert first["suggested_mapping"]["要素1"] == "A"
    assert first["suggested_mapping"]["要素2"] == "B"
    assert first["suggested_mapping"]["单位"] == "E"
    assert first["suggested_mapping"]["数量"] == "F"
    assert first["suggested_mapping"]["实物工作费调整系数"] == "G"
    assert payload["sheets"][1]["enabled"] is False


def test_workload_field_preferences_are_temporary_and_used_for_inspect(tmp_path, monkeypatch):
    preferences_path = tmp_path / "workload-field-preferences.json"
    monkeypatch.setattr(main_module, "DEFAULT_WORKLOAD_FIELD_PREFERENCES_PATH", preferences_path)

    input_path = tmp_path / "custom-workload-inspect.xlsx"
    workbook = Workbook()
    sheet = workbook.active
    sheet.title = "委托方工作量"
    sheet.append(["项目大类", "作业名称", "计量单位", "工程数量", "甲方备注"])
    sheet.append(["控制测量", "平面控制", "km", 12, "委托方给定"])
    workbook.save(input_path)
    workbook.close()

    client = TestClient(app)
    save_response = client.post(
        "/api/workload-capture/field-preferences",
        json={
            "preferences": {
                "要素1": ["项目大类"],
                "要素2": ["作业名称"],
                "单位": ["计量单位"],
                "数量": ["工程数量"],
                "委托方备注": ["甲方备注"],
            }
        },
    )

    assert save_response.status_code == 200
    assert not preferences_path.exists()
    assert save_response.json()["preferences"]["要素1"] == ["项目大类"]
    assert save_response.json()["adjacent_fallback_enabled"] is True
    assert save_response.json()["element_sequence_enabled"] is True

    get_response = client.get("/api/workload-capture/field-preferences")
    assert get_response.status_code == 200
    assert get_response.json()["file_path"].replace("\\", "/").endswith("config/project-default-settings.json")
    assert get_response.json()["adjacent_fallback_enabled"] is True
    assert get_response.json()["element_sequence_enabled"] is True

    with input_path.open("rb") as handle:
        inspect_response = client.post(
            "/api/workload-capture/inspect",
            data={
                "role": "source",
                "field_preferences": json.dumps(
                    {
                        "要素1": ["项目大类"],
                        "要素2": ["作业名称"],
                        "单位": ["计量单位"],
                        "数量": ["工程数量"],
                        "委托方备注": ["甲方备注"],
                    },
                    ensure_ascii=False,
                ),
            },
            files={
                "file": (
                    input_path.name,
                    handle,
                    "application/vnd.openxmlformats-officedocument.spreadsheetml.sheet",
                )
            },
        )

    assert inspect_response.status_code == 200
    mapping = inspect_response.json()["sheets"][0]["suggested_mapping"]
    assert mapping["要素1"] == "A"
    assert mapping["要素2"] == "B"
    assert mapping["单位"] == "C"
    assert mapping["数量"] == "D"
    assert mapping["委托方备注"] == "E"


def test_workload_adjacent_fallback_setting_can_be_disabled_for_inspect(tmp_path, monkeypatch):
    preferences_path = tmp_path / "workload-field-preferences.json"
    monkeypatch.setattr(main_module, "DEFAULT_WORKLOAD_FIELD_PREFERENCES_PATH", preferences_path)

    input_path = tmp_path / "workload-adjacent-disabled.xlsx"
    workbook = Workbook()
    sheet = workbook.active
    sheet.title = "委托方工作量"
    sheet.append(["项目", "未命名列", "类别", "比例尺", "单位", "数量"])
    sheet.append(["控制测量", "平面控制", "一级", "1:500", "点", 1])
    workbook.save(input_path)
    workbook.close()

    client = TestClient(app)

    with input_path.open("rb") as handle:
        inspect_response = client.post(
            "/api/workload-capture/inspect",
            data={
                "role": "source",
                "adjacent_fallback_enabled": "false",
                "element_sequence_enabled": "false",
            },
            files={
                "file": (
                    input_path.name,
                    handle,
                    "application/vnd.openxmlformats-officedocument.spreadsheetml.sheet",
                )
            },
        )

    assert inspect_response.status_code == 200
    mapping = inspect_response.json()["sheets"][0]["suggested_mapping"]
    assert mapping["要素1"] == "A"
    assert mapping["要素2"] == ""
    assert mapping["要素3"] == "C"


def test_workload_target_field_preferences_are_temporary_and_used_for_inspect(tmp_path, monkeypatch):
    preferences_path = tmp_path / "workload-target-field-preferences.json"
    monkeypatch.setattr(main_module, "DEFAULT_WORKLOAD_TARGET_FIELD_PREFERENCES_PATH", preferences_path)

    input_path = tmp_path / "custom-target-inspect.xlsx"
    workbook = Workbook()
    sheet = workbook.active
    sheet.title = "表2 测量"
    sheet.append(["要素1", "单位", "工程量-测试", "实物工作费调整系数-测试", "抓取日志-测试"])
    sheet.append(["控制测量", "km", None, None, None])
    workbook.save(input_path)
    workbook.close()

    client = TestClient(app)
    save_response = client.post(
        "/api/workload-capture/target-field-preferences",
        json={
            "preferences": {
                "数量(信息抓取)": ["工程量-测试"],
                "实物工作费调整系数(信息抓取)": ["实物工作费调整系数-测试"],
                "抓取日志": ["抓取日志-测试"],
            }
        },
    )

    assert save_response.status_code == 200
    assert not preferences_path.exists()
    assert save_response.json()["preferences"]["实物工作费调整系数(信息抓取)"] == ["实物工作费调整系数-测试"]
    assert save_response.json()["adjacent_fallback_enabled"] is True
    assert save_response.json()["element_sequence_enabled"] is False

    get_response = client.get("/api/workload-capture/target-field-preferences")
    assert get_response.status_code == 200
    assert get_response.json()["file_path"].replace("\\", "/").endswith("config/project-default-settings.json")
    assert get_response.json()["adjacent_fallback_enabled"] is True
    assert get_response.json()["element_sequence_enabled"] is False

    with input_path.open("rb") as handle:
        inspect_response = client.post(
            "/api/workload-capture/inspect",
            data={
                "role": "target",
                "field_preferences": json.dumps(
                    {
                        "数量(信息抓取)": ["工程量-测试"],
                        "实物工作费调整系数(信息抓取)": ["实物工作费调整系数-测试"],
                        "抓取日志": ["抓取日志-测试"],
                    },
                    ensure_ascii=False,
                ),
            },
            files={
                "file": (
                    input_path.name,
                    handle,
                    "application/vnd.openxmlformats-officedocument.spreadsheetml.sheet",
                )
            },
        )

    assert inspect_response.status_code == 200
    mapping = inspect_response.json()["sheets"][0]["suggested_mapping"]
    assert mapping["数量(信息抓取)"] == "C"
    assert mapping["实物工作费调整系数(信息抓取)"] == "D"
    assert mapping["抓取日志"] == "E"


def test_workload_capture_run_endpoint_fills_and_downloads_workbooks(tmp_path):
    source_path = tmp_path / "workload-source.xlsx"
    source_book = Workbook()
    source = source_book.active
    source.title = "委托方工作量"
    source.append(["要素1", "要素2", "单位", "数量", "实物工作费调整系数", "技术工作费调整系数", "委托方备注"])
    source.append(["控制测量", "平面控制", "km", 12, 1.1, 0.8, "委托方给定"])
    source_book.save(source_path)
    source_book.close()

    target_path = tmp_path / "target.xlsx"
    target_book = Workbook()
    target = target_book.active
    target.title = "表2 测量"
    target.append(["要素1", "要素2", "单位", "数量(信息抓取)", "实物工作费调整系数(信息抓取)", "技术工作费调整系数(信息抓取)"])
    target.append(["控制测量", "平面控制", "km", None, None, None])
    target_book.save(target_path)
    target_book.close()

    source_configs = [
        {
            "sheet_name": "委托方工作量",
            "enabled": True,
            "header_row": 1,
            "column_mapping": {
                "要素1": "A",
                "要素2": "B",
                "单位": "C",
                "数量": "D",
                "实物工作费调整系数": "E",
                "技术工作费调整系数": "F",
                "委托方备注": "G",
            },
        }
    ]
    target_configs = [
        {
            "sheet_name": "表2 测量",
            "enabled": True,
            "header_row": 1,
                "column_mapping": {
                    "要素1": "A",
                    "要素2": "B",
                    "单位": "C",
                    "数量(信息抓取)": "D",
                    "实物工作费调整系数(信息抓取)": "E",
                    "技术工作费调整系数(信息抓取)": "F",
                },
            }
        ]

    client = TestClient(app)
    with source_path.open("rb") as source_handle, target_path.open("rb") as target_handle:
        response = client.post(
            "/api/workload-capture/run",
            data={
                "selected_fields": json.dumps(["数量(信息抓取)", "实物工作费调整系数(信息抓取)", "技术工作费调整系数(信息抓取)"], ensure_ascii=False),
                "source_sheet_configs": json.dumps(source_configs, ensure_ascii=False),
                "target_sheet_configs": json.dumps(target_configs, ensure_ascii=False),
            },
            files={
                "workload_file": (
                    source_path.name,
                    source_handle,
                    "application/vnd.openxmlformats-officedocument.spreadsheetml.sheet",
                ),
                "target_file": (
                    target_path.name,
                    target_handle,
                    "application/vnd.openxmlformats-officedocument.spreadsheetml.sheet",
                ),
            },
        )

    assert response.status_code == 200
    payload = response.json()
    assert payload["summary"]["filled_rows"] == 1
    assert payload["summary"]["warning_rows"] == 0
    assert re.fullmatch(r"【临时】-控制价计算表（填好数量后）-\d{8}-\d{4}\.xlsx", payload["summary"]["output_target"])
    assert re.fullmatch(r"【临时】-原表-\(工作量信息抓取后标注符合用\)-\d{8}-\d{4}\.xlsx", payload["summary"]["output_workload"])

    target_response = client.get(payload["downloads"]["target"])
    workload_response = client.get(payload["downloads"]["workload"])
    assert target_response.status_code == 200
    assert workload_response.status_code == 200
    filled_path = tmp_path / "filled-target.xlsx"
    filled_path.write_bytes(target_response.content)
    filled_book = load_workbook(filled_path, data_only=True)
    filled_sheet = filled_book["表2 测量"]
    headers = [cell.value for cell in filled_sheet[1]]
    assert filled_sheet.cell(row=2, column=headers.index("数量(信息抓取)") + 1).value == 12
    assert filled_sheet.cell(row=2, column=headers.index("实物工作费调整系数(信息抓取)") + 1).value == 1.1
    assert filled_sheet.cell(row=2, column=headers.index("技术工作费调整系数(信息抓取)") + 1).value == 0.8
    assert "抓取成功" in filled_sheet.cell(row=2, column=headers.index("抓取日志") + 1).value
    filled_book.close()


def test_workload_capture_apply_to_current_updates_preview_and_keeps_marked_source_download(tmp_path, monkeypatch):
    runtime_dir = tmp_path / "runtime"
    monkeypatch.setattr(main_module, "RUNTIME_DIR", runtime_dir)
    job_id = "current-workload-job"
    job_dir = runtime_dir / job_id
    job_dir.mkdir(parents=True)

    source_path = tmp_path / "workload-source-current.xlsx"
    source_book = Workbook()
    source = source_book.active
    source.title = "委托方工作量"
    source.append(["要素1", "要素2", "单位", "数量", "实物工作费调整系数", "技术工作费调整系数", "委托方备注"])
    source.append(["控制测量", "平面控制", "km", 12, 1.1, 0.8, "委托方给定"])
    source_book.save(source_path)
    source_book.close()

    output_path = job_dir / "output.xlsx"
    output_book = Workbook()
    output = output_book.active
    output.title = "表2 测量"
    output.append(["要素1", "要素2", "单位", "数量(信息抓取)", "实物工作费调整系数(信息抓取)", "技术工作费调整系数(信息抓取)", "委托方备注(信息抓取)"])
    output.append(["控制测量", "平面控制", "km", None, None, None, None])
    output_book.save(output_path)
    output_book.close()

    summary = FillSummary(
        total_data_rows=1,
        price_column="",
        filled_rows=0,
        matched_rows=0,
        unchanged_rows=0,
        review_rows=0,
        conflict_rows=0,
        output_excel=output_path.name,
        output_report="report.docx",
        report_text="测试报告",
        table_preview={
            "sheet_name": "表2 测量",
            "header_row": 1,
            "headers": ["要素1", "要素2", "单位", "数量(信息抓取)"],
            "rows": [["控制测量", "平面控制", "km", None]],
            "row_numbers": [2],
        },
    )
    main_module._save_process_state(job_dir, "input.xlsx", None, output_path, job_dir / "report.docx", summary)

    source_configs = [
        {
            "sheet_name": "委托方工作量",
            "enabled": True,
            "header_row": 1,
            "column_mapping": {
                "要素1": "A",
                "要素2": "B",
                "单位": "C",
                "数量": "D",
                "实物工作费调整系数": "E",
                "技术工作费调整系数": "F",
                "委托方备注": "G",
            },
        }
    ]
    target_configs = [
        {
            "sheet_name": "表2 测量",
            "enabled": True,
            "header_row": 1,
            "column_mapping": {
                "要素1": "A",
                "要素2": "B",
                "单位": "C",
                "数量(信息抓取)": "D",
                "实物工作费调整系数(信息抓取)": "E",
                "技术工作费调整系数(信息抓取)": "F",
                "委托方备注(信息抓取)": "G",
            },
        }
    ]

    client = TestClient(app)
    inspect_response = client.post("/api/workload-capture/inspect-current-target", data={"job_id": job_id})
    assert inspect_response.status_code == 200
    assert inspect_response.json()["sheets"][0]["suggested_mapping"]["数量(信息抓取)"] == "D"

    with source_path.open("rb") as source_handle:
        response = client.post(
            "/api/workload-capture/apply-to-current",
            data={
                "job_id": job_id,
                "selected_fields": json.dumps(
                    [
                        "数量(信息抓取)",
                        "实物工作费调整系数(信息抓取)",
                        "技术工作费调整系数(信息抓取)",
                        "委托方备注(信息抓取)",
                    ],
                    ensure_ascii=False,
                ),
                "source_sheet_configs": json.dumps(source_configs, ensure_ascii=False),
                "target_sheet_configs": json.dumps(target_configs, ensure_ascii=False),
                "write_mode": "conservative",
            },
            files={
                "workload_file": (
                    source_path.name,
                    source_handle,
                    "application/vnd.openxmlformats-officedocument.spreadsheetml.sheet",
                ),
            },
        )

    assert response.status_code == 200
    payload = response.json()
    assert payload["workload_summary"]["filled_rows"] == 1
    assert payload["workload_summary"]["written_cells"] == 4
    assert payload["summary"]["table_preview"]["rows"][0][3] == 12
    marked_response = client.get(payload["workload_downloads"]["workload"])
    assert marked_response.status_code == 200

    output_book = load_workbook(output_path, data_only=True)
    try:
        sheet = output_book["表2 测量"]
        assert sheet["D2"].value == 12
        assert sheet["E2"].value == 1.1
        assert sheet["F2"].value == 0.8
        assert sheet["G2"].value == "委托方给定"
        assert "抓取成功" in sheet["H2"].value
    finally:
        output_book.close()


def test_workload_capture_run_endpoint_filters_rows_by_selected_field(tmp_path):
    source_path = tmp_path / "workload-source-filter.xlsx"
    source_book = Workbook()
    source = source_book.active
    source.title = "委托方工作量"
    source.append(["要素1", "要素2", "单位", "数量", "实物工作费调整系数"])
    source.append(["控制测量", "平面控制", "km", 12, 1.1])
    source.append(["控制测量", "导线测量", "km", None, 1.3])
    source_book.save(source_path)
    source_book.close()

    target_path = tmp_path / "target-filter.xlsx"
    target_book = Workbook()
    target = target_book.active
    target.title = "表2 测量"
    target.append(["要素1", "要素2", "单位"])
    target.append(["控制测量", "平面控制", "km"])
    target.append(["控制测量", "导线测量", "km"])
    target_book.save(target_path)
    target_book.close()

    source_configs = [
        {
            "sheet_name": "委托方工作量",
            "enabled": True,
            "header_row": 1,
            "column_mapping": {
                "要素1": "A",
                "要素2": "B",
                "单位": "C",
                "数量": "D",
                "实物工作费调整系数": "E",
            },
        }
    ]
    target_configs = [
        {
            "sheet_name": "表2 测量",
            "enabled": True,
            "header_row": 1,
            "column_mapping": {
                "要素1": "A",
                "要素2": "B",
                "单位": "C",
            },
        }
    ]

    client = TestClient(app)
    with source_path.open("rb") as source_handle, target_path.open("rb") as target_handle:
        response = client.post(
            "/api/workload-capture/run",
            data={
                "selected_fields": json.dumps(["数量(信息抓取)"], ensure_ascii=False),
                "source_sheet_configs": json.dumps(source_configs, ensure_ascii=False),
                "target_sheet_configs": json.dumps(target_configs, ensure_ascii=False),
                "only_capture_rows_with_value": "true",
                "value_filter_field": "数量",
            },
            files={
                "workload_file": (
                    source_path.name,
                    source_handle,
                    "application/vnd.openxmlformats-officedocument.spreadsheetml.sheet",
                ),
                "target_file": (
                    target_path.name,
                    target_handle,
                    "application/vnd.openxmlformats-officedocument.spreadsheetml.sheet",
                ),
            },
        )

    assert response.status_code == 200
    payload = response.json()
    assert payload["summary"]["source_rows"] == 1
    assert payload["summary"]["filled_rows"] == 1

    target_response = client.get(payload["downloads"]["target"])
    assert target_response.status_code == 200
    filled_path = tmp_path / "filtered-target.xlsx"
    filled_path.write_bytes(target_response.content)
    filled_book = load_workbook(filled_path, data_only=True)
    filled_sheet = filled_book["表2 测量"]
    headers = [cell.value for cell in filled_sheet[1]]
    quantity_col = headers.index("数量(信息抓取)") + 1
    log_col = headers.index("抓取日志") + 1
    assert filled_sheet.cell(row=2, column=quantity_col).value == 12
    assert filled_sheet.cell(row=3, column=quantity_col).value is None
    assert "模式A+B均未命中" in filled_sheet.cell(row=3, column=log_col).value
    filled_book.close()


def test_inspect_endpoint_returns_headers_and_default_mapping(tmp_path):
    input_path = tmp_path / "headers.xlsx"
    workbook = Workbook()
    sheet = workbook.active
    sheet.append(["专业", "要素2", "单位", "基价测试列", "要素1", "要素3", "要素4", "要素5"])
    sheet.append(["x", "y", "m", "空单价", "z", "", "a", "b"])
    workbook.save(input_path)

    client = TestClient(app)
    with input_path.open("rb") as handle:
        response = client.post(
            "/api/inspect",
            files={
                "file": (
                    input_path.name,
                    handle,
                    "application/vnd.openxmlformats-officedocument.spreadsheetml.sheet",
                )
            },
        )

    assert response.status_code == 200
    payload = response.json()
    assert payload["header_row"] == 1
    assert payload["headers"] == ["专业", "要素2", "单位", "基价测试列", "要素1", "要素3", "要素4", "要素5"]
    assert payload["columns"][0] == {"letter": "A", "header": "专业", "label": "A列 - 专业"}
    assert payload["suggested_mapping"]["要素1"] == "E"
    assert payload["suggested_mapping"]["要素2"] == "B"
    assert payload["suggested_mapping"]["单位"] == "C"
    assert payload["suggested_mapping"]["价格列"] == "D"


def test_inspect_endpoint_scans_first_four_rows_for_mapping_row(tmp_path):
    input_path = tmp_path / "headers-row3.xlsx"
    workbook = Workbook()
    sheet = workbook.active
    sheet.append(["项目名称", "管勘智算测试"])
    sheet.append(["说明", "前两行不是表头"])
    sheet.append(["价格输入", "专业", "工作项", "比例尺", "复杂程度", "单位列", "空列", "要素1"])
    sheet.append(["空单价", "岩土工程勘察", "地质测绘", "比例-1:500", "复杂", "km2", "", ""])
    workbook.save(input_path)

    client = TestClient(app)
    with input_path.open("rb") as handle:
        response = client.post(
            "/api/inspect",
            files={
                "file": (
                    input_path.name,
                    handle,
                    "application/vnd.openxmlformats-officedocument.spreadsheetml.sheet",
                )
            },
        )

    assert response.status_code == 200
    payload = response.json()
    assert payload["header_row"] == 3
    assert payload["headers"] == ["价格输入", "专业", "工作项", "比例尺", "复杂程度", "单位列", "空列", "要素1"]
    assert payload["suggested_mapping"]["要素1"] == "H"
    assert payload["suggested_mapping"]["价格列"] == "A"


def test_inspect_endpoint_returns_candidate_sheet_configs(tmp_path):
    input_path = tmp_path / "candidate-sheets.xlsx"
    workbook = Workbook()
    sheet = workbook.active
    sheet.title = "说明"
    sheet.append(["说明"])
    table_sheet = workbook.create_sheet("表2 测量")
    table_sheet.append(["标题"])
    table_sheet.append(["说明"])
    table_sheet.append(["说明"])
    table_sheet.append([
        "映射行",
        "要素1",
        "要素2",
        "单位",
        "单价匹配-测试",
        "匹配报告预留位置",
        "实物工作费调整系数",
        "技术工作费调整系数",
    ])
    workbook.create_sheet("费用汇总")
    workbook.save(input_path)

    client = TestClient(app)
    with input_path.open("rb") as handle:
        response = client.post(
            "/api/inspect",
            files={
                "file": (
                    input_path.name,
                    handle,
                    "application/vnd.openxmlformats-officedocument.spreadsheetml.sheet",
                )
            },
        )

    assert response.status_code == 200
    payload = response.json()
    assert [sheet["sheet_name"] for sheet in payload["sheets"]] == ["表2 测量"]
    config = payload["sheets"][0]
    assert config["enabled"] is True
    assert config["header_row"] == 4
    assert config["suggested_mapping"]["要素1"] == "B"
    assert config["suggested_mapping"]["要素3"] == "空元素列"
    assert config["suggested_mapping"]["输出-价格列"] == "E"
    assert config["suggested_mapping"]["输出-实物工作费调整系数"] == "G"
    assert config["suggested_mapping"]["输出-技术工作费调整系数"] == "H"


def test_demo_load_sample_uses_candidate_sheet_configs(tmp_path, monkeypatch):
    data_dir = tmp_path / "03-知识库-二维数据库制作"
    data_dir.mkdir()
    kb_path = data_dir / "【数据库】【导入】.xlsx"
    kb_workbook = Workbook()
    kb_sheet = kb_workbook.active
    kb_sheet.append(["序号", "要素1", "要素2", "要素3", "要素4", "要素5", "单位", "基价"])
    kb_sheet.append([1, "控制测量", "首级控制测量", None, "GPS测量C级", "中等", "个", 123])
    kb_workbook.save(kb_path)

    sample_path = data_dir / "输入100 和 空单价100.xlsx"
    workbook = Workbook()
    summary = workbook.active
    summary.title = "费用汇总"
    summary.append(["油气长输管道工程勘察测量费用测算表"])
    detail = workbook.create_sheet("表2-通用工程测量费用")
    detail.append(["附表2油气长输管道通用工程测量费用测算表"])
    detail.append(["说明"])
    detail.append(["说明"])
    detail.append(["映射行", "要素1", "要素2", "要素3", "要素4", "要素5", "单位", "单价匹配-测试", "数量"])
    detail.append([1, "控制测量", "首级控制测量", None, "GPS测量C级", "中等", "个", None, 1])
    workbook.save(sample_path)

    monkeypatch.setattr(main_module, "DEFAULT_KB_PATH", kb_path)
    monkeypatch.setattr(main_module, "DEFAULT_INPUT_FIELD_PREFERENCES_PATH", tmp_path / "missing-preferences.json")
    monkeypatch.setattr(main_module, "RUNTIME_DIR", tmp_path / "runtime")
    monkeypatch.setattr(main_module, "recalculate_workbook", lambda _path: None)

    client = TestClient(app)
    response = client.post("/api/demo/load-sample")

    assert response.status_code == 200
    payload = response.json()
    assert payload["demo_mode"] is True
    assert payload["sample_file"] == sample_path.name
    assert payload["summary"]["total_data_rows"] == 1
    assert payload["summary"]["filled_rows"] == 1


def test_write_report_creates_docx(tmp_path):
    summary = FillSummary(
        total_data_rows=3,
        price_column="基价测试列",
        filled_rows=2,
        matched_rows=2,
        unchanged_rows=1,
        review_rows=1,
        conflict_rows=0,
        output_excel="filled.xlsx",
        output_report="report.docx",
        report_text="输入3行，匹配2行。",
        table_preview={"headers": ["要素1"], "rows": [["测试"]]},
        price_logs=["第 4 行：待复核，候选数量 0，未找到匹配项。"],
        review_details=[
            ReviewRow(
                excel_row=4,
                status="not_found",
                message="未找到匹配项",
                values={"要素1": "测试"},
            )
        ],
    )
    output_path = tmp_path / "report.docx"

    write_report(output_path, "input.xlsx", summary)

    assert output_path.exists()
    doc = Document(output_path)
    text = "\n".join(p.text for p in doc.paragraphs)
    assert "五、其他需要注意的事项" in text
    assert "待复核提示" in text
    assert "基价/单价匹配 2 行" in text
    assert "第 4 行：未找到匹配项" in text
    markdown_text = output_path.with_suffix(".md").read_text(encoding="utf-8")
    assert "## 价格识别日志" in markdown_text
    assert "第 4 行：待复核" in markdown_text


def test_report_markdown_includes_experience_warning_section():
    summary = FillSummary(
        total_data_rows=1,
        price_column="基价",
        filled_rows=1,
        matched_rows=1,
        unchanged_rows=0,
        review_rows=0,
        conflict_rows=0,
        output_excel="filled.xlsx",
        output_report="report.docx",
        report_text="输入1行，匹配1行。",
        table_preview={"headers": [], "rows": []},
        warning_summary={
            "pool_enabled": True,
            "checked_rows": 1,
            "warning_rows": 1,
            "high_rows": 0,
            "low_rows": 1,
            "medium_rows": 1,
            "summary_text": "经验池预警：输入候选 1 行，可比选 1 行，未找到同类 0 行，发现 1 条预警，其中高风险 0 条、低风险 1 条；匹配模式：字段完全匹配 1 行",
        },
        warning_details=[
            {
                "sheet_name": "表2",
                "excel_row": 5,
                "metric": "基价",
                "current_value": 5000,
                "experience_values": [4274, 4300],
                "experience_average": 4287,
                "experience_min": 4274,
                "experience_max": 4300,
                "experience_range_text": "4274~4300",
                "sample_count": 2,
                "deviation_percent": 16.632844,
                "severity": "low",
                "severity_label": "低风险",
                "message": "基价 当前值 5000；经验池平均值 4287；经验范围 4274~4300；实际偏离率 16.632844%；已触发低风险预警。",
                "suggested_action": "建议复核：基价 当前值 5000 相对经验池平均值 4287 偏离 16.632844%，建议结合项目条件与历史范围 4274~4300 复核。",
                "source_rows": [],
            }
        ],
    )

    markdown = build_report_markdown("input.xlsx", summary)

    assert "## 经验池预警" in markdown
    assert "经验池预警：输入候选 1 行，可比选 1 行，未找到同类 0 行，发现 1 条预警" in markdown
    assert "表2 第 5 行：基价" in markdown


def test_experience_warning_settings_endpoint_validates_and_persists(tmp_path, monkeypatch):
    settings_path = tmp_path / "experience-warning-settings.json"
    monkeypatch.setattr(main_module, "DEFAULT_EXPERIENCE_WARNING_SETTINGS_PATH", settings_path)
    client = TestClient(app)

    get_response = client.get("/api/experience-warnings/settings")
    assert get_response.status_code == 200
    assert get_response.json()["settings"] == {
        "low_risk_warning_ratio": 5.0,
        "high_risk_warning_ratio": 20.0,
        "only_check_rows_with_value": True,
        "value_filter_field": "数量",
    }
    assert get_response.json()["filter_fields"] == ["数量"]

    save_response = client.post(
        "/api/experience-warnings/settings",
        json={"settings": {"low_risk_warning_ratio": 8, "high_risk_warning_ratio": 25, "only_check_rows_with_value": False, "value_filter_field": "数量"}},
    )
    assert save_response.status_code == 200
    assert save_response.json()["settings"] == {
        "low_risk_warning_ratio": 8.0,
        "high_risk_warning_ratio": 25.0,
        "only_check_rows_with_value": False,
        "value_filter_field": "数量",
    }

    stored = json.loads(settings_path.read_text(encoding="utf-8"))
    assert stored["settings"] == {
        "low_risk_warning_ratio": 8.0,
        "high_risk_warning_ratio": 25.0,
        "only_check_rows_with_value": False,
        "value_filter_field": "数量",
    }

    invalid_response = client.post(
        "/api/experience-warnings/settings",
        json={"settings": {"low_risk_warning_ratio": 30, "high_risk_warning_ratio": 20}},
    )
    assert invalid_response.status_code == 400
    assert invalid_response.json()["detail"] == "高风险预警比率必须大于等于低风险预警比率"


def test_write_report_fills_template_placeholders_and_summary_table(tmp_path, monkeypatch):
    template_path = tmp_path / "控制价报告模板-yyyy-mm-dd.docx"
    template = Document()
    template.add_paragraph("控制价报告模板-yyyy-mm-dd")
    template.add_paragraph("【项目名称】")
    template.add_paragraph("【yyyy】年【mm】月【dd】日")
    table = template.add_table(rows=2, cols=2)
    table.cell(0, 0).text = "合计"
    table.cell(0, 1).text = "【费用汇总-合计（不含税）】"
    table.cell(1, 0).text = "增值税"
    table.cell(1, 1).text = "【费用汇总-增值税】"
    template.add_paragraph("采购计划金额 【采购计划金额】 万元")
    template.add_heading("五、其他需要注意的事项", level=1)
    template.save(template_path)

    workbook = Workbook()
    ws = workbook.active
    ws.title = "费用汇总"
    ws.append(["项目名称：", "铜梁江津-遵义-贵阳项目"])
    ws.append(["合计（含税）", 123456.78])
    ws.append(["合计（不含税）", 116468.66])
    ws.append(["增值税", 6789.01])
    workbook_path = tmp_path / "filled.xlsx"
    workbook.save(workbook_path)

    summary = FillSummary(
        total_data_rows=10,
        price_column="H-基价",
        filled_rows=8,
        matched_rows=8,
        unchanged_rows=1,
        review_rows=2,
        conflict_rows=1,
        output_excel="filled.xlsx",
        output_report="report.docx",
        report_text="输入10行，匹配8行。",
        table_preview={"headers": ["要素1"], "rows": [["测试"]]},
        price_logs=[],
        physical_matched_rows=0,
        physical_experience_rows=6,
        physical_review_rows=4,
        technical_matched_rows=7,
        technical_experience_rows=1,
        technical_review_rows=2,
    )
    output_path = tmp_path / "控制价报告模板-yyyy-mm-dd-处理报告-【codex】.docx"

    monkeypatch.setattr(report_module, "REPORT_TEMPLATE_PATH", template_path)
    report_path = write_report(output_path, "input.xlsx", summary, output_excel_path=workbook_path, report_date=date(2026, 6, 16))

    assert report_path.name == "控制价报告模板-2026-06-16-处理报告-【codex】.docx"
    assert report_path.exists()
    document = Document(report_path)
    text = "\n".join(p.text for p in document.paragraphs)
    table_text = "\n".join(cell.text for table in document.tables for row in table.rows for cell in row.cells)
    assert "yyyy-mm-dd" not in text
    assert "【项目名称】" not in text
    assert "2026年06月16日" in text
    assert "铜梁江津-遵义-贵阳项目" in text
    assert "116468.66" in table_text
    assert "6789.01" in table_text
    assert "采购计划金额 116468.66 万元" in text
    assert "实物工作费调整系数：第一层命中 0 行，第二层经验 6 行，待复核 4 行。" in text
    assert "技术工作费调整系数：第一层命中 7 行，第二层经验 1 行，待复核 2 行。" in text
    report_heading = next(p for p in document.paragraphs if p.text == "造价智算匹配报告")
    assert report_heading.runs[0].font.size.pt == 14
    assert report_heading.runs[0].bold is True


def test_write_report_resolves_fee_summary_formulas_without_cached_values(tmp_path, monkeypatch):
    template_path = tmp_path / "控制价报告模板-yyyy-mm-dd.docx"
    template = Document()
    template.add_paragraph("【项目名称】")
    table = template.add_table(rows=2, cols=2)
    table.cell(0, 0).text = "合计"
    table.cell(0, 1).text = "【费用汇总-合计（不含税）】"
    table.cell(1, 0).text = "增值税"
    table.cell(1, 1).text = "【费用汇总-增值税】"
    template.add_paragraph("采购计划金额 【采购计划金额】 万元")
    template.save(template_path)

    workbook = Workbook()
    detail = workbook.active
    detail.title = "表2-通用工程测量费用"
    detail["M5"] = 8804632.263536
    summary_sheet = workbook.create_sheet("费用汇总")
    summary_sheet.append(["项目名称：", "铜梁江津-遵义-贵阳项目"])
    summary_sheet.append(["合计（含税）", "=ROUND('表2-通用工程测量费用'!M5/10000+39.356,4)"])
    summary_sheet.append(["合计（不含税）", "=ROUND(B2/1.06,4)"])
    summary_sheet.append(["增值税", "=B2-B3"])
    workbook_path = tmp_path / "filled.xlsx"
    workbook.save(workbook_path)
    workbook.close()

    summary = FillSummary(
        total_data_rows=1,
        price_column="H-基价",
        filled_rows=1,
        matched_rows=1,
        unchanged_rows=0,
        review_rows=0,
        conflict_rows=0,
        output_excel="filled.xlsx",
        output_report="report.docx",
        report_text="输入1行，匹配1行。",
        table_preview={"headers": [], "rows": []},
    )
    output_path = tmp_path / "控制价报告模板-yyyy-mm-dd-处理报告-【codex】.docx"

    monkeypatch.setattr(report_module, "REPORT_TEMPLATE_PATH", template_path)
    report_path = write_report(output_path, "input.xlsx", summary, output_excel_path=workbook_path, report_date=date(2026, 6, 20))

    document = Document(report_path)
    table_text = "\n".join(cell.text for table in document.tables for row in table.rows for cell in row.cells)
    text = "\n".join(p.text for p in document.paragraphs)
    assert "867.75" in table_text
    assert "52.07" in table_text
    assert "采购计划金额 867.75 万元" in text


def test_process_endpoint_generates_excel_and_report():
    if INPUT_PATH is None:
        pytest.skip("缺少 100 行空单价输入测试表")
    client = TestClient(app)

    with INPUT_PATH.open("rb") as handle:
        response = client.post(
            "/api/process",
            files={
                "file": (
                    INPUT_PATH.name,
                    handle,
                    "application/vnd.openxmlformats-officedocument.spreadsheetml.sheet",
                )
            },
        )

    assert response.status_code == 200
    payload = response.json()
    assert payload["summary"]["filled_rows"] == 100
    assert payload["summary"]["review_rows"] == 0
    assert payload["summary"]["report_text"] == "输入100行，匹配100行。"
    assert payload["summary"]["table_preview"]["headers"][:7] == [
        "要素1",
        "要素2",
        "要素3",
        "要素4",
        "要素5",
        "单位",
        "基价测试列",
    ]
    assert "匹配状态" in payload["summary"]["table_preview"]["headers"]
    assert "候选数量" in payload["summary"]["table_preview"]["headers"]
    assert "匹配说明" in payload["summary"]["table_preview"]["headers"]
    price_index = payload["summary"]["table_preview"]["headers"].index("基价测试列")
    status_index = payload["summary"]["table_preview"]["headers"].index("匹配状态")
    assert payload["summary"]["table_preview"]["rows"][0][price_index] == 17213
    assert payload["summary"]["table_preview"]["rows"][0][status_index] == "已匹配"

    excel_response = client.get(payload["downloads"]["excel"])
    report_response = client.get(payload["downloads"]["report"])
    assert excel_response.status_code == 200
    assert report_response.status_code == 200


def test_experience_warning_endpoint_writes_warnings_after_process(tmp_path, monkeypatch):
    input_path = tmp_path / "warning-input.xlsx"
    workbook = Workbook()
    sheet = workbook.active
    sheet.title = "表2 测量"
    sheet.append(["要素1", "要素2", "要素3", "要素4", "要素5", "单位", "数量", "基价"])
    sheet.append(["岩土工程勘察", "地质测绘", "", "比例-1:500", "复杂", "km2", 2, "空单价"])
    workbook.save(input_path)
    workbook.close()

    pool_path = tmp_path / "经验池.xlsx"
    pool_workbook = Workbook()
    pool_sheet = pool_workbook.active
    pool_sheet.title = "经验池"
    pool_sheet.append(EXPERIENCE_POOL_HEADERS)
    record = {header: "" for header in EXPERIENCE_POOL_HEADERS}
    record.update(
        {
            "来源文件": "历史控制价.xlsx",
            "来源sheet": "表2 测量",
            "来源行": 88,
            "要素1": "岩土工程勘察",
            "要素2": "地质测绘",
            "要素3": "",
            "要素4": "比例-1:500",
            "要素5": "复杂",
            "单位": "km2",
            "基价": 1,
        }
    )
    pool_sheet.append([record.get(header) for header in EXPERIENCE_POOL_HEADERS])
    pool_workbook.save(pool_path)
    pool_workbook.close()
    monkeypatch.setattr(main_module, "DEFAULT_EXPERIENCE_POOL_PATH", pool_path)
    monkeypatch.setattr(main_module, "LEGACY_EXPERIENCE_POOL_PATH", tmp_path / "missing-legacy.xlsx")

    client = TestClient(app)
    with input_path.open("rb") as handle:
        response = client.post(
            "/api/process",
            data={"only_match_rows_with_value": "false"},
            files={
                "file": (
                    input_path.name,
                    handle,
                    "application/vnd.openxmlformats-officedocument.spreadsheetml.sheet",
                )
            },
        )

    assert response.status_code == 200
    payload = response.json()
    preview_headers = payload["summary"]["table_preview"]["headers"]
    assert payload["summary"]["warning_summary"]["executed"] is False
    assert WARNING_PARAMETER_FIELD not in preview_headers
    assert WARNING_DETAIL_FIELD not in preview_headers

    warning_response = client.post("/api/experience-warnings/run", data={"job_id": payload["job_id"]})
    assert warning_response.status_code == 200
    warning_payload = warning_response.json()
    preview_headers = warning_payload["summary"]["table_preview"]["headers"]
    assert warning_payload["summary"]["warning_summary"]["executed"] is True
    assert warning_payload["summary"]["warning_summary"]["match_mode_counts"]["字段完全匹配"] == 1
    assert warning_payload["summary"]["warning_summary"]["warning_rows"] == 1
    assert warning_payload["summary"]["warning_summary"]["low_rows"] + warning_payload["summary"]["warning_summary"]["high_rows"] == 1
    assert WARNING_PARAMETER_FIELD in preview_headers
    assert WARNING_DETAIL_FIELD in preview_headers
    parameter_index = preview_headers.index(WARNING_PARAMETER_FIELD)
    detail_index = preview_headers.index(WARNING_DETAIL_FIELD)
    assert warning_payload["summary"]["table_preview"]["rows"][0][parameter_index] == "基价"
    assert "匹配模式：字段完全匹配" in warning_payload["summary"]["table_preview"]["rows"][0][detail_index]
    assert "经验池平均值 1" in warning_payload["summary"]["table_preview"]["rows"][0][detail_index]
    assert "低风险阈值 5%" in warning_payload["summary"]["table_preview"]["rows"][0][detail_index]
    assert "历史控制价.xlsx / 表2 测量 第88行 的 基价 为 1" in warning_payload["summary"]["table_preview"]["rows"][0][detail_index]
    progress_response = client.get(f"/api/experience-warnings/progress/{payload['job_id']}")
    assert progress_response.status_code == 200
    assert progress_response.json() == {
        "status": "completed",
        "processed_rows": 1,
        "total_rows": 1,
        "matched_rows": 1,
        "warning_rows": 1,
    }

    excel_response = client.get(warning_payload["downloads"]["excel"])
    output_path = tmp_path / "warning-output.xlsx"
    output_path.write_bytes(excel_response.content)
    output_book = load_workbook(output_path, data_only=True)
    try:
        output_sheet = output_book["表2 测量"]
        headers = [cell.value for cell in output_sheet[1]]
        assert output_sheet.cell(row=2, column=headers.index(WARNING_PARAMETER_FIELD) + 1).value == "基价"
        assert "经验池平均值 1" in output_sheet.cell(
            row=2,
            column=headers.index(WARNING_DETAIL_FIELD) + 1,
        ).value
        assert "历史控制价.xlsx / 表2 测量 第88行 的 基价 为 1" in output_sheet.cell(
            row=2,
            column=headers.index(WARNING_DETAIL_FIELD) + 1,
        ).value
    finally:
        output_book.close()


def test_experience_warning_settings_affect_run_result(tmp_path, monkeypatch):
    kb_path = tmp_path / "kb.xlsx"
    kb_workbook = Workbook()
    kb_sheet = kb_workbook.active
    kb_sheet.append(
        [
            "序号",
            "要素1",
            "要素2",
            "要素3",
            "要素4",
            "要素5",
            "单位",
            "基价",
            "备注",
            "【经验数】实物工作费调整系数",
            "【经验数解释】-实物工作费调整系数",
            "【经验数】技术工作费调整系数",
            "【经验数解释】-技术工作费调整系数",
        ]
    )
    kb_sheet.append([1, "控制测量", "首级控制测量", "", "GPS测量C级", "中等", "个", 100, "", None, "", None, ""])
    kb_workbook.save(kb_path)
    kb_workbook.close()

    input_path = tmp_path / "warning-settings-input.xlsx"
    workbook = Workbook()
    sheet = workbook.active
    sheet.title = "表2 测量"
    sheet.append(["要素1", "要素2", "要素3", "要素4", "要素5", "单位", "数量", "基价"])
    sheet.append(["控制测量", "首级控制测量", "", "GPS测量C级", "中等", "个", 8, "空单价"])
    workbook.save(input_path)
    workbook.close()

    pool_path = tmp_path / "经验池.xlsx"
    pool_workbook = Workbook()
    pool_sheet = pool_workbook.active
    pool_sheet.title = "经验池"
    pool_sheet.append(EXPERIENCE_POOL_HEADERS)
    record = {header: "" for header in EXPERIENCE_POOL_HEADERS}
    record.update(
        {
            "来源文件": "历史控制价.xlsx",
            "来源sheet": "表2 测量",
            "来源行": 18,
            "要素1": "控制测量",
            "要素2": "首级控制测量",
            "要素3": "",
            "要素4": "GPS测量C级",
            "要素5": "中等",
            "单位": "个",
            "基价": 100,
        }
    )
    pool_sheet.append([record.get(header) for header in EXPERIENCE_POOL_HEADERS])
    pool_workbook.save(pool_path)
    pool_workbook.close()

    settings_path = tmp_path / "experience-warning-settings.json"
    settings_path.write_text(
        json.dumps(
            {
                "settings": {
                    "low_risk_warning_ratio": 40,
                    "high_risk_warning_ratio": 60,
                    "only_check_rows_with_value": True,
                    "value_filter_field": "数量",
                }
            },
            ensure_ascii=False,
        ),
        encoding="utf-8",
    )

    monkeypatch.setattr(main_module, "DEFAULT_KB_PATH", kb_path)
    monkeypatch.setattr(main_module, "DEFAULT_EXPERIENCE_POOL_PATH", pool_path)
    monkeypatch.setattr(main_module, "LEGACY_EXPERIENCE_POOL_PATH", tmp_path / "missing-legacy.xlsx")
    monkeypatch.setattr(main_module, "DEFAULT_EXPERIENCE_WARNING_SETTINGS_PATH", settings_path)

    client = TestClient(app)
    with input_path.open("rb") as handle:
        response = client.post(
            "/api/process",
            data={"only_match_rows_with_value": "false"},
            files={
                "file": (
                    input_path.name,
                    handle,
                    "application/vnd.openxmlformats-officedocument.spreadsheetml.sheet",
                )
            },
        )

    assert response.status_code == 200
    payload = response.json()
    warning_response = client.post("/api/experience-warnings/run", data={"job_id": payload["job_id"]})
    assert warning_response.status_code == 200
    warning_payload = warning_response.json()
    assert warning_payload["summary"]["warning_summary"]["warning_rows"] == 0
    assert warning_payload["summary"]["warning_summary"]["low_risk_threshold_percent"] == 40.0
    assert warning_payload["summary"]["warning_summary"]["high_risk_threshold_percent"] == 60.0
    preview_headers = warning_payload["summary"]["table_preview"]["headers"]
    detail_index = preview_headers.index(WARNING_DETAIL_FIELD)
    assert "当前阈值：低风险 40%；高风险 60%" in warning_payload["summary"]["table_preview"]["rows"][0][detail_index]


def test_experience_warning_endpoint_returns_clear_error_when_filter_field_is_missing(tmp_path, monkeypatch):
    kb_path = tmp_path / "kb-filter-missing.xlsx"
    kb_workbook = Workbook()
    kb_sheet = kb_workbook.active
    kb_sheet.append(
        [
            "序号",
            "要素1",
            "要素2",
            "要素3",
            "要素4",
            "要素5",
            "单位",
            "基价",
            "备注",
            "【经验数】实物工作费调整系数",
            "【经验数解释】-实物工作费调整系数",
            "【经验数】技术工作费调整系数",
            "【经验数解释】-技术工作费调整系数",
        ]
    )
    kb_sheet.append([1, "控制测量", "首级控制测量", "", "GPS测量C级", "中等", "个", 100, "", None, "", None, ""])
    kb_workbook.save(kb_path)
    kb_workbook.close()

    input_path = tmp_path / "warning-filter-missing-input.xlsx"
    workbook = Workbook()
    sheet = workbook.active
    sheet.title = "表2 测量"
    sheet.append(["要素1", "要素2", "要素3", "要素4", "要素5", "单位", "基价"])
    sheet.append(["控制测量", "首级控制测量", "", "GPS测量C级", "中等", "个", "空单价"])
    workbook.save(input_path)
    workbook.close()

    pool_path = tmp_path / "warning-filter-missing-pool.xlsx"
    pool_workbook = Workbook()
    pool_sheet = pool_workbook.active
    pool_sheet.title = "经验池"
    pool_sheet.append(EXPERIENCE_POOL_HEADERS)
    record = {header: "" for header in EXPERIENCE_POOL_HEADERS}
    record.update(
        {
            "来源文件": "历史控制价.xlsx",
            "来源sheet": "表2 测量",
            "来源行": 18,
            "要素1": "控制测量",
            "要素2": "首级控制测量",
            "要素3": "",
            "要素4": "GPS测量C级",
            "要素5": "中等",
            "单位": "个",
            "基价": 100,
        }
    )
    pool_sheet.append([record.get(header) for header in EXPERIENCE_POOL_HEADERS])
    pool_workbook.save(pool_path)
    pool_workbook.close()

    settings_path = tmp_path / "experience-warning-settings.json"
    settings_path.write_text(
        json.dumps(
            {
                "settings": {
                    "low_risk_warning_ratio": 5,
                    "high_risk_warning_ratio": 20,
                    "only_check_rows_with_value": True,
                    "value_filter_field": "数量",
                }
            },
            ensure_ascii=False,
        ),
        encoding="utf-8",
    )

    monkeypatch.setattr(main_module, "DEFAULT_KB_PATH", kb_path)
    monkeypatch.setattr(main_module, "DEFAULT_EXPERIENCE_POOL_PATH", pool_path)
    monkeypatch.setattr(main_module, "LEGACY_EXPERIENCE_POOL_PATH", tmp_path / "missing-legacy.xlsx")
    monkeypatch.setattr(main_module, "DEFAULT_EXPERIENCE_WARNING_SETTINGS_PATH", settings_path)

    client = TestClient(app)
    with input_path.open("rb") as handle:
        response = client.post(
            "/api/process",
            data={"only_match_rows_with_value": "false"},
            files={
                "file": (
                    input_path.name,
                    handle,
                    "application/vnd.openxmlformats-officedocument.spreadsheetml.sheet",
                )
            },
        )

    assert response.status_code == 200
    payload = response.json()
    warning_response = client.post("/api/experience-warnings/run", data={"job_id": payload["job_id"]})
    assert warning_response.status_code == 400
    assert warning_response.json()["detail"] == "预警 sheet 表2 测量 未映射过滤字段：数量"


def test_experience_warning_endpoint_writes_no_warning_rows_after_process(tmp_path, monkeypatch):
    kb_path = tmp_path / "kb.xlsx"
    kb_workbook = Workbook()
    kb_sheet = kb_workbook.active
    kb_sheet.append(
        [
            "序号",
            "要素1",
            "要素2",
            "要素3",
            "要素4",
            "要素5",
            "单位",
            "基价",
            "备注",
            "【经验数】实物工作费调整系数",
            "【经验数解释】-实物工作费调整系数",
            "【经验数】技术工作费调整系数",
            "【经验数解释】-技术工作费调整系数",
        ]
    )
    kb_sheet.append([1, "控制测量", "首级控制测量", "", "GPS测量C级", "中等", "个", 4274, "", None, "", None, ""])
    kb_workbook.save(kb_path)
    kb_workbook.close()

    input_path = tmp_path / "warning-none-input.xlsx"
    workbook = Workbook()
    sheet = workbook.active
    sheet.title = "表2 测量"
    sheet.append(["要素1", "要素2", "要素3", "要素4", "要素5", "单位", "数量", "基价", "实物工作费调整系数", "技术工作费调整系数"])
    sheet.append(["控制测量", "首级控制测量", "", "GPS测量C级", "中等", "个", 9, "空单价", 0.6, 0.22])
    workbook.save(input_path)
    workbook.close()

    pool_path = tmp_path / "经验池.xlsx"
    pool_workbook = Workbook()
    pool_sheet = pool_workbook.active
    pool_sheet.title = "经验池"
    pool_sheet.append(EXPERIENCE_POOL_HEADERS)
    record = {header: "" for header in EXPERIENCE_POOL_HEADERS}
    record.update(
        {
            "来源文件": "历史控制价.xlsx",
            "来源sheet": "表2 测量",
            "来源行": 18,
            "要素1": "控制测量",
            "要素2": "首级控制测量",
            "要素3": "",
            "要素4": "GPS测量C级",
            "要素5": "中等",
            "单位": "个",
            "基价": 4274,
            "实物工作费调整系数": 0.6,
            "技术工作费调整系数": 0.22,
        }
    )
    pool_sheet.append([record.get(header) for header in EXPERIENCE_POOL_HEADERS])
    pool_workbook.save(pool_path)
    pool_workbook.close()

    monkeypatch.setattr(main_module, "DEFAULT_KB_PATH", kb_path)
    monkeypatch.setattr(main_module, "DEFAULT_EXPERIENCE_POOL_PATH", pool_path)
    monkeypatch.setattr(main_module, "LEGACY_EXPERIENCE_POOL_PATH", tmp_path / "missing-legacy.xlsx")

    client = TestClient(app)
    with input_path.open("rb") as handle:
        response = client.post(
            "/api/process",
            files={
                "file": (
                    input_path.name,
                    handle,
                    "application/vnd.openxmlformats-officedocument.spreadsheetml.sheet",
                )
            },
        )

    assert response.status_code == 200
    payload = response.json()
    warning_response = client.post("/api/experience-warnings/run", data={"job_id": payload["job_id"]})
    assert warning_response.status_code == 200
    warning_payload = warning_response.json()
    preview_headers = warning_payload["summary"]["table_preview"]["headers"]
    parameter_index = preview_headers.index(WARNING_PARAMETER_FIELD)
    detail_index = preview_headers.index(WARNING_DETAIL_FIELD)
    assert warning_payload["summary"]["warning_summary"]["warning_rows"] == 0
    assert warning_payload["summary"]["table_preview"]["rows"][0][parameter_index] == "无预警"
    assert "已比对参数：" in warning_payload["summary"]["table_preview"]["rows"][0][detail_index]
    assert "基价" in warning_payload["summary"]["table_preview"]["rows"][0][detail_index]
    assert "当前阈值：低风险 5%；高风险 20%" in warning_payload["summary"]["table_preview"]["rows"][0][detail_index]
    assert "结论：未超过阈值，故无预警。" in warning_payload["summary"]["table_preview"]["rows"][0][detail_index]


def test_process_endpoint_accepts_column_mapping_for_shuffled_headers(tmp_path):
    input_path = tmp_path / "api-shuffled.xlsx"
    workbook = Workbook()
    sheet = workbook.active
    sheet.append(["价格输入", "专业", "工作项", "比例尺", "复杂程度", "单位列", "空列"])
    sheet.append(["空单价", "岩土工程勘察", "地质测绘", "比例-1:500", "复杂", "km2", ""])
    workbook.save(input_path)

    client = TestClient(app)
    mapping = {
        "要素1": "专业",
        "要素2": "工作项",
        "要素3": "空列",
        "要素4": "比例尺",
        "要素5": "复杂程度",
        "单位": "单位列",
        "价格列": "价格输入",
    }

    with input_path.open("rb") as handle:
        response = client.post(
            "/api/process",
            data={
                "column_mapping": json.dumps(mapping, ensure_ascii=False),
                "only_match_rows_with_value": "false",
            },
            files={
                "file": (
                    input_path.name,
                    handle,
                    "application/vnd.openxmlformats-officedocument.spreadsheetml.sheet",
                )
            },
        )

    assert response.status_code == 200
    payload = response.json()
    assert payload["summary"]["filled_rows"] == 1
    assert payload["summary"]["review_rows"] == 0
    assert re.fullmatch(r"【输出】-控制价计算表-\d{8}-\d{4}\.xlsx", payload["summary"]["output_excel"])
    assert re.fullmatch(r"【输出】-控制价报告-\d{8}-\d{4}\.docx", payload["summary"]["output_report"])

    excel_response = client.get(payload["downloads"]["excel"])
    assert excel_response.status_code == 200
    output_path = tmp_path / "download.xlsx"
    output_path.write_bytes(excel_response.content)
    ws = load_workbook(output_path, data_only=True).active
    headers = [cell.value for cell in ws[1]]
    assert ws.cell(row=2, column=headers.index("价格输入") + 1).value == 17213
    assert ws.cell(row=2, column=headers.index("匹配状态") + 1).value == "已匹配"


def test_process_endpoint_accepts_header_row_for_mapping_line(tmp_path):
    input_path = tmp_path / "api-header-row.xlsx"
    workbook = Workbook()
    sheet = workbook.active
    sheet.append(["项目名称", "测试"])
    sheet.append(["说明", "非表头"])
    sheet.append(["价格输入", "专业", "工作项", "比例尺", "复杂程度", "单位列", "空列"])
    sheet.append(["空单价", "岩土工程勘察", "地质测绘", "比例-1:500", "复杂", "km2", ""])
    workbook.save(input_path)

    client = TestClient(app)
    mapping = {
        "要素1": "B",
        "要素2": "C",
        "要素3": "G",
        "要素4": "D",
        "要素5": "E",
        "单位": "F",
        "价格列": "A",
    }

    with input_path.open("rb") as handle:
        response = client.post(
            "/api/process",
            data={
                "column_mapping": json.dumps(mapping, ensure_ascii=False),
                "header_row": "3",
                "only_match_rows_with_value": "false",
            },
            files={
                "file": (
                    input_path.name,
                    handle,
                    "application/vnd.openxmlformats-officedocument.spreadsheetml.sheet",
                )
            },
        )

    assert response.status_code == 200
    payload = response.json()
    assert payload["summary"]["total_data_rows"] == 1
    assert payload["summary"]["filled_rows"] == 1

    excel_response = client.get(payload["downloads"]["excel"])
    output_path = tmp_path / "download-header-row.xlsx"
    output_path.write_bytes(excel_response.content)
    ws = load_workbook(output_path, data_only=True).active
    headers = [cell.value for cell in ws[3]]
    assert ws.cell(row=4, column=1).value == 17213
    assert ws.cell(row=4, column=headers.index("匹配状态") + 1).value == "已匹配"


def test_process_endpoint_accepts_empty_element_column_mapping(tmp_path):
    input_path = tmp_path / "api-empty-element.xlsx"
    workbook = Workbook()
    sheet = workbook.active
    sheet.append(["价格输入", "专业", "工作项", "比例尺", "复杂程度", "单位列"])
    sheet.append(["空单价", "岩土工程勘察", "地质测绘", "比例-1:500", "复杂", "km2"])
    workbook.save(input_path)

    client = TestClient(app)
    mapping = {
        "要素1": "B",
        "要素2": "C",
        "要素3": "空元素列",
        "要素4": "D",
        "要素5": "E",
        "单位": "F",
        "价格列": "A",
    }

    with input_path.open("rb") as handle:
        response = client.post(
            "/api/process",
            data={
                "column_mapping": json.dumps(mapping, ensure_ascii=False),
                "only_match_rows_with_value": "false",
            },
            files={
                "file": (
                    input_path.name,
                    handle,
                    "application/vnd.openxmlformats-officedocument.spreadsheetml.sheet",
                )
            },
        )

    assert response.status_code == 200
    payload = response.json()
    assert payload["summary"]["filled_rows"] == 1
    assert payload["summary"]["review_rows"] == 0


def test_process_endpoint_accepts_merged_cell_rule_flags(tmp_path):
    input_path = tmp_path / "api-merge-flags.xlsx"
    workbook = Workbook()
    sheet = workbook.active
    sheet.append(["价格输入", "要素1", "要素2", "要素3", "要素4", "要素5", "单位"])
    sheet.append(["空单价", "岩土工程勘察", "地质测绘", "", "比例-1:500", "复杂", "km2"])
    sheet.append(["空单价", "岩土工程勘察", None, "", "比例-1:500", "复杂", "km2"])
    sheet.merge_cells("C2:C3")
    workbook.save(input_path)

    client = TestClient(app)
    with input_path.open("rb") as handle:
        response = client.post(
            "/api/process",
            data={
                "merge_vertical_cells": "false",
                "merge_horizontal_cells": "true",
                "only_match_rows_with_value": "false",
            },
            files={
                "file": (
                    input_path.name,
                    handle,
                    "application/vnd.openxmlformats-officedocument.spreadsheetml.sheet",
                )
            },
        )

    assert response.status_code == 200
    payload = response.json()
    assert payload["summary"]["filled_rows"] == 1
    assert payload["summary"]["review_rows"] == 1


def test_process_endpoint_accepts_multi_sheet_configs(tmp_path):
    input_path = tmp_path / "api-multi-sheet.xlsx"
    workbook = Workbook()
    sheet = workbook.active
    sheet.title = "表2 测量"
    sheet.append(["单价匹配-测试", "要素1", "要素2", "要素4", "要素5", "单位"])
    sheet.append(["空单价", "岩土工程勘察", "地质测绘", "比例-1:500", "复杂", "km2"])
    skipped = workbook.create_sheet("表3 跳过")
    skipped.append(["单价匹配-测试", "要素1", "要素2", "要素4", "要素5", "单位"])
    skipped.append(["空单价", "岩土工程勘察", "地质测绘", "比例-1:500", "复杂", "km2"])
    workbook.save(input_path)

    client = TestClient(app)
    sheet_configs = [
        {
            "sheet_name": "表2 测量",
            "enabled": True,
            "header_row": 1,
            "column_mapping": {
                "输出-价格列": "A",
                "要素1": "B",
                "要素2": "C",
                "要素3": "空元素列",
                "要素4": "D",
                "要素5": "E",
                "单位": "F",
            },
        },
        {
            "sheet_name": "表3 跳过",
            "enabled": False,
            "header_row": 1,
            "column_mapping": {
                "输出-价格列": "A",
                "要素1": "B",
                "要素2": "C",
                "要素3": "空元素列",
                "要素4": "D",
                "要素5": "E",
                "单位": "F",
            },
        },
    ]

    with input_path.open("rb") as handle:
        response = client.post(
            "/api/process",
            data={
                "sheet_configs": json.dumps(sheet_configs, ensure_ascii=False),
                "only_match_rows_with_value": "false",
            },
            files={
                "file": (
                    input_path.name,
                    handle,
                    "application/vnd.openxmlformats-officedocument.spreadsheetml.sheet",
                )
            },
        )

    assert response.status_code == 200
    payload = response.json()
    assert payload["summary"]["total_data_rows"] == 1
    assert payload["summary"]["filled_rows"] == 1

    excel_response = client.get(payload["downloads"]["excel"])
    output_path = tmp_path / "download-multi-sheet.xlsx"
    output_path.write_bytes(excel_response.content)
    output_book = load_workbook(output_path, data_only=True)
    assert output_book["表2 测量"].cell(row=2, column=1).value == 17213
    assert output_book["表3 跳过"].cell(row=2, column=1).value == "空单价"


def test_llm_config_reads_api_key_from_backend_environment(monkeypatch):
    monkeypatch.setenv("DEEPSEEK_API_KEY", "secret-from-backend-env")

    assert LlmConfig().resolved_api_key() == "secret-from-backend-env"


def test_llm_config_reads_api_key_from_local_env_file(tmp_path, monkeypatch):
    monkeypatch.delenv("DEEPSEEK_API_KEY", raising=False)
    (tmp_path / ".env.local").write_text("DEEPSEEK_API_KEY=secret-from-local-file\n", encoding="utf-8")
    monkeypatch.setattr(llm_module, "PROJECT_ROOT", tmp_path)

    assert LlmConfig().resolved_api_key() == "secret-from-local-file"


def test_llm_config_keeps_siliconflow_api_key_compatibility(tmp_path, monkeypatch):
    monkeypatch.delenv("SILICONFLOW_API_KEY", raising=False)
    (tmp_path / ".env.local").write_text("SILICONFLOW_API_KEY=secret-from-siliconflow\n", encoding="utf-8")
    monkeypatch.setattr(llm_module, "PROJECT_ROOT", tmp_path)

    config = LlmConfig(
        provider="siliconflow",
        model="deepseek-ai/DeepSeek-V4-Flash",
        base_url="https://api.siliconflow.cn/v1/chat/completions",
    )

    assert config.resolved_api_key() == "secret-from-siliconflow"
    assert config.chat_completions_url() == "https://api.siliconflow.cn/v1/chat/completions"


def test_deepseek_default_uses_official_chat_completions_url():
    config = LlmConfig()

    assert config.provider == "deepseek"
    assert config.model == "deepseek-v4-flash"
    assert config.chat_completions_url() == "https://api.deepseek.com/chat/completions"


def test_risk_report_rejects_when_task_is_missing():
    client = TestClient(app)
    response = client.post("/api/risk-report", data={"job_id": "missing"})

    assert response.status_code == 404


def test_risk_report_returns_prompt_debug_without_api_key(tmp_path, monkeypatch):
    import app.main as main_module

    captured = {}

    def fake_call_chat_completion(config, messages):
        captured["model"] = config.model
        captured["messages"] = messages
        return "风险提示正文"

    job_id = "debug-risk"
    job_dir = tmp_path / job_id
    job_dir.mkdir()
    workbook = Workbook()
    workbook.active.append(["项目", "单价"])
    workbook.active.append(["测试", 123])
    workbook.save(job_dir / "样例-填价结果-【codex】.xlsx")
    document = Document()
    document.add_heading("五、其他需要注意的事项", level=1)
    document.save(job_dir / "样例-处理报告-【codex】.docx")
    (job_dir / "样例-处理报告-【codex】.md").write_text("## 五、其他需要注意的事项\n", encoding="utf-8")

    monkeypatch.setenv("DEEPSEEK_API_KEY", "secret-from-backend-env")
    monkeypatch.setattr(main_module, "RUNTIME_DIR", tmp_path)
    monkeypatch.setattr(main_module, "call_chat_completion", fake_call_chat_completion)

    client = TestClient(app)
    response = client.post("/api/risk-report", data={"job_id": job_id, "model": "risk-model"})

    assert response.status_code == 200
    payload = response.json()
    assert payload["risk_report"] == "风险提示正文"
    assert payload["debug"]["model"] == "risk-model"
    assert payload["debug"]["messages"][0]["role"] == "system"
    assert "造价与招标控制价复核助手" in payload["debug"]["messages"][0]["content"]
    assert "【处理摘要】" in payload["debug"]["messages"][1]["content"]
    assert "【费用汇总】" in payload["debug"]["messages"][1]["content"]
    assert "【知识库依据】" in payload["debug"]["messages"][1]["content"]
    assert "不要逐行复述匹配情况" in payload["debug"]["messages"][1]["content"]
    assert "第二层经验提示" in payload["debug"]["messages"][1]["content"] or payload["knowledge_sources"] == []
    assert "knowledge_sources" in payload
    assert payload["debug"]["prompt_markdown"].endswith("提示词-【codex】.md")
    assert Path(payload["debug"]["prompt_markdown"]).exists()
    assert "secret-from-backend-env" not in str(payload["debug"])
    assert captured["messages"] == payload["debug"]["messages"]


def test_llm_chat_endpoint_returns_model_answer(monkeypatch):
    import app.main as main_module

    captured = {}

    def fake_call_chat_completion(config, messages):
        captured["model"] = config.model
        captured["messages"] = messages
        return "模型回答"

    monkeypatch.setenv("DEEPSEEK_API_KEY", "secret-from-backend-env")
    monkeypatch.setattr(main_module, "call_chat_completion", fake_call_chat_completion)

    client = TestClient(app)
    response = client.post("/api/llm-chat", data={"message": "请用一句话介绍管勘智算", "model": "demo-model"})

    assert response.status_code == 200
    payload = response.json()
    assert payload["answer"] == "模型回答"
    assert payload["debug"]["provider"] == "deepseek"
    assert payload["debug"]["model"] == "demo-model"
    assert payload["debug"]["temperature"] == 0.2
    assert payload["debug"]["max_tokens"] == 1800
    assert payload["debug"]["prompt_markdown"].endswith("提示词-【codex】.md")
    assert Path(payload["debug"]["prompt_markdown"]).exists()
    assert payload["debug"]["messages"][0] == {
        "role": "system",
        "content": "你是造价智算本地原型的大模型测试助手，回答应简洁、准确，避免编造未提供的事实。",
    }
    assert payload["debug"]["messages"][-1]["content"] == "请用一句话介绍管勘智算"
    assert "secret-from-backend-env" not in str(payload["debug"])
    assert captured["model"] == "demo-model"
    assert captured["messages"][-1]["content"] == "请用一句话介绍管勘智算"


def test_llm_chat_endpoint_rejects_blank_message():
    client = TestClient(app)
    response = client.post("/api/llm-chat", data={"message": "   "})

    assert response.status_code == 400


def test_knowledge_search_finds_technical_fee_022(tmp_path, monkeypatch):
    import app.main as main_module

    monkeypatch.setattr(main_module, "DEFAULT_INDEX_PATH", tmp_path / "knowledge-index.json", raising=False)
    client = TestClient(app)
    response = client.post("/api/knowledge/search", json={"question": "0.22 是哪来的？", "limit": 5})

    assert response.status_code == 200
    payload = response.json()
    assert payload["evidence_found"] is True
    joined = "\n".join(result["snippet"] for result in payload["results"])
    assert "技术工作费" in joined
    assert "0.22" in joined or "22%" in joined


def test_knowledge_search_includes_original_standard_materials(tmp_path, monkeypatch):
    import app.main as main_module

    monkeypatch.setattr(main_module, "DEFAULT_INDEX_PATH", tmp_path / "knowledge-index.json", raising=False)
    client = TestClient(app)
    response = client.post("/api/knowledge/search", json={"question": "计价格2002工程勘察设计收费标准使用手册", "limit": 5})

    assert response.status_code == 200
    payload = response.json()
    assert payload["evidence_found"] is True
    assert any(result["source_type"] == "standard" for result in payload["results"])
    assert any("01-原始资料" in result["source_file"] for result in payload["results"])


def test_knowledge_search_prioritizes_price_database_for_specific_price_question(tmp_path, monkeypatch):
    import app.main as main_module

    monkeypatch.setattr(main_module, "DEFAULT_INDEX_PATH", tmp_path / "knowledge-index.json", raising=False)
    client = TestClient(app)
    response = client.post(
        "/api/knowledge/search",
        json={
            "question": "@知识库，地形图测绘（地形测量） 山岭隧道（洞身） 复杂 1:2000 一般单价多少？",
            "limit": 5,
        },
    )

    assert response.status_code == 200
    payload = response.json()
    assert payload["forced_knowledge"] is True
    assert payload["evidence_found"] is True
    first = payload["results"][0]
    assert "03-知识库-二维数据库制作/【数据库】【导入】.xlsx" in first["source_file"]
    assert "序号：118" in first["snippet"]
    assert "基价：14244" in first["snippet"]
    assert "要素4：复杂" in first["snippet"]
    assert "比例-1:2000" in first["snippet"]


def test_knowledge_search_keeps_space_separated_price_terms_for_gps_question(tmp_path, monkeypatch):
    import app.main as main_module

    monkeypatch.setattr(main_module, "DEFAULT_INDEX_PATH", tmp_path / "knowledge-index.json", raising=False)
    client = TestClient(app)
    response = client.post(
        "/api/knowledge/search",
        json={
            "question": "@知识库，控制测量 首级控制测量 GPS测量E级 中等 多少钱？",
            "limit": 5,
        },
    )

    assert response.status_code == 200
    payload = response.json()
    assert payload["forced_knowledge"] is True
    assert payload["evidence_found"] is True
    first = payload["results"][0]
    assert "03-知识库-二维数据库制作/【数据库】【导入】.xlsx" in first["source_file"]
    assert "序号：2" in first["snippet"]
    assert "要素1：控制测量" in first["snippet"]
    assert "要素2：首级控制测量" in first["snippet"]
    assert "GPS测量E级" in first["snippet"]
    assert "要素5：中等" in first["snippet"]
    assert "基价：3203" in first["snippet"]


def test_knowledge_search_returns_price_database_candidates_for_vague_price_question(tmp_path, monkeypatch):
    import app.main as main_module

    monkeypatch.setattr(main_module, "DEFAULT_INDEX_PATH", tmp_path / "knowledge-index.json", raising=False)
    client = TestClient(app)
    response = client.post(
        "/api/knowledge/search",
        json={"question": "@知识库，控制测量 GPS E级 多少钱？", "limit": 6},
    )

    assert response.status_code == 200
    payload = response.json()
    assert payload["forced_knowledge"] is True
    assert payload["evidence_found"] is True
    database_results = [
        result
        for result in payload["results"]
        if "03-知识库-二维数据库制作/【数据库】【导入】.xlsx" in result["source_file"]
    ]
    assert len(database_results) >= 3
    joined = "\n".join(result["snippet"] for result in database_results)
    assert "GPS测量E级" in joined
    assert "首级控制测量" in joined
    assert "基价：2821" in joined
    assert "基价：4123" in joined


def test_knowledge_search_finds_factor_combination_rule():
    client = TestClient(app)
    response = client.post("/api/knowledge/search", json={"question": "附加调整系数为什么不能连乘？", "limit": 5})

    assert response.status_code == 200
    payload = response.json()
    assert payload["evidence_found"] is True
    joined = "\n".join(result["snippet"] for result in payload["results"])
    assert "附加调整系数" in joined
    assert "连乘" in joined or "相加" in joined or "总则" in joined


def test_knowledge_search_understands_physical_factor_shorthand_with_force_prefix_comma(tmp_path, monkeypatch):
    import app.main as main_module

    monkeypatch.setattr(main_module, "DEFAULT_INDEX_PATH", tmp_path / "knowledge-index.json", raising=False)
    client = TestClient(app)
    response = client.post("/api/knowledge/search", json={"question": "@知识库，实物工作系数如何确定", "limit": 5})

    assert response.status_code == 200
    payload = response.json()
    assert payload["forced_knowledge"] is True
    assert payload["evidence_found"] is True
    assert payload["results"][0]["module"] == "实物工作费调整系数"
    joined = "\n".join(result["snippet"] for result in payload["results"])
    assert "实物工作费" in joined
    assert "附加调整系数" in joined or "调整系数" in joined


def test_knowledge_search_finds_second_layer_experience_rule():
    client = TestClient(app)
    response = client.post("/api/knowledge/search", json={"question": "第二层经验提示是什么意思？", "limit": 5})

    assert response.status_code == 200
    payload = response.json()
    assert payload["evidence_found"] is True
    joined = "\n".join(result["snippet"] for result in payload["results"])
    assert "第二层" in joined
    assert "经验提示" in joined or "经验数" in joined


def test_knowledge_search_finds_risk_report_generation_card(tmp_path, monkeypatch):
    import app.main as main_module

    monkeypatch.setattr(main_module, "DEFAULT_INDEX_PATH", tmp_path / "knowledge-index.json", raising=False)
    client = TestClient(app)
    response = client.post("/api/knowledge/search", json={"question": "输出风险报告怎么生成？", "limit": 5})

    assert response.status_code == 200
    payload = response.json()
    assert payload["evidence_found"] is True
    joined = "\n".join(result["snippet"] for result in payload["results"])
    assert "输出风险报告" in joined or "风险报告" in joined
    assert "问问智算" in joined
    assert "Word" in joined or "审查" in joined


def test_knowledge_ask_returns_no_evidence_without_calling_model(monkeypatch):
    import app.main as main_module

    def fail_call_chat_completion(config, messages):
        raise AssertionError("没有证据时不应调用大模型")

    monkeypatch.setattr(main_module, "call_chat_completion", fail_call_chat_completion)
    client = TestClient(app)
    response = client.post("/api/knowledge/ask", json={"question": "火星土豆怎么收费？"})

    assert response.status_code == 200
    payload = response.json()
    assert payload["evidence_found"] is False
    assert payload["sources"] == []
    assert payload["answer"] == "当前知识库未找到明确依据，需要人工复核。"


def test_knowledge_ask_uses_evidence_bounded_prompt(monkeypatch):
    import app.main as main_module

    captured = {}

    def fake_call_chat_completion(config, messages):
        captured["messages"] = messages
        return "智算解释：0.22 来自技术工作费依据。\n\n依据来源：已列出。"

    monkeypatch.setenv("DEEPSEEK_API_KEY", "secret-from-backend-env")
    monkeypatch.setattr(main_module, "call_chat_completion", fake_call_chat_completion)

    client = TestClient(app)
    response = client.post(
        "/api/knowledge/ask",
        json={
            "question": "0.22 是哪来的？",
            "row_context": {
                "sheet_name": "表2 测量",
                "row_number": 9,
                "values": {"技术工作费调整系数": "0.22", "匹配状态": "已匹配"},
            },
            "model": "demo-model",
        },
    )

    assert response.status_code == 200
    payload = response.json()
    assert payload["evidence_found"] is True
    assert payload["answer"].startswith("智算解释")
    assert payload["sources"]
    assert payload["debug"]["model"] == "demo-model"
    system_prompt = captured["messages"][0]["content"]
    user_prompt = captured["messages"][1]["content"]
    assert "只能基于【已检索资料】和【当前行上下文】回答" in system_prompt
    assert "不得直接裁决基价、实物工作费调整系数、技术工作费调整系数" in system_prompt
    assert "表2 测量" in user_prompt
    assert "0.22 是哪来的？" in user_prompt
    assert "secret-from-backend-env" not in str(payload["debug"])


def test_knowledge_ask_strips_force_knowledge_prefix(monkeypatch):
    import app.main as main_module

    captured = {}

    def fake_call_chat_completion(config, messages):
        captured["messages"] = messages
        return "智算解释：已按知识库依据回答。\n\n依据来源：已列出。"

    monkeypatch.setenv("DEEPSEEK_API_KEY", "secret-from-backend-env")
    monkeypatch.setattr(main_module, "call_chat_completion", fake_call_chat_completion)

    client = TestClient(app)
    response = client.post("/api/knowledge/ask", json={"question": "查库：0.22 是哪来的？", "model": "demo-model"})

    assert response.status_code == 200
    payload = response.json()
    assert payload["forced_knowledge"] is True
    assert payload["evidence_found"] is True
    assert payload["sources"]
    user_prompt = captured["messages"][1]["content"]
    assert "0.22 是哪来的？" in user_prompt
    assert "查库：" not in user_prompt


def test_llm_chat_force_knowledge_prefix_uses_knowledge_path(monkeypatch):
    import app.main as main_module

    captured = {}

    def fake_call_chat_completion(config, messages):
        captured["messages"] = messages
        return "智算解释：普通问答入口已切换到知识库依据回答。\n\n依据来源：已列出。"

    monkeypatch.setenv("DEEPSEEK_API_KEY", "secret-from-backend-env")
    monkeypatch.setattr(main_module, "call_chat_completion", fake_call_chat_completion)

    client = TestClient(app)
    response = client.post("/api/llm-chat", data={"message": "@知识库 0.22 是哪来的？", "model": "demo-model"})

    assert response.status_code == 200
    payload = response.json()
    assert payload["forced_knowledge"] is True
    assert payload["evidence_found"] is True
    assert payload["sources"]
    assert payload["debug"]["model"] == "demo-model"
    system_prompt = captured["messages"][0]["content"]
    user_prompt = captured["messages"][1]["content"]
    assert "只能基于【已检索资料】和【当前行上下文】回答" in system_prompt
    assert "0.22 是哪来的？" in user_prompt
    assert "@知识库" not in user_prompt



