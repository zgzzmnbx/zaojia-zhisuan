from __future__ import annotations

import json
import re
import sys
from itertools import zip_longest
from pathlib import Path
from typing import Any

from docx import Document
from fastapi.testclient import TestClient
from openpyxl import load_workbook


PROJECT_ROOT = Path(__file__).resolve().parents[1]
BACKEND_ROOT = PROJECT_ROOT / "backend"
if str(BACKEND_ROOT) not in sys.path:
    sys.path.insert(0, str(BACKEND_ROOT))

from app import main as main_module  # noqa: E402


BASELINE_DIR = PROJECT_ROOT / "08-测试与演示" / "测试-V3.1-260802"
INPUT_PATH = BASELINE_DIR / "【项目例子】【测试输入】可行性研究勘察测量控制价计算 -v3.1【批注-完备】.xlsx"
ANSWER_EXCEL = BASELINE_DIR / "【答案-对应版本v3.1】【输出】-控制价计算表-20260802-0123-隐藏空行.xlsx"
ANSWER_WORD = BASELINE_DIR / "【答案-对应版本v3.1】【输出】-控制价报告-.docx"
EVIDENCE_DIR = PROJECT_ROOT / "Codex-Temp" / "acceptance" / "task-p0-v31"
RUNTIME_DIR = EVIDENCE_DIR / "runtime"
TARGET_HEADERS = ("基价", "单价", "实物工作费调整系数", "技术工作费调整系数")
STATUS_HEADERS = ("匹配状态", "输出-匹配状态", "候选数量", "输出-候选数量")


def _compact(value: object) -> str:
    return re.sub(r"\s+", "", str(value or "")).replace("（", "(").replace("）", ")")


def _header_map(sheet: Any) -> tuple[int, dict[str, int]]:
    best: tuple[int, dict[str, int]] = (0, {})
    wanted = set(TARGET_HEADERS + STATUS_HEADERS)
    for row_number in range(1, min(sheet.max_row, 20) + 1):
        mapping: dict[str, int] = {}
        for column_number in range(1, sheet.max_column + 1):
            text = _compact(sheet.cell(row_number, column_number).value)
            for header in wanted:
                if text == _compact(header):
                    mapping[header] = column_number
        if len(mapping) > len(best[1]):
            best = (row_number, mapping)
    return best


def _fill_signature(cell: Any) -> tuple[object, ...]:
    color = cell.fill.fgColor
    return (
        cell.fill.fill_type,
        color.type,
        color.rgb,
        color.indexed,
        color.theme,
        color.tint,
    )


def compare_excel(actual_path: Path, answer_path: Path) -> dict[str, object]:
    actual = load_workbook(actual_path, data_only=False)
    answer = load_workbook(answer_path, data_only=False)
    try:
        common_sheets = [name for name in answer.sheetnames if name in actual.sheetnames]
        missing_sheets = [name for name in answer.sheetnames if name not in actual.sheetnames]
        target_differences: list[dict[str, object]] = []
        formula_differences: list[dict[str, object]] = []
        fill_differences: list[dict[str, object]] = []
        status_differences: list[dict[str, object]] = []
        hidden_row_differences: list[dict[str, object]] = []
        full_value_difference_count = 0
        compared_target_cells = 0
        for sheet_name in common_sheets:
            actual_sheet = actual[sheet_name]
            answer_sheet = answer[sheet_name]
            answer_header_row, answer_headers = _header_map(answer_sheet)
            actual_header_row, actual_headers = _header_map(actual_sheet)
            max_row = max(answer_sheet.max_row, actual_sheet.max_row)
            max_column = max(answer_sheet.max_column, actual_sheet.max_column)
            for row_number in range(1, max_row + 1):
                if bool(actual_sheet.row_dimensions[row_number].hidden) != bool(answer_sheet.row_dimensions[row_number].hidden):
                    hidden_row_differences.append({"sheet": sheet_name, "row": row_number})
                for column_number in range(1, max_column + 1):
                    actual_value = actual_sheet.cell(row_number, column_number).value
                    answer_value = answer_sheet.cell(row_number, column_number).value
                    if actual_value != answer_value:
                        full_value_difference_count += 1
                    if (
                        isinstance(actual_value, str) and actual_value.startswith("=")
                    ) or (
                        isinstance(answer_value, str) and answer_value.startswith("=")
                    ):
                        if actual_value != answer_value:
                            formula_differences.append({
                                "sheet": sheet_name, "cell": actual_sheet.cell(row_number, column_number).coordinate,
                                "actual": actual_value, "answer": answer_value,
                            })
            if not answer_header_row or not actual_header_row:
                continue
            for header in TARGET_HEADERS:
                if header not in answer_headers or header not in actual_headers:
                    continue
                for row_number in range(answer_header_row + 1, max_row + 1):
                    answer_cell = answer_sheet.cell(row_number, answer_headers[header])
                    actual_cell = actual_sheet.cell(row_number, actual_headers[header])
                    compared_target_cells += 1
                    if actual_cell.value != answer_cell.value:
                        target_differences.append({
                            "sheet": sheet_name, "row": row_number, "field": header,
                            "actual": actual_cell.value, "answer": answer_cell.value,
                        })
                    if _fill_signature(actual_cell) != _fill_signature(answer_cell):
                        fill_differences.append({"sheet": sheet_name, "row": row_number, "field": header})
            for header in STATUS_HEADERS:
                if header not in answer_headers or header not in actual_headers:
                    continue
                for row_number in range(answer_header_row + 1, max_row + 1):
                    answer_value = answer_sheet.cell(row_number, answer_headers[header]).value
                    actual_value = actual_sheet.cell(row_number, actual_headers[header]).value
                    if actual_value != answer_value:
                        status_differences.append({
                            "sheet": sheet_name, "row": row_number, "field": header,
                            "actual": actual_value, "answer": answer_value,
                        })
        return {
            "common_sheets": common_sheets,
            "missing_sheets": missing_sheets,
            "compared_target_cells": compared_target_cells,
            "three_number_difference_count": len(target_differences),
            "three_number_differences": target_differences[:50],
            "formula_difference_count": len(formula_differences),
            "formula_differences": formula_differences[:50],
            "target_fill_difference_count": len(fill_differences),
            "target_fill_differences": fill_differences[:50],
            "review_status_difference_count": len(status_differences),
            "review_status_differences": status_differences[:50],
            "hidden_row_difference_count": len(hidden_row_differences),
            "hidden_row_differences": hidden_row_differences[:50],
            "full_value_difference_count": full_value_difference_count,
        }
    finally:
        actual.close()
        answer.close()


def _word_lines(path: Path) -> list[str]:
    document = Document(path)
    lines: list[str] = []

    def add_paragraphs(paragraphs: Any) -> None:
        for paragraph in paragraphs:
            text = re.sub(r"\s+", " ", paragraph.text).strip()
            if text:
                lines.append(text)

    add_paragraphs(document.paragraphs)
    for table in document.tables:
        for row in table.rows:
            lines.append(" | ".join(re.sub(r"\s+", " ", cell.text).strip() for cell in row.cells))
    for section in document.sections:
        add_paragraphs(section.header.paragraphs)
        add_paragraphs(section.footer.paragraphs)
    return lines


def compare_word(actual_path: Path, answer_path: Path) -> dict[str, object]:
    actual_lines = _word_lines(actual_path)
    answer_lines = _word_lines(answer_path)
    differences = [
        {"index": index, "actual": actual, "answer": answer}
        for index, (actual, answer) in enumerate(zip_longest(actual_lines, answer_lines, fillvalue=""), start=1)
        if actual != answer
    ]
    anchor = "造价智算匹配报告"
    actual_core_index = max((index for index, line in enumerate(actual_lines) if line == anchor), default=-1)
    answer_core_index = max((index for index, line in enumerate(answer_lines) if line == anchor), default=-1)
    actual_appendix_index = next((index for index, line in enumerate(actual_lines) if line == "大模型风险提示"), actual_core_index)
    answer_appendix_index = next((index for index, line in enumerate(answer_lines) if line == "大模型风险提示"), answer_core_index)

    def normalized_cover(lines: list[str], end: int) -> list[str]:
        return [
            re.sub(r"^\d{4}年\d{2}月\d{2}日$", "[报告日期]", line)
            for line in lines[:max(0, end)]
        ]

    actual_cover = normalized_cover(actual_lines, actual_appendix_index)
    answer_cover = normalized_cover(answer_lines, answer_appendix_index)
    cover_differences = [
        {"index": index, "actual": actual, "answer": answer}
        for index, (actual, answer) in enumerate(zip_longest(actual_cover, answer_cover, fillvalue=""), start=1)
        if actual != answer
    ]
    actual_core = actual_lines[actual_core_index:] if actual_core_index >= 0 else []
    answer_core = answer_lines[answer_core_index:] if answer_core_index >= 0 else []
    core_differences = [
        {"index": index, "actual": actual, "answer": answer}
        for index, (actual, answer) in enumerate(zip_longest(actual_core, answer_core, fillvalue=""), start=1)
        if actual != answer
    ]
    actual_appendix = actual_lines[actual_appendix_index:actual_core_index] if 0 <= actual_appendix_index < actual_core_index else []
    answer_appendix = answer_lines[answer_appendix_index:answer_core_index] if 0 <= answer_appendix_index < answer_core_index else []
    appendix_text = "\n".join(actual_appendix)
    required_risk_facts = ["500000", "4274", "1.3", "1.15", "第二层经验"]
    missing_risk_facts = [fact for fact in required_risk_facts if fact not in appendix_text]
    return {
        "actual_line_count": len(actual_lines),
        "answer_line_count": len(answer_lines),
        "difference_count": len(differences),
        "differences": differences[:50],
        "cover_difference_count_excluding_date": len(cover_differences),
        "cover_differences": cover_differences[:20],
        "business_core_difference_count": len(core_differences),
        "business_core_differences": core_differences[:20],
        "actual_llm_appendix_line_count": len(actual_appendix),
        "answer_llm_appendix_line_count": len(answer_appendix),
        "missing_llm_risk_facts": missing_risk_facts,
        "business_content_passed": (
            not cover_differences
            and not core_differences
            and bool(actual_appendix)
            and not missing_risk_facts
        ),
    }


def _require(response: Any, label: str) -> dict[str, Any]:
    if response.status_code != 200:
        raise RuntimeError(f"{label}失败：HTTP {response.status_code} {response.text[:500]}")
    return response.json()


def run_acceptance() -> dict[str, object]:
    for path in (INPUT_PATH, ANSWER_EXCEL, ANSWER_WORD):
        if not path.is_file():
            raise FileNotFoundError(path)
    EVIDENCE_DIR.mkdir(parents=True, exist_ok=True)
    RUNTIME_DIR.mkdir(parents=True, exist_ok=True)
    main_module.RUNTIME_DIR = RUNTIME_DIR
    client = TestClient(main_module.app)

    with INPUT_PATH.open("rb") as handle:
        inspect = _require(client.post(
            "/api/inspect",
            files={"file": (INPUT_PATH.name, handle, "application/vnd.openxmlformats-officedocument.spreadsheetml.sheet")},
        ), "结构识别")
    sheet_configs = [
        {
            "sheet_name": sheet["sheet_name"],
            "enabled": bool(sheet.get("enabled", True)),
            "header_row": int(sheet["header_row"]),
            "column_mapping": sheet["suggested_mapping"],
            "output_match_report": True,
            "merge_vertical_cells": True,
            "merge_horizontal_cells": True,
            "only_match_rows_with_value": True,
            "match_value_filter_field": "数量",
        }
        for sheet in inspect.get("sheets", [])
    ]
    with INPUT_PATH.open("rb") as handle:
        process = _require(client.post(
            "/api/process",
            data={
                "project_name": "V3.1 Task P0 正式业务回归",
                "source_type": "web",
                "sheet_configs": json.dumps(sheet_configs, ensure_ascii=False),
                "only_match_rows_with_value": "true",
                "match_value_filter_field": "数量",
            },
            files={"file": (INPUT_PATH.name, handle, "application/vnd.openxmlformats-officedocument.spreadsheetml.sheet")},
        ), "正式处理")
    job_id = str(process["job_id"])
    task_id = str((process.get("task_tracking") or {}).get("task", {}).get("task_id") or "")
    if not task_id:
        raise RuntimeError("正式处理未返回业务 Task")

    risk = _require(client.post("/api/experience-warnings/run", data={"job_id": job_id}), "风险检查")
    llm_risk_response = client.post("/api/risk-report", data={"job_id": job_id})
    llm_risk_result = (
        {"status": "completed", "http_status": llm_risk_response.status_code}
        if llm_risk_response.status_code == 200
        else {
            "status": "failed",
            "http_status": llm_risk_response.status_code,
            "detail": str((llm_risk_response.json() if llm_risk_response.headers.get("content-type", "").startswith("application/json") else {}).get("detail") or "风险报告生成失败")[:300],
        }
    )
    excel_response = client.get(
        f"/api/download/{job_id}/excel",
        params={"hide_empty_rows": "true", "value_filter_field": "数量"},
    )
    word_response = client.get(f"/api/download/{job_id}/report")
    if excel_response.status_code != 200 or word_response.status_code != 200:
        raise RuntimeError("正式成果下载失败")
    actual_excel = EVIDENCE_DIR / "V3.1-Task-P0-实际输出.xlsx"
    actual_word = EVIDENCE_DIR / "V3.1-Task-P0-实际输出.docx"
    actual_excel.write_bytes(excel_response.content)
    actual_word.write_bytes(word_response.content)

    task = _require(client.get(f"/api/tasks/{task_id}"), "Task 详情")
    timeline = _require(client.get(f"/api/tasks/{task_id}/timeline"), "Task 时间线")
    project_id = str((process.get("project_tracking") or {}).get("project_id") or "")
    project_tasks = _require(client.get(f"/api/projects/{project_id}/tasks"), "项目 Task") if project_id else {"items": []}
    excel_comparison = compare_excel(actual_excel, ANSWER_EXCEL)
    word_comparison = compare_word(actual_word, ANSWER_WORD)
    real_events = [item for item in timeline["items"] if not item.get("is_placeholder")]
    report = {
        "status": "passed" if (
            excel_comparison["three_number_difference_count"] == 0
            and excel_comparison["formula_difference_count"] == 0
            and excel_comparison["target_fill_difference_count"] == 0
            and excel_comparison["review_status_difference_count"] == 0
            and word_comparison["business_content_passed"] is True
        ) else "differences_found",
        "input": str(INPUT_PATH.relative_to(PROJECT_ROOT)),
        "answers": [str(ANSWER_EXCEL.relative_to(PROJECT_ROOT)), str(ANSWER_WORD.relative_to(PROJECT_ROOT))],
        "outputs": [str(actual_excel.relative_to(PROJECT_ROOT)), str(actual_word.relative_to(PROJECT_ROOT))],
        "job_id": job_id,
        "task_id": task_id,
        "project_id": project_id,
        "run_id": str((process.get("project_tracking") or {}).get("run_id") or ""),
        "skill": task.get("skill_snapshot"),
        "task_status": task.get("status"),
        "task_stage": task.get("stage"),
        "artifact_version": task.get("artifact_version"),
        "review_round": task.get("review_round"),
        "project_task_count": len(project_tasks.get("items") or []),
        "timeline": {
            "actual_event_count": timeline.get("actual_event_count"),
            "real_events": [
                {"event_type": item["event_type"], "status": item["status"], "source_module": item["source_module"]}
                for item in real_events
            ],
            "placeholders": [
                {"event_type": item["event_type"], "status": item["status"]}
                for item in timeline["items"] if item.get("is_placeholder")
            ],
        },
        "risk_summary": risk["summary"].get("warning_summary"),
        "llm_risk_report": llm_risk_result,
        "excel_comparison": excel_comparison,
        "word_comparison": word_comparison,
    }
    report_path = EVIDENCE_DIR / "V3.1-Task-P0-验收结果.json"
    report_path.write_text(json.dumps(report, ensure_ascii=False, indent=2), encoding="utf-8")
    return report


if __name__ == "__main__":
    result = run_acceptance()
    print(json.dumps(result, ensure_ascii=False, indent=2))
    raise SystemExit(0 if result["status"] == "passed" else 1)
